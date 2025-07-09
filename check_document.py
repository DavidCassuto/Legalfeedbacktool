import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
import re
import os
import google.generativeai as genai
import textwrap
from docx.opc.exceptions import PackageNotFoundError
import datetime
from docx.oxml.ns import qn # qn is from docx.oxml.ns
from docx.oxml import OxmlElement # OxmlElement is from docx.oxml (re-exports it from lxml.oxml)
from docx.text.run import Run
from docx.enum.text import WD_COLOR_INDEX

# --- Configure the Gemini API ---
# SECURITY: For production applications, retrieve the API key from environment variables or a secure configuration file!
# Hardcoding as below is ONLY for this example.
# >>> REPLACE THE EMPTY STRING BELOW WITH YOUR OWN GEMINI API KEY <<<
GEMINI_API_KEY = "AIzaSyBvLmbtZ5DDRg3j5OYCAQhx6-ChOp5TJaU" # <--- ENTER YOUR GEMINI API KEY HERE BETWEEN THE QUOTES

# Initialize model to None before the try block
model = None

if GEMINI_API_KEY == "":
    print("WARNING: Your Gemini API key is not set. AI analysis will be skipped.")
    # model remains None
else:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        # Choose a model. 'gemini-1.5-flash-latest' is often cost-effective and fast.
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        print("INFO: Gemini API configured and model loaded.")
    except Exception as e:
        print(f"ERROR: While configuring Gemini API or loading the model: {e}")
        print("Please ensure your API key is correct, the 'google-generativeai' library is installed, and your internet connection is working.")
        model = None # Ensure model is None upon error


# --- CHECK SETTINGS (Hardcoded for now - later from web UI) ---
# Define here which checks are active in which sections.
# Keys: Names of the checks (must match the logic below)
# Values: List of section number patterns (strings or regex) or 'all'
CHECK_SETTINGS = {
    'title_rule': ['all'], # Title check
    'legal_citations': ['all'], # Legal text check
    'first_ik': ['all'], # 'Ik' (I/me) check
    'deelvragen_structure': ['all'], # Sub-questions structure check (looks for "Deelvragen" heading, so 'all' is fine)
    'section_1_3_intro': ['1.3'], # Introductory sentences in 1.3
    'section_1_4_content': ['1.4'], # Objective/Output/Outcome in 1.4
    'paragraph_length': ['all'], # Paragraph length check (internally excluded for 1.3)
    'paragraph_separation': ['all'], # Paragraph separation check
    # AI Checks - these are already section-specific in their implementation,
    # but we can decide here whether they are executed at all.
    'ai_1_1_1_2': ['1.1', '1.2'],
    'ai_1_3_questions': ['1.3'],
    'ai_chapter_2_analysis': ['2'], # New check for Chapter 2
    # General AI analysis per chunk
    'ai_general_analysis': ['all'] # Can also be specific sections, e.g., ['2', '3', '4']
}

# --- Helper function to check if a check is active for a section ---
def is_check_active_for_section(check_name, current_section):
    """Checks if a check is active for the current section based on CHECK_SETTINGS."""
    if check_name not in CHECK_SETTINGS:
        return False # Check not defined, default off

    allowed_sections = CHECK_SETTINGS[check_name]

    if 'all' in allowed_sections:
        return True # Active for all sections

    if current_section is None:
        return False # Not in a numbered section and not 'all'

    # Check if the current section matches one of the allowed patterns
    for allowed_pattern in allowed_sections:
        # Try as exact match
        if current_section == allowed_pattern:
            return True
        # Try as regex match (e.g., '^1\.\d+$' for all 1.x subsections)
        try:
            if re.match(allowed_pattern, current_section):
                 return True
        except re.error:
            # Invalid regex pattern in settings
            print(f"Warning: Invalid regex pattern '{allowed_pattern}' for check '{check_name}' in CHECK_SETTINGS.")
            pass # Continue with other patterns

    return False # No match found


# --- Function to find the first non-empty paragraph ---
def find_first_non_empty_paragraph(document, start_index=0):
    """Finds the index and object of the first non-empty paragraph from start_index."""
    for idx in range(start_index, len(document.paragraphs)):
        para = document.paragraphs[idx]
        if para.text.strip():
            return para, idx
    return None, -1

# --- Function to find a paragraph based on (part of) the text ---
def find_paragraph_by_text(document, search_text):
    """Finds the first paragraph containing the specified text and returns its index."""
    search_text_stripped_lower = search_text.strip().lower()
    for idx, para in enumerate(document.paragraphs):
        # Check if the paragraph text contains the search text (case-insensitive)
        if search_text_stripped_lower in para.text.strip().lower():
            return para, idx
    return None, -1

# --- Function to find a paragraph based on a regular expression ---
def find_paragraph_by_regex(document, pattern):
    """Finds the first paragraph matching the regex and returns its index."""
    compiled_pattern = re.compile(pattern)
    for idx, para in enumerate(document.paragraphs):
        if compiled_pattern.search(para.text):
            return para, idx
    return None, -1

# Helper function to find the index of the next numbered heading
def find_next_numbered_heading(start_index, document):
    """Finds the index of the next paragraph styled as a Heading with a number."""
    # Updated regex to be more flexible with spacing and punctuation after the number
    heading_pattern = re.compile(r'^\s*(\d+(\.\d+)*)[\s\.\:]?', re.IGNORECASE)
    for idx in range(start_index + 1, len(document.paragraphs)):
        para = document.paragraphs[idx]
        # Check if style is a Heading and text starts with a number pattern (e.g., "2 ", "3.1 ")
        if para.style.name.startswith('Heading') and heading_pattern.search(para.text.strip()):
            print(f"DEBUG: Potential next numbered heading found at index {idx}: '{para.text.strip()[:50]}'") # Added debug
            return idx
    print(f"DEBUG: No next numbered heading found after index {start_index}.") # Added debug
    return len(document.paragraphs) # Use document.paragraphs


# --- Function to check if a paragraph is a "real paragraph" (for separation check) ---
def is_real_paragraph_for_separation(para):
    """Checks if a paragraph is a 'real paragraph' (not empty, > 60 characters, not a heading, not a list, no special style, no comment)."""
    text_content = para.text.strip()
    style_name = para.style.name if para.style else "None"

    if not para or not text_content:
        return False
    # Updated: Also exclude paragraphs matching the heading detection pattern
    heading_detection_pattern = re.compile(r'^\s*(\d+(\.\d+)*)[\s\.\:]?', re.IGNORECASE)
    if heading_detection_pattern.search(text_content):
        return False
    if style_name.startswith('Heading'):
        return False # Keep this check too for safety
    if style_name in ['Title', 'Subtitle', 'Caption', 'Quote']:
        return False
    style_name_lower = style_name.lower()
    list_indicators = ['list', 'bullet', 'number', 'opsomming', 'lijst']
    if any(indicator in style_name_lower for indicator in list_indicators):
        return False
    # AI comments and layout comments are excluded
    if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
         return False

    if len(text_content) < 60:
        return False # Keep original length check for separation logic

    return True

# --- Function to check if a paragraph is a "text block" (for length check) ---
def is_text_block_for_length_check(para):
    """Checks if a paragraph is a text block for length check (not empty, not a heading, not a list, no special style, no comment)."""
    text_content = para.text.strip()
    style_name = para.style.name if para.style else "None"

    if not para or not text_content:
        return False
    # Updated: Also exclude paragraphs matching the heading detection pattern
    heading_detection_pattern = re.compile(r'^\s*(\d+(\.\d+)*)[\s\.\:]?', re.IGNORECASE)
    if heading_detection_pattern.search(text_content):
        return False
    if style_name.startswith('Heading'):
        return False # Keep this check too for safety
    if style_name in ['Title', 'Subtitle', 'Caption', 'Quote']:
        return False
    style_name_lower = style_name.lower()
    list_indicators = ['list', 'bullet', 'number', 'opsomming', 'lijst']
    if any(indicator in style_name_lower for indicator in list_indicators):
        return False
    # Exclude comments themselves
    if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
         return False

    return True


# --- Function for checking the title rule ---
# This check is now called from the main loop, but still operates on the entire document.
# The 'all' setting in CHECK_SETTINGS ensures it is always executed.
def check_title_rule(document):
    errors = []
    print("\nDEBUG: Checking title rule...")
    title_paragraph, title_index = find_first_non_empty_paragraph(document)

    if title_paragraph is None:
        errors.append({"location": document, "message": "Opmerking: Document is leeg of bevat alleen lege paragrafen. Geen titel gevonden."})
        print("DEBUG: Document is empty, no title found.")
        return errors, title_index

    title_text = title_paragraph.text.strip()

    forbidden_titles = ["scriptie", "plan van aanpak"]
    if title_text.lower() in forbidden_titles:
        errors.append({
            "location": title_paragraph,
            "message": f"Opmerking: De titel is te algemeen: '{title_text}'. Gebruik een specifieke en informatieve titel voor je document."
        })
        print(f"DEBUG: Found forbidden title: '{title_text}' at index {title_index}.")

    min_title_font_size_pt = 18
    actual_font_size = None
    if title_paragraph.runs:
        for run in title_paragraph.runs:
            if run.font.size:
                try:
                    if actual_font_size is None or run.font.size.pt > actual_font_size:
                        actual_font_size = run.font.size.pt
                except AttributeError:
                    pass

    if actual_font_size is None or (actual_font_size is not None and actual_font_size < min_title_font_size_pt):
        errors.append({
            "location": title_paragraph,
            "message": f"Opmerking: De titel '{title_text}' heeft een te klein lettertype ({actual_font_size if actual_font_size is not None else 'onbekend'}pt). De titel moet minimaal {min_title_font_size_pt}pt zijn."
        })
        print(f"DEBUG: Title font size too small at index {title_index}. Found {actual_font_size}pt, required {min_title_font_size_pt}pt.")

    return errors, title_index

# --- Function to check for potentially literal legal texts (WITH LIMIT) ---
# This check is now called from the main loop for each paragraph.
def check_legal_citations_in_paragraph(para, para_index):
    errors = []
    citation_patterns = [
        r"Op grond van art\.?",
        r"Artikel \d+",
        r"art\. \d+",
        r"lid \d+:",
        r"in artikel \d+",
        r"conform artikel \d+"
    ]
    compiled_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in citation_patterns]

    text_content = para.text.strip()
    style_name = para.style.name if para.style else "None"
    style_name_lower = style_name.lower()
    list_indicators = ['list', 'bullet', 'number', 'opsomming', 'lijst']
    if not text_content or style_name.startswith('Heading') or any(indicator in style_name_lower for indicator in list_indicators) or style_name in ['Title', 'Subtitle', 'Caption', 'Quote']:
         return errors # No errors in this type of paragraph
    # Exclude comments themselves
    if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
         return errors # No errors in comment paragraphs


    for pattern in compiled_patterns:
        if pattern.search(text_content):
            # The limit logic is now handled in the main loop
            errors.append({
                "location": para,
                "message": "Opmerking: wettekst niet letterlijk citeren."
            })
            print(f"DEBUG: Found potential legal citation pattern in paragraph {para_index}. Reporting error.")
            break # Stop after finding the first pattern in this paragraph

    return errors


# --- Check first "Ik" (I/me) ---
# This check is now called from the main loop for each paragraph.
def check_first_ik_in_paragraph(para, para_index, ik_already_found):
    errors = []
    # Use a regex pattern to find 'ik', 'mij', 'me' as whole words or starting/ending a word
    # case-insensitive, exclude occurrences within AI/Layout comments
    ik_pattern = re.compile(r"\b(ik|mij|me)\b", re.IGNORECASE)

    text_content = para.text.strip()
    # Exclude comments themselves from this check
    if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
         return errors, ik_already_found # No errors in comment paragraphs, return status

    if ik_pattern.search(text_content):
        if not ik_already_found: # Only report the first time
            errors.append({
                "location": para,
                "message": "Opmerking: Schrijf dit hele document in een zakelijke stijl. Dat betekent geen ik of mij of me. De noodzaak van het onderzoek zou immers op geen enkele manier van jou afhankelijk moeten zijn."
            })
            print(f"DEBUG: Found first 'Ik/Mij/Me' in paragraph {para_index}: '{text_content[:50]}...'")
            return errors, True # Error found, status is now True
        else:
            return errors, ik_already_found # Already found, no new error, status remains True
    else:
        return errors, ik_already_found # Not found, status remains the same


# --- Check Sub-questions structure (numbering check removed) ---
# This check is now called from the main loop when the "Deelvragen" section is found.
def check_deelvragen_structure_in_section(document, deelvragen_header_index, end_of_deelvragen_section_index):
    errors = []
    print("\nDEBUG: Checking deelvraag structure in section...")

    # Find potential deelvraag paragraphs within the section boundaries
    potential_deelvraag_paras = []

    for idx in range(deelvragen_header_index + 1, end_of_deelvragen_section_index):
        para = document.paragraphs[idx]
        text_content = para.text.strip()

        # Exclude comments and empty lines
        if not text_content or text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
             continue

        # Simple check if it looks like a question (contains a question mark)
        question_pattern = re.compile(r"[\?]$", re.IGNORECASE) # Check if it ends with a question mark

        if question_pattern.search(text_content):
             potential_deelvraag_paras.append({"para": para, "index": idx})
             print(f"DEBUG: Found potential deelvraag at index {idx} within section: '{text_content[:50]}'")


    if not potential_deelvraag_paras:
         print("DEBUG: No potential deelvraag paragraphs found in section after the header.")
         # We could add an error here if sub-questions are mandatory
         return errors

    # Check structure (Rule 2) - Simplified check for multiple main clauses
    print("DEBUG: Checking deelvraag structure (simplified) within section...")
    # Pattern to find text with a question mark, followed by 'en' (as a word boundary), followed by another question mark
    structure_error_pattern = re.compile(r"[\?].*\b(en|maar|of)\b.*[\?]", re.IGNORECASE) # Added 'maar' and 'of' as common conjunctions
    multiple_question_marks_pattern = re.compile(r"[\?].*[\?]") # Multiple question marks in one line

    for dv_info in potential_deelvraag_paras:
        para = dv_info['para']
        idx = dv_info['index']
        text_content = para.text.strip()

        # Check if the text contains indicators of multiple main clauses within one question
        if structure_error_pattern.search(text_content) or multiple_question_marks_pattern.search(text_content):
             errors.append({
                 "location": para,
                 "message": "Opmerking: Een deelvraag mag niet uit twee aparte hoofdzinnen bestaan (mogelijke indicatie door verbindingswoord na vraagteken of meerdere vraagtekens)."
             })
             print(f"DEBUG: Deelvraag structure error detected at index {idx} within section: '{text_content[:50]}...'")


    return errors

# --- Check Paragraph 1.3 content (Intro sentences) ---
# This check is now called from the main loop when section "1.3" is found.
def check_section_1_3_intro_in_section(document, section_1_3_index, end_of_1_3_section_index):
    errors = []
    print("\nDEBUG: Checking section 1.3 intro content...")

    section_1_3_content_paras = document.paragraphs[section_1_3_index + 1 : end_of_1_3_section_index]

    if not section_1_3_content_paras:
         # errors.append({"location": document.paragraphs[section_1_3_index], "message": "Opmerking: Paragraaf 1.3 (Probleemstelling/Vragen) lijkt geen inhoud te hebben na de titel."}) # This check can be elsewhere, or if AI finds no questions
         print("DEBUG: Section 1.3 has no content paragraphs for intro check.")
         return errors # No content to check intro


    # --- Check for introductory sentences before the first question ---
    intro_sentences_found = False
    first_question_index_in_1_3_content_list = -1 # Index within the section_1_3_content_paras list

    # Find the first potential question within the 1.3 content paragraphs
    first_question_original_index = -1
    question_pattern = re.compile(r"[\?]") # Simple pattern to find lines with question marks
    for i, para in enumerate(section_1_3_content_paras):
         text_content = para.text.strip()
         # Exclude comments from the text sent to AI
         if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
              continue

         if question_pattern.search(text_content):
              if first_question_original_index == -1: # Capture the first one found
                   first_question_original_index = section_1_3_index + 1 + i
                   first_question_index_in_1_3_content_list = i
              # We don't need to find all questions for this check, only the first one
              # break # Commented out break to find all questions for the AI check later


    if first_question_original_index != -1:
         # Check paragraphs from 1.3 header + 1 up to the paragraph before the first question
         # Ensure there's at least one paragraph *before* the first question to check
         if first_question_index_in_1_3_content_list > 0:
              intro_check_paras = section_1_3_content_paras[0 : first_question_index_in_1_3_content_list]
              for para in intro_check_paras:
                  text_content = para.text.strip()
                  if text_content and not (text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:")):
                      # Crude check for sentence length vs single word/short phrase
                      if len(text_content) > 30: # Threshold can be adjusted
                           intro_sentences_found = True
                           print(f"DEBUG: Found potential intro sentence before first question at index {section_1_3_index + 1 + intro_check_paras.index(para)}.")
                           break # Found an intro sentence, no need to check further


         # Check if the very first content paragraph is a question and thus has no preceding intro in this section
         elif first_question_index_in_1_3_content_list == 0:
             # If the first content paragraph is a question, there's no text before it in this section
             intro_sentences_found = False
             print("DEBUG: First content paragraph in 1.3 is a question. No preceding intro found in this section.")


         # If there are paragraphs before the first question (i.e., the first content paragraph IS the first question)
         # then intro_sentences_found should remain False if no preceding text was found.
         # If there are paragraphs before the first question but they are all empty or comments, intro_sentences_found is also False.

    # Add error if intro sentences are missing BEFORE the first question
    # This check is only relevant if at least one question was found (assuming that's the main question)
    # We now need to check if there are any questions at all in 1.3 (the AI check also does this)
    # But for this specific check (intro) we only look at the structure before the first question.
    # We no longer have potential_question_paras_in_1_3 here, we use first_question_original_index
    if first_question_original_index != -1 and not intro_sentences_found:
         errors.append({
             "location": document.paragraphs[section_1_3_index], # Link to 1.3 header
             "message": "Opmerking: De hoofdvraag in Paragraaf 1.3 moet worden ingeleid met een of meer zinnen. Er lijkt geen inleidende tekst te zijn vóór de eerste vraag."
         })
         print("DEBUG: Missing intro sentences before first question in 1.3.")


    return errors


# --- Check Paragraph 1.4 content (Objective, Output, Outcome) ---
# This check is now called from the main loop when section "1.4" is found.
def check_section_1_4_content_in_section(document, section_1_4_index, end_of_1_4_section_index):
    errors = []
    print("\nDEBUG: Checking section 1.4 content (Objective, Output, Outcome)...")

    section_1_4_content_paras = document.paragraphs[section_1_4_index + 1 : end_of_1_4_section_index]

    if not section_1_4_content_paras:
         errors.append({"location": document.paragraphs[section_1_4_index], "message": "Opmerking: Paragraaf 1.4 lijkt geen inhoud te hebben na de titel."})
         print("DEBUG: Section 1.4 has no content paragraphs.")
         return errors


    # Check for presence of Objective, Output, Outcome (using text matching)
    doelstelling_found = False
    output_found = False
    outcome_found = False

    # Patterns for finding the required elements (case-insensitive, whole words or starting word)
    doelstelling_pattern = re.compile(r"\bdoelstelling\b|\bdoel\b", re.IGNORECASE) # Check for both "doelstelling" and "doel"
    output_pattern = re.compile(r"\boutput\b", re.IGNORECASE)
    outcome_pattern = re.compile(r"\boutcome\b", re.IGNORECASE)

    for i, para in enumerate(section_1_4_content_paras):
        text_content = para.text.strip()
         # Exclude comments
        if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
             continue

        if doelstelling_pattern.search(text_content):
             doelstelling_found = True
             print(f"DEBUG: 'Doelstelling'/'Doel' found in paragraph index {section_1_4_index + 1 + i} (within 1.4 content).")
        if output_pattern.search(text_content):
             output_found = True
             print(f"DEBUG: 'Output' found in paragraph index {section_1_4_index + 1 + i} (within 1.4 content).")
        if outcome_pattern.search(text_content):
             outcome_found = True
             print(f"DEBUG: 'Outcome' found in paragraph index {section_1_4_index + 1 + i} (within 1.4 content).")


    if not doelstelling_found:
         errors.append({
             "location": document.paragraphs[section_1_4_index], # Link to 1.4 header
             "message": "Opmerking: De doelstelling lijkt te ontbreken in Paragraaf 1.4."
         })
         print("DEBUG: 'Doelstelling' not found in section 1.4.")

    if not output_found:
         errors.append({
             "location": document.paragraphs[section_1_4_index], # Link to 1.4 header
             "message": "Opmerking: Het kopje of de vermelding van 'Output' lijkt te ontbreken in Paragraaf 1.4."
         })
         print("DEBUG: 'Output' not found in section 1.4.")

    if not outcome_found:
         errors.append({
             "location": document.paragraphs[section_1_4_index], # Link to 1.4 header
             "message": "Opmerking: Het kopje of de vermelding van 'Outcome' lijkt te ontbreken in Paragraaf 1.4."
         })
         print("DEBUG: 'Outcome' not found in section 1.4.")

    return errors


# --- Check Paragraph Lengths ---
# This check is now called from the main loop for each paragraph.
def check_paragraph_length_in_paragraph(para, para_index, current_section): # Added current_section
    errors = []
    # print("DEBUG: Checking paragraph length...") # Too verbose in main loop

    # Exclude section 1.3 from this check
    if current_section == '1.3':
        print(f"DEBUG: Skipping paragraph length check for section 1.3 (paragraph {para_index}).")
        return errors

    SHORT_PARA_MIN = 150
    SHORT_PARA_MAX = 350
    LONG_PARA_MIN = 1200

    # Only check paragraphs that are considered text blocks for this purpose
    if not is_text_block_for_length_check(para):
         return errors

    text_content = para.text.strip()
    text_length = len(text_content)

    if text_length >= SHORT_PARA_MIN and text_length <= SHORT_PARA_MAX:
         errors.append({
             "location": para,
             "message": f"Opmerking: Nogal korte alinea ({text_length} tekens). Overweeg wijziging."
         })
         print(f"DEBUG: Short paragraph detected at index {para_index} ({text_length} chars): '{text_content[:50]}...'")
    elif text_length > LONG_PARA_MIN:
         errors.append({
             "location": para,
             "message": f"Opmerking: Nogal lange alinea ({text_length} tekens). Overweeg andere indeling."
         })
         print(f"DEBUG: Long paragraph detected at index {para_index} ({text_length} chars): '{text_content[:50]}...'")

    return errors

# --- Check paragraph separation ---
# This check is performed at the end because it needs the entire list of paragraphs.
def check_paragraph_separation(document):
    errors = []
    print("\nDEBUG: Checking paragraph separation...")

    last_real_paragraph_index = -1 # Track the index of the last encountered "real" paragraph

    for i, para in enumerate(document.paragraphs):
        if is_real_paragraph_for_separation(para):
            if last_real_paragraph_index != -1:
                # Found a new "real" paragraph, check the spacing since the last one
                intermediate_paragraphs_count = i - last_real_paragraph_index - 1

                if intermediate_paragraphs_count != 1:
                    # The paragraph *before* the current one (at last_real_paragraph_index) is where the issue starts
                    errors.append({
                        "location": document.paragraphs[last_real_paragraph_index], # Link error to the paragraph *before* the gap
                        "message": "Opmerking: Onjuiste alineascheiding: gebruik exact één witregel."
                    })
                    print(f"DEBUG: Incorrect paragraph separation detected between paragraph {last_real_paragraph_index} and {i}. Found {intermediate_paragraphs_count} intermediate paragraphs.")

            # Update the index of the last encountered "real" paragraph
            last_real_paragraph_index = i

    print(f"DEBUG: Finished paragraph separation check.")
    return errors


# --- NEW AI CHECK: Analyze 1.1 and 1.2 content ---
# This check is now called from the main loop when section "1.1" is found.
def analyze_section_1_1_1_2_content_in_section(document, section_1_1_index, end_of_1_2_section_index):
    errors = []
    if model is None:
        print("DEBUG: AI model not available for 1.1/1.2 analysis.")
        errors.append({"location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document, "message": "Opmerking: AI analyse voor sectie 1.1 en 1.2 overgeslagen: Model niet beschikbaar."})
        return errors

    print("\n--- Start AI analysis for section 1.1 and 1.2 ---")

    # Extract text content for 1.1 and 1.2
    section_1_1_1_2_text = ""
    # We need the paragraph objects and their info for potential error linking later
    section_1_1_1_2_paragraphs_info = []
    for idx in range(section_1_1_index, end_of_1_2_section_index):
        para = document.paragraphs[idx]
        text_content = para.text.strip()
        # Exclude comments from the text sent to AI
        if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
             continue
        section_1_1_1_2_text += text_content + "\n"
        if hasattr(para, '_element'):
             section_1_1_1_2_paragraphs_info.append((para, idx, para._element))
        else:
             print(f"WARNING: Paragraph {idx} in 1.1/1.2 has no '_element' attribute. Cannot reliably process it for AI analysis placement.")


    if not section_1_1_1_2_text.strip():
        errors.append({"location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document, "message": "Inhoudelijke feedback: Sectie 1.1 en 1.2 lijken geen relevante tekstuele inhoud te bevatten voor analyse."})
        print("DEBUG: No relevant text content found in 1.1/1.2 for AI analysis.")
        return errors


    prompt = textwrap.dedent(f"""
    Analyze the following text from sections 1.1 and 1.2 of a juridical scriptie concept. Evaluate the following points in Dutch:
    1.  Is the juridical problem clearly and concretely described?
    2.  Is the relevance of the problem for juridical practice or theory sufficiently motivated?
    3.  Is it clear what juridical research is necessary to solve the practical problem mentioned? (Based on the 'handelingsprobleem' in 1.2)
    4.  Is it clear what practical research is necessary to solve the practical problem mentioned? (Based on the 'handelingsprobleem' in 1.2)

    For each point, provide a brief assessment in Dutch. If there are issues or if the point is not addressed, explain it. If the point is well addressed, state that.
    Format your response as a list of comments, each on a new line, starting with "Inhoudelijke feedback: ".

    Example format:
    Inhoudelijke feedback: Juridisch probleem is duidelijk beschreven.
    Inhoudelijke feedback: Relevantie voor de praktijk kan nog beter worden gemotiveerd.
    Inhoudelijke feedback: Het benodigde juridisch onderzoek wordt nog niet duidelijk.

    Text to analyze (Sections 1.1 and 1.2):
    ---
    {section_1_1_1_2_text.strip()}
    ---
    """)

    print(f"DEBUG: Sending 1.1/1.2 text (approx {len(section_1_1_1_2_text)} chars) to Gemini API for analysis...")
    try:
        response = model.generate_content(prompt, request_options={"timeout": 120}, generation_config=genai.GenerationConfig(temperature=0))

        if hasattr(response, 'text'):
            response_text = response.text.strip()
            print(f"DEBUG_AI: Raw AI response for 1.1/1.2: {response_text[:200]}...") # Added debug print


            if response._result.prompt_feedback and response._result.prompt_feedback.block_reason:
                print(f"WARNING: AI analysis for 1.1/1.2 blocked: {response._result.prompt_feedback.block_reason}") # Minimal debug
                errors.append({"location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document, "message": f"Inhoudelijke feedback: Analyse van sectie 1.1 en 1.2 geblokkeerd door veiligheidsfilter."})
                return errors

            if response_text and response_text.upper() != "OK": # Check for non-OK response
                # Parse the AI response - expect lines starting with "Inhoudelijke feedback: "
                lines = response_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith("Inhoudelijke feedback: "):
                        # Link the comment to the start of section 1.1 for now
                        # For section-specific AI comments, we link to the section header paragraph
                        errors.append({
                            "location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document,
                            "original_index": section_1_1_index if len(document.paragraphs) > section_1_1_index else -1,
                            "element": document.paragraphs[section_1_1_index]._element if len(document.paragraphs) > section_1_1_index and hasattr(document.paragraphs[section_1_1_index], '_element') else None,
                            "message": line # Keep the "Inhoudelijke feedback: " prefix
                        })
                        print(f"DEBUG: Parsed 1.1/1.2 section AI comment: '{line[:100]}...'")
                    elif line.strip(): # Handle any non-empty line that doesn't start with Inhoudelijke feedback:
                        # Handle unexpected response format
                        errors.append({
                            "location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document,
                            "original_index": section_1_1_index if len(document.paragraphs) > section_1_1_index else -1,
                             "element": document.paragraphs[section_1_1_index]._element if len(document.paragraphs) > section_1_1_index and hasattr(document.paragraphs[section_1_1_index], '_element') else None,
                            "message": f"Inhoudelijke feedback: Onverwacht AI reactieformaat voor 1.1/1.2: '{line[:100]}...'"
                        })
                        print(f"WARNING: Unexpected AI format for 1.1/1.2: '{line[:50]}...'") # Minimal debug

            elif response_text.upper() == "OK":
                print("INFO: AI analysis for 1.1/1.2 was OK.") # Added debug print
            else:
                print("INFO: AI analysis for 1.1/1.2 returned no textual response.") # Minimal debug
                errors.append({"location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document, "message": "Inhoudelijke feedback: Geen analyse resultaten ontvangen voor sectie 1.1 en 1.2."})


        else:
            print("INFO: AI analysis for 1.1/1.2 response object had no text attribute.") # Minimal debug
            errors.append({"location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document, "message": "Inhoudelijke feedback: Geen analyse resultaten ontvangen voor sectie 1.1 en 1.2."})


    except Exception as e:
        print(f"ERROR: While calling the Gemini API for 1.1/1.2: {e}") # Minimal debug
        errors.append({"location": document.paragraphs[section_1_1_index] if len(document.paragraphs) > section_1_1_index else document, "message": f"Inhoudelijke feedback: Fout bij AI analyse voor sectie 1.1 en 1.2: {e}"})

    print("--- End AI analysis for section 1.1 and 1.2 ---")
    return errors


# --- NEW AI CHECK: Analyze 1.3 questions content ---
# This check is now called from the main loop when section "1.3" is found.
def analyze_section_1_3_questions_content_in_section(document, section_1_3_index, end_of_1_3_section_index):
    errors = []
    if model is None:
        print("DEBUG: AI model not available for 1.3 questions analysis.")
        errors.append({"location": document.paragraphs[section_1_3_index] if len(document.paragraphs) > section_1_3_index else document, "message": "Opmerking: AI analyse voor vragen in sectie 1.3 overgeslagen: Model niet beschikbaar."})
        return errors

    print("\n--- Start AI analysis for questions in section 1.3 ---")

    section_1_3_content_paras = document.paragraphs[section_1_3_index + 1 : end_of_1_3_section_index]

    # Extract text content for questions in 1.3
    questions_text = ""
    question_paras_info = [] # Keep track of paragraphs info (para_object, original_idx, para_element) for potential error linking
    question_pattern = re.compile(r"[\?]") # Simple pattern to find lines with question marks

    for idx, para in enumerate(section_1_3_content_paras):
        text_content = para.text.strip()
        # Exclude comments from the text sent to AI
        if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
             continue

        if question_pattern.search(text_content):
             questions_text += text_content + "\n"
             original_idx = section_1_3_index + 1 + idx # Calculate original index
             if hasattr(para, '_element'):
                  question_paras_info.append((para, original_idx, para._element))
             else:
                  print(f"WARNING: Question paragraph {original_idx} in 1.3 has no '_element' attribute. Cannot reliably process it for AI analysis placement.")


    if not questions_text.strip():
        errors.append({"location": document.paragraphs[section_1_3_index] if len(document.paragraphs) > section_1_3_index else document, "message": "Inhoudelijke feedback: Geen vragen gevonden in sectie 1.3 voor analyse."})
        print("DEBUG: No question marks found in 1.3 for AI analysis.")
        return errors # No questions found in 1.3


    # Use the first found question paragraph as the location for the AI comment, if available
    # Otherwise, use the 1.3 header paragraph
    error_location_para = question_paras_info[0][0] if question_paras_info else (document.paragraphs[section_1_3_index] if len(document.paragraphs) > section_1_3_index else document)
    error_location_idx = question_paras_info[0][1] if question_paras_info else (section_1_3_index if len(document.paragraphs) > section_1_3_index else -1)
    error_location_element = question_paras_info[0][2] if question_paras_info and question_paras_info[0][2] is not None else (document.paragraphs[section_1_3_index]._element if len(document.paragraphs) > section_1_3_index and hasattr(document.paragraphs[section_1_3_index], '_element') else None)


    prompt = textwrap.dedent(f"""
    Analyze the following questions from section 1.3 of a juridical scriptie concept. Evaluate the following points in Dutch:
    1.  Does the main question meet the requirements for a juridical scriptie (specific, relevant, researchable)?
    2.  Do the sub-questions sufficiently cover the main question?
    3.  Is there a logical coherence between the main question and the sub-questions?
    4.  De student formuleert open deelvragen die onderzoekbaar zijn en voldoende specifiek. De deelvragen zijn onderling consistent en tevens consistent met de hoofdvraag. De deelvragen vertonen een opbouw in complexiteit.

    Assume the first question is the main question, followed by sub-questions.

    For each point, provide a brief assessment in Dutch. If there are issues or if the point is not addressed, explain it. If the point is well addressed, state that.
    Format your response as a list of comments, each on a new line, starting with "Inhoudelijke feedback: ".

    Example format:
    Inhoudelijke feedback: Hoofdvraag is specifiek en relevant.
    Inhoudelijke feedback: Deelvragen dekken de hoofdvraag goed af.
    Inhoudelijke feedback: De samenhang tussen hoofd- en deelvragen kan nog duidelijker.
    Inhoudelijke feedback: De deelvragen zijn open, onderzoekbaar en consistent met de hoofdvraag.

    Questions to analyze (from Section 1.3):
    ---
    {questions_text.strip()}
    ---
    """)

    print(f"DEBUG: Sending 1.3 questions text (approx {len(questions_text)} chars) to Gemini API for analysis...")
    try:
        response = model.generate_content(prompt, request_options={"timeout": 120}, generation_config=genai.GenerationConfig(temperature=0))

        if hasattr(response, 'text'):
            response_text = response.text.strip()
            print(f"DEBUG_AI: Raw AI response for 1.3 questions: {response_text[:200]}...") # Added debug print

            if response._result.prompt_feedback and response._result.prompt_feedback.block_reason:
                print(f"WARNING: AI analysis for 1.3 questions blocked: {response._result.prompt_feedback.block_reason}") # Minimal debug
                errors.append({"location": error_location_para, "original_index": error_location_idx, "element": error_location_element, "message": f"Inhoudelijke feedback: Analyse van vragen in sectie 1.3 geblokkeerd door veiligheidsfilter."})
                return errors


            if response_text and response_text.upper() != "OK": # Check for non-OK response
                # Parse the AI response - expect lines starting with "Inhoudelijke feedback: "
                lines = response_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith("Inhoudelijke feedback: "):
                        # Link the comment to the first question paragraph found, or the 1.3 header
                        errors.append({
                            "location": error_location_para, # Use the determined error location paragraph
                            "original_index": error_location_idx, # Use the determined error location index
                            "element": error_location_element, # Use the determined error location element
                            "message": line # Keep the "Inhoudelijke feedback: " prefix
                        })
                        print(f"DEBUG: Parsed 1.3 questions section AI comment: '{line[:100]}...'")

                    elif line.strip(): # Handle any non-empty line that doesn't start with Inhoudelijke feedback:
                        # Handle unexpected response format
                        errors.append({
                            "location": error_location_para,
                            "original_index": error_location_idx,
                            "element": error_location_element,
                            "message": f"Inhoudelijke feedback: Onverwacht AI reactieformaat voor 1.3 vragen: '{line[:100]}...'"
                        })
                        print(f"WARNING: Unexpected AI format for 1.3 questions: '{line[:50]}...'") # Minimal debug

            elif response_text.upper() == "OK":
                print("INFO: AI analysis for 1.3 questions was OK.") # Added debug print
            else:
                print("INFO: AI analysis for 1.3 questions returned no textual response.") # Minimal debug
                errors.append({"location": error_location_para, "original_index": error_location_idx, "element": error_location_element, "message": "Inhoudelijke feedback: Geen analyse resultaten ontvangen voor vragen in sectie 1.3."})


        else:
            print("INFO: AI analysis for 1.3 questions response object had no text attribute.") # Minimal debug
            errors.append({"location": error_location_para, "original_index": error_location_idx, "element": error_location_element, "message": "Inhoudelijke feedback: Geen analyse resultaten ontvangen voor vragen in sectie 1.3."})


    except Exception as e:
        print(f"ERROR: While calling the Gemini API for 1.3 questions: {e}") # Minimal debug
        errors.append({"location": error_location_para, "original_index": error_location_idx, "element": error_location_element, "message": f"Inhoudelijke feedback: Fout bij AI analyse voor vragen in sectie 1.3: {e}"})

    print("--- End AI analysis for questions in section 1.3 ---")
    return errors

# --- NEW AI CHECK: Analyze Chapter 2 content ---
def analyze_chapter_2_content_in_section(document, section_2_index, end_of_2_section_index):
    errors = []
    if model is None:
        print("DEBUG: AI model not available for Chapter 2 analysis.")
        errors.append({"location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document, "message": "Opmerking: AI analyse voor Hoofdstuk 2 overgeslagen: Model niet beschikbaar."})
        return errors

    print("\n--- Start AI analysis for Chapter 2 ---")

    chapter_2_text = ""
    chapter_2_paragraphs_info = []
    for idx in range(section_2_index, end_of_2_section_index):
        para = document.paragraphs[idx]
        text_content = para.text.strip()
        # Exclude comments from the text sent to AI
        if text_content.startswith("Inhoudelijke feedback:") or text_content.startswith("Feedback op tekst:") or text_content.startswith("Opmerking:") or text_content.startswith("AI analyse voor dit stuk tekst") or text_content.startswith("DC:") or text_content.startswith("[opmerking]:") or text_content.startswith("AI opmerking:"):
             continue
        chapter_2_text += text_content + "\n"
        if hasattr(para, '_element'):
             chapter_2_paragraphs_info.append((para, idx, para._element))
        else:
             print(f"WARNING: Paragraph {idx} in Chapter 2 has no '_element' attribute. Cannot reliably process it for AI analysis placement.")


    if not chapter_2_text.strip():
        errors.append({"location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document, "message": "Inhoudelijke feedback: Hoofdstuk 2 lijkt geen relevante tekstuele inhoud te bevatten voor analyse."})
        print("DEBUG: No relevant text content found in Chapter 2 for AI analysis.")
        return errors

    prompt = textwrap.dedent(f"""
    Analyze the following text from Chapter 2 (Juridische Context) of a juridical scriptie concept. Evaluate the following points in Dutch:
    1.  Is the student able to independently conduct a juridical exploration, tailored to the problem description and in relation to the action problem?
    2.  Does the student provide sufficient background information?
    3.  Is the provided information relevant and substantively anchored?
    4.  Pay close attention to the relevance of the given sources.

    For each point, provide a brief assessment in Dutch. If there are issues or if the point is not addressed, explain it. If the point is well addressed, state that.
    Format your response as a list of comments, each on a new line, starting with "Inhoudelijke feedback: ".

    Example format:
    Inhoudelijke feedback: De student is in staat zelfstandig een juridische verkenning te maken.
    Inhoudelijke feedback: De achtergrondinformatie is voldoende.
    Inhoudelijke feedback: De relevantie van de bronnen kan nog beter worden toegelicht.

    Text to analyze (Chapter 2):
    ---
    {chapter_2_text.strip()}
    ---
    """)

    print(f"DEBUG: Sending Chapter 2 text (approx {len(chapter_2_text)} chars) to Gemini API for analysis...")
    try:
        response = model.generate_content(prompt, request_options={"timeout": 120}, generation_config=genai.GenerationConfig(temperature=0))

        if hasattr(response, 'text'):
            response_text = response.text.strip()
            print(f"DEBUG_AI: Raw AI response for Chapter 2: {response_text[:200]}...")

            if response._result.prompt_feedback and response._result.prompt_feedback.block_reason:
                print(f"WARNING: AI analysis for Chapter 2 blocked: {response._result.prompt_feedback.block_reason}")
                errors.append({"location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document, "message": f"Inhoudelijke feedback: Analyse van Hoofdstuk 2 geblokkeerd door veiligheidsfilter."})
                return errors

            if response_text and response_text.upper() != "OK":
                lines = response_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith("Inhoudelijke feedback: "):
                        errors.append({
                            "location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document,
                            "original_index": section_2_index if len(document.paragraphs) > section_2_index else -1,
                            "element": document.paragraphs[section_2_index]._element if len(document.paragraphs) > section_2_index and hasattr(document.paragraphs[section_2_index], '_element') else None,
                            "message": line
                        })
                        print(f"DEBUG: Parsed Chapter 2 AI comment: '{line[:100]}...'")
                    elif line.strip():
                        errors.append({
                            "location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document,
                            "original_index": section_2_index if len(document.paragraphs) > section_2_index else -1,
                            "element": document.paragraphs[section_2_index]._element if len(document.paragraphs) > section_2_index and hasattr(document.paragraphs[section_2_index], '_element') else None,
                            "message": f"Inhoudelijke feedback: Onverwacht AI reactieformaat voor Hoofdstuk 2: '{line[:100]}...'"
                        })
                        print(f"WARNING: Unexpected AI format for Chapter 2: '{line[:50]}...'")

            elif response_text.upper() == "OK":
                print("INFO: AI analysis for Chapter 2 was OK.")
            else:
                print("INFO: AI analysis for Chapter 2 returned no textual response.")
                errors.append({"location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document, "message": "Inhoudelijke feedback: Geen analyse resultaten ontvangen voor Hoofdstuk 2."})
        else:
            print("INFO: AI analysis for Chapter 2 response object had no text attribute.")
            errors.append({"location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document, "message": "Inhoudelijke feedback: Geen analyse resultaten ontvangen voor Hoofdstuk 2."})
    except Exception as e:
        print(f"ERROR: While calling the Gemini API for Chapter 2: {e}")
        errors.append({"location": document.paragraphs[section_2_index] if len(document.paragraphs) > section_2_index else document, "message": f"Inhoudelijke feedback: Fout bij AI analyse voor Hoofdstuk 2: {e}"})

    print("--- End AI analysis for Chapter 2 ---")
    return errors


# --- Function to split document paragraphs into text chunks for AI analysis (general analysis) ---
# This function now receives a list of (paragraph_object, original_index, paragraph_element) tuples
def chunk_paragraphs_for_ai(paragraphs_info, max_chars_per_chunk=4000 * 3.5):
    chunks = []
    current_chunk_paras_info = [] # Store (para_object, original_idx, para_element) tuples
    current_chunk_text = ""

    for para, original_idx, para_element in paragraphs_info:
        # Skip comments from being added to chunks for AI analysis
        if para.text.strip().startswith("Inhoudelijke feedback:") or para.text.strip().startswith("Feedback op tekst:") or para.text.strip().startswith("Opmerking:") or para.text.strip().startswith("AI analyse voor dit stuk tekst") or para.text.strip().startswith("DC:") or para.text.strip().startswith("[opmerking]:") or para.text.strip().startswith("AI opmerking:"):
            continue

        para_text = para.text

        # Check if adding the current paragraph exceeds the max chunk size
        if len(current_chunk_text) + len(para_text) + 1 > max_chars_per_chunk and current_chunk_text.strip():
            # Add the current chunk to the list
            chunks.append({
                "text": current_chunk_text.strip(),
                "paragraphs_info": current_chunk_paras_info, # Store the list of (para, index, element) tuples
                "start_para_original_idx": current_chunk_paras_info[0][1], # Get index from the first tuple
                "end_para_original_idx": current_chunk_paras_info[-1][1] # Get index from the last tuple
            })

            # Start a new chunk with the current paragraph
            current_chunk_text = para_text
            current_chunk_paras_info = [(para, original_idx, para_element)]

        else:
            # Add the current paragraph to the current chunk
            if current_chunk_text:
                current_chunk_text += "\n" + para_text
            else:
                current_chunk_text = para_text
            current_chunk_paras_info.append((para, original_idx, para_element))


    # Add the last chunk if it's not empty
    if current_chunk_text.strip():
        chunks.append({
            "text": current_chunk_text.strip(),
            "paragraphs_info": current_chunk_paras_info,
            "start_para_original_idx": current_chunk_paras_info[0][1] if current_chunk_paras_info else 0,
            "end_para_original_idx": current_chunk_paras_info[-1][1] if current_chunk_paras_info else 0 # Fallback for end index
        })

    return chunks


# --- Function to perform AI analysis on a text chunk (general analysis) ---
def analyze_chunk_with_gemini(chunk_text):
    if model is None:
        print("DEBUG: AI model not available for chunk analysis.")
        return "[AI Analysis skipped: Model not available]"

    print(f"INFO: Starting general AI analysis for chunk (first 100 characters): {chunk_text[:100]}...") # Added debug print

    # Adjusted prompt slightly for clarity to AI
    prompt = textwrap.dedent(f"""
    Analyze the following text chunk from a document for logical flaws and unclear, confusing, or sentences with awkward phrasing or structure.
    Focus only on the *content* and *structure* of the main prose within the chunk. Do NOT critique bibliography entries, section titles, or other non-prose elements unless they are part of a sentence being analyzed.

    For each issue you find:
    1. Quote the *exact* sentence or phrase from the text that contains the issue. Enclose the quote in double quotes ".
    2. Briefly explain the nature of the issue in Dutch (e.g., 'illogische stap', 'verwarrende formulering', 'probleem met zinsstructuur').
    3. Present each finding on a new line, starting with the quoted text, followed by a colon and space, then the explanation.
    4. Ensure each line STRICTLY follows the format "Quote" : Explanation and contains ONLY one critique. Do not add any other text or numbering outside this format within your critique list.

    Example format:
    "Deze zin is onduidelijk hoe die aansluit." : verwarrende formulering
    "Hierdoor is de conclusie onjuist op basis van..." : illogische stap

    If no issues are found or the text is too short/irellevant for these checks, respond ONLY with the word "OK".

    Text to analyze:
    ---
    {chunk_text}
    ---
    """)

    try:
        response = model.generate_content(prompt, request_options={"timeout": 120}, generation_config=genai.GenerationConfig(temperature=0))

        if hasattr(response, 'text'):
            response_text = response.text.strip()
            print(f"DEBUG_AI: Raw AI response for general analysis: {response_text[:200]}...") # Added debug print

            if response._result.prompt_feedback and response._result.prompt_feedback.block_reason:
                print(f"WARNING: General AI analysis blocked: {response._result.prompt_feedback.block_reason}")
                return "[AI Analysis blocked by safety filter]"

            # Check if the response text is just "OK" (case-insensitive)
            if response_text.upper() == "OK":
                print("INFO: General AI analysis: No issues found in chunk.")
                return "OK" # Return "OK" to indicate no issues found

            # This 'if' block handles cases where the AI returned text that is not "OK"
            if response_text:
                return response_text.strip()
            else:
                print("INFO: General AI analysis returned no textual response.")
                if response._result.candidates:
                    for candidate in response._result.candidates:
                        if candidate.finish_reason:
                            if candidate.finish_reason == 'SAFETY':
                                print("WARNING: General AI analysis blocked (finish reason).")
                                return "[AI Analysis blocked by safety filter]"
                        if hasattr(candidate, 'safety_ratings') and candidate.safety_ratings:
                            for rating in candidate.safety_ratings:
                                if hasattr(rating, 'probability') and rating.probability > 0.5:
                                    print("WARNING: General AI analysis blocked (safety rating).")
                                    return "[AI Analysis blocked by safety filter]"
                return "[No textual response or analysis problem from AI]"

        else:
            print("INFO: General AI analysis response object had no text attribute.")
            return "[No textual response or analysis problem from AI]"


    except Exception as e:
        if "API key not valid" in str(e):
            print(f"ERROR: While calling the Gemini API: API key invalid. {e}")
            return f"[AI Analysis error: API key invalid. {e}]"
        else:
            print(f"ERROR: While calling the Gemini API: {e}")
            print("Check your internet connection, API key, and API status.")
            return f"[AI Analysis error: {e}]"

# --- Helper function to find AI quote in paragraph objects and their original indices ---
# This function now receives the list of (paragraph_object, original_index, paragraph_element) tuples for the chunk
def find_quote_paragraph(quote_text, chunk_paragraphs_info):
    """Attempts to find the paragraph object, its original index, and element for a given quote within chunk paragraphs_info."""
    quote_text_lower_stripped = quote_text.strip().lower()
    found_locations = []

    for para, original_idx, para_element in chunk_paragraphs_info:
        para_text_lower_stripped = para.text.strip().lower()
        # Search if the stripped, lowercase quote appears in the stripped, lowercase paragraph text
        if quote_text_lower_stripped in para_text_lower_stripped:
            found_locations.append({"paragraph": para, "original_index": original_idx, "element": para_element})

    if len(found_locations) == 1:
        # Unique match found
        return found_locations[0]['paragraph'], found_locations[0]['original_index'], found_locations[0]['element'], "unique"
    elif len(found_locations) > 1:
        # Multiple matches found
        print(f"DEBUG_AI_PLACE: Found multiple locations ({len(found_locations)}) for quote: '{quote_text[:50]}...' in chunk.")
        # For now, we use the first match, and mark it as ambiguous.
        return found_locations[0]['paragraph'], found_locations[0]['original_index'], found_locations[0]['element'], "ambiguous"
    else:
        # No match found
        print(f"DEBUG_AI_PLACE: No location found for quote: '{quote_text[:50]}...' in chunk.")
        return None, -1, None, "not_found"


# --- Function to parse AI response and collect errors (adapted for individual critiques) ---
# This function now receives the list of (paragraph_object, original_index, paragraph_element) tuples for the chunk
def parse_gemini_response(document, response_text, chunk_paragraphs_info, chunk_start_original_idx, chunk_end_original_idx):
    """Parses AI response into individual critique errors, attempting to place them."""
    errors = []
    individual_ai_intro = "Feedback op tekst: " # Changed prefix
    unplaced_critiques = [] # To collect critiques that could not be placed

    if response_text.strip().upper() == "OK":
        print("DEBUG: AI response was OK, no critiques to parse.")
        return [] # No critiques found

    # Handle analysis errors like blockages or no text response
    if response_text.startswith("[AI Analysis skipped]") or response_text.startswith("[No textual response]") or response_text.startswith("[AI Analysis error:"):
        # Use the start of the chunk for the error location
        error_location_para = document.paragraphs[chunk_start_original_idx] if len(document.paragraphs) > chunk_start_original_idx else document
        errors.append({"location": error_location_para, "original_index": chunk_start_original_idx, "element": error_location_para._element if hasattr(error_location_para, '_element') else None, "message": f"Opmerking: Probleem met AI analyse voor dit stuk tekst (ongeveer vanaf paragraaf {chunk_start_original_idx}): {response_text}"})
        print(f"DEBUG: Parsed bundled AI problem message for chunk start {chunk_start_original_idx}.")
        return errors


    lines = response_text.strip().split('\n')
    # Updated regex to match double quotes instead of triple backticks
    critique_pattern = re.compile(r"\s*\"(.*?)\"\s*:\s*(.*)", re.DOTALL)

    for line in lines:
        line = line.strip()
        if not line: continue # Skip empty lines

        match = critique_pattern.match(line)
        if match:
            quoted_text = match.group(1).strip()
            explanation = match.group(2).strip()

            # Try to find the paragraph for this quote within the chunk's paragraphs
            # Now pass the list of (paragraph, index, element) tuples
            found_para, original_idx, found_element, status = find_quote_paragraph(quoted_text, chunk_paragraphs_info)

            if found_para:
                 # Critique successfully linked to a specific paragraph
                 # Store the critique in a format that add_error_to_document can parse for display
                 # Ensure message_text uses double quotes
                 message_text = f"{individual_ai_intro}\"" + quoted_text + "\" : " + explanation
                 if status == "ambiguous":
                      message_text += " [Plaatsing mogelijk ambigu - controleer]" # Add note for ambiguity

                 errors.append({
                      "location": found_para, # Link to the specific paragraph object
                      "original_index": original_idx, # Store the original index
                      "element": found_element, # Store the element
                      "message": message_text
                 })
                 print(f"DEBUG: Parsed individual AI critique for quote '{quoted_text[:50]}...' linked to original index {original_idx}.")


            else:
                 # Quote not found, add to list of unplaceable critiques
                 # Ensure unplaced_critiques uses double quotes
                 unplaced_critiques.append(f"\"" + quoted_text + "\" : " + explanation)
                 print(f"DEBUG: Parsed individual AI critique for quote '{quoted_text[:50]}...' could not be placed.")

        else:
            # Line did not match expected critique format - report as parsing error
            # Link the parsing error to the beginning of the chunk
            error_location_para = document.paragraphs[chunk_start_original_idx] if len(document.paragraphs) > chunk_start_original_idx else document
            errors.append({
                "location": error_location_para, # Link to the chunk's starting paragraph object
                "original_index": chunk_start_original_idx, # Store the original index
                "element": error_location_para._element if hasattr(error_location_para, '_element') else None, # Store the element if available
                "message": f"Opmerking: [AI Parseerfout] Kon kritiekregel niet verwerken in chunk start {chunk_start_original_idx}: '{line[:100]}...'"
            })
            print(f"DEBUG: Parsed AI parse error for chunk start {chunk_start_original_idx}: '{line[:50]}...'")


    # If there are unplaceable critiques, add them as a separate block, linked to the beginning of the chunk
    if unplaced_critiques:
        unplaced_message_header = f"AI analyse voor dit stuk tekst (kon {len(unplaced_critiques)} opmerking(en) niet plaatsen):"
        unplaced_message_content = "\n\n" + "\n".join(unplaced_critiques) # Use double newline for separation

        unplaced_location_para = document.paragraphs[chunk_start_original_idx] if len(document.paragraphs) > chunk_start_original_idx else document
        errors.append({
             "location": unplaced_location_para, # Link to the chunk's starting paragraph object
             "original_index": chunk_start_original_idx, # Store the original index
             "element": unplaced_location_para._element if hasattr(unplaced_location_para, '_element') else None, # Store the element if available
             "message": unplaced_message_header + unplaced_message_content # Use the multi-line format for the unplaceable block
        })
        print(f"DEBUG: Parsed bundled unplaced AI critiques ({len(unplaced_critiques)}) for chunk start {chunk_start_original_idx}.")


    return errors


# --- Helper function to find the index of a paragraph for sorting ---
# This function now uses the original index if available in the error structure
def get_paragraph_index(error_item, doc_obj):
    """Gets the original index of the paragraph associated with an error item for sorting."""
    # Check if the error item is a dictionary and has an 'original_index' key
    if isinstance(error_item, dict) and 'original_index' in error_item and error_item['original_index'] is not None:
        return error_item['original_index']
    # Fallback for errors that are not linked to a specific paragraph with original_index
    # This includes document-level errors or errors where placement failed.
    # We can sort these to the beginning or end. Let's sort them to the end for now.
    return len(doc_obj.paragraphs) + 1000


# --- Function to add an error message to the document ---
# This function is crucial for making the feedback visible in the Word document.
def add_error_to_document(document, message, error_info):
    """
    Adds an error message as a new paragraph directly after the paragraph where the error was detected,
    using XML manipulation for more reliable insertion.
    error_info is expected to be a dictionary containing 'location' (paragraph object),
    'original_index' (int), and 'element' (lxml element).
    """
    new_paragraph = document.add_paragraph()

    # Define common styles
    blue_color = RGBColor(0x00, 0x00, 0xFF) # Blue
    black_color = RGBColor(0x00, 0x00, 0x00) # Black
    font_size_pt = Pt(10)

    # Handle document-level errors first (they are added at the end)
    if error_info.get('original_index', -1) == -1 or error_info.get('location') is document:
        run = new_paragraph.add_run(f"Opmerking: [Algemeen] {message}")
        run.font.color.rgb = blue_color
        run.font.bold = True # Prefix is bold
        run.font.size = font_size_pt
        run.font.italic = False
        print(f"DEBUG: Added document-level error message: '{message[:100]}...' at the end.")
        return

    # Handle fallback for missing element (also added at the end)
    target_paragraph = error_info['location']
    target_element = error_info['element']
    if target_element is None:
        print(f"WARNING: Could not place comment after paragraph {error_info['original_index']} ('{target_paragraph.text.strip()[:50]}...') because '_element' is missing. Adding to the end.")
        run = new_paragraph.add_run(f"Opmerking: [Plaatsingsfout - na index {error_info['original_index']}] {message}")
        run.font.color.rgb = blue_color
        run.font.bold = True # Prefix is bold
        run.font.size = font_size_pt
        run.font.italic = False
        return

    # Now, process comments that are inserted after a specific paragraph
    ai_critique_prefix = "Feedback op tekst: "
    ai_general_prefix = "Inhoudelijke feedback: "
    layout_prefix = "Opmerking: " # This is the new prefix for layout comments

    quote_start_char = "\""
    quote_end_char = "\""

    if message.startswith(ai_critique_prefix):
        # Attempt to parse the AI critique message for multi-color formatting
        # Example format: Feedback op tekst:"quoted text" : explanation
        try:
            # Find the start of the quoted text (after "Feedback op tekst:")
            remaining_message = message[len(ai_critique_prefix):]
            quote_start_idx = remaining_message.find(quote_start_char)
            quote_end_idx = remaining_message.find(quote_end_char, quote_start_idx + 1)

            if quote_start_idx != -1 and quote_end_idx != -1:
                # Part 1: "Feedback op tekst: "
                run_prefix = new_paragraph.add_run(ai_critique_prefix)
                run_prefix.font.color.rgb = blue_color
                run_prefix.font.bold = True # Prefix is bold
                run_prefix.font.italic = False
                run_prefix.font.size = font_size_pt

                # Part 2: Opening quote char "
                run_open_quote = new_paragraph.add_run(quote_start_char)
                run_open_quote.font.color.rgb = black_color
                run_open_quote.font.bold = False
                run_open_quote.font.italic = False
                run_open_quote.font.size = font_size_pt

                # Part 3: Quoted text
                quoted_text = remaining_message[quote_start_idx + 1 : quote_end_idx]
                run_quote_text = new_paragraph.add_run(quoted_text)
                run_quote_text.font.color.rgb = black_color
                run_quote_text.font.bold = False
                run_quote_text.font.italic = False
                run_quote_text.font.size = font_size_pt

                # Part 4: Closing quote char "
                run_close_quote = new_paragraph.add_run(quote_end_char)
                run_close_quote.font.color.rgb = black_color
                run_close_quote.font.bold = False
                run_close_quote.font.italic = False
                run_close_quote.font.size = font_size_pt

                # Part 5: " : " and explanation
                explanation_part = remaining_message[quote_end_idx + 1:]
                run_explanation = new_paragraph.add_run(explanation_part)
                run_explanation.font.color.rgb = blue_color
                run_explanation.font.bold = False # Explanation is not bold
                run_explanation.font.italic = False
                run_explanation.font.size = font_size_pt
            else:
                # Fallback if parsing fails for this specific AI format, treat as general AI message
                run = new_paragraph.add_run(message)
                run.font.color.rgb = blue_color
                run.font.bold = True # Prefix is bold
                run.font.italic = False
                run.font.size = font_size_pt
        except Exception as e:
            print(f"WARNING: Error parsing AI critique for formatting: {e}. Adding as general AI comment.")
            run = new_paragraph.add_run(message)
            run.font.color.rgb = blue_color
            run.font.bold = True # Prefix is bold
            run.font.italic = False
            run.font.size = font_size_pt

    elif message.startswith(ai_general_prefix) or message.startswith("AI analyse voor dit stuk tekst"): # General AI messages
        run = new_paragraph.add_run(message)
        run.font.color.rgb = blue_color
        run.font.bold = True # Prefix is bold
        run.font.italic = False
        run.font.size = font_size_pt

    elif message.startswith(layout_prefix): # Layout/structure comments (e.g., "Opmerking:")
        run = new_paragraph.add_run(message)
        run.font.color.rgb = blue_color
        run.font.bold = True # Prefix is bold
        run.font.size = font_size_pt
        run.font.italic = False

    else: # Fallback for any other unexpected message format, treat as general layout comment
        run = new_paragraph.add_run(message)
        run.font.color.rgb = blue_color
        run.font.bold = True # Prefix is bold
        run.font.size = font_size_pt
        run.font.italic = False


    # Get the XML element of the new paragraph
    new_paragraph_element = new_paragraph._element

    # Get the parent of the target paragraph element (usually the body)
    parent = target_element.getparent()

    # Find the index of the target paragraph element within its parent
    index = parent.index(target_element)

    # Insert the new paragraph element directly after the target paragraph element
    parent.insert(index + 1, new_paragraph_element)

    print(f"DEBUG: Inserted comment after paragraph with original index {error_info['original_index']} using XML manipulation.")


# --- Helper function to find the index of a paragraph for sorting ---
# This function now uses the original index if available in the error structure
def get_paragraph_index(error_item, doc_obj):
    """Gets the original index of the paragraph associated with an error item for sorting."""
    # Check if the error item is a dictionary and has an 'original_index' key
    if isinstance(error_item, dict) and 'original_index' in error_item and error_item['original_index'] is not None:
        return error_item['original_index']
    # Fallback for errors that are not linked to a specific paragraph with original_index
    # This includes document-level errors or errors where placement failed.
    # We can sort these to the beginning or end. Let's sort them to the end for now.
    return len(doc_obj.paragraphs) + 1000


# --- Main script to process the document ---

# Ask the user for the file name
input_doc_filename = input("Enter the name of the document (e.g., 'document.docx'): ")
input_doc_path = os.path.join(os.getcwd(), input_doc_filename)

# Generate the output file name
name_part, ext_part = os.path.splitext(input_doc_filename)
output_doc_path = os.path.join(os.getcwd(), f"{name_part}_checked{ext_part}")

print(f"Starting check of '{input_doc_path}'...")
print("INFO: This is the version of the script with section-aware AI analysis (Fix Element Insertion).") # Added version indicator

doc = None
try:
    doc = Document(input_doc_path)
    print(f"Document '{input_doc_path}' loaded successfully.")
except (FileNotFoundError, PackageNotFoundError) as e:
    print(f"Error: '{input_doc_path}' not found. Ensure the document is in the same folder as this script and the name is spelled correctly.")
    try:
        with open("error_log.txt", "a") as log_file:
            log_file.write(f"{datetime.datetime.now()} - Error loading document '{input_doc_path}': {e}\n")
    except Exception as log_e:
        pass
    exit()
except Exception as e:
    print(f"An unexpected error occurred while loading or creating the document: {e}")
    try:
        with open("error_log.txt", "a") as log_file:
            log_file.write(f"{datetime.datetime.now()} - Unexpected error loading document '{input_doc_path}': {e}\n")
    except Exception as log_e:
        pass
    exit()


all_errors = []
ik_already_found = False # State variable for the 'Ik' check
last_legal_citation_error_para_idx = -1000 # State variable for legal citation limit

# --- Perform checks per paragraph and per section ---
print("\nStarting checks...")

# Title check (document-wide)
if is_check_active_for_section('title_rule', 'all'):
    title_errors, _ = check_title_rule(doc)
    for error in title_errors:
        if isinstance(error['location'], Paragraph):
            error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
            error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
        elif error['location'] is doc:
            error['original_index'] = -1 # Indicate document level
            error['element'] = None # No element for document level
        all_errors.append(error)


current_section = None
section_start_index = 0
previous_chapter_number = None # To track chapter changes for page breaks

# Updated regex to be more flexible with spacing and punctuation after the number
heading_detection_pattern = re.compile(r'^\s*(\d+(\.\d+)*)[\s\.\:]?', re.IGNORECASE)

for idx, para in enumerate(doc.paragraphs):
    text_content = para.text.strip()
    style_name = para.style.name if para.style else "None"

    # Detect section headings (numbering like 1, 2, 3.1, etc.)
    section_match = heading_detection_pattern.search(text_content)
    if section_match:
        detected_section_number = section_match.group(1)
        print(f"DEBUG: Potential Section Heading detected by pattern at index {idx}: Style='{style_name}', Text='{text_content[:50]}', Detected Number='{detected_section_number}'")

        # Determine if it's a new top-level chapter (e.g., '1', '2', '3')
        is_top_level_chapter = False
        chapter_number_match = re.match(r'^(\d+)$', detected_section_number)
        if chapter_number_match:
            current_chapter_number = int(chapter_number_match.group(1))
            if previous_chapter_number is not None and current_chapter_number > previous_chapter_number:
                is_top_level_chapter = True
            previous_chapter_number = current_chapter_number
        else:
            # If it's a sub-section (e.g., 1.1), reset previous_chapter_number to allow new chapter detection later
            # This handles cases like 1.1 -> 2.1
            if '.' in detected_section_number:
                main_chapter_num = int(detected_section_number.split('.')[0])
                if previous_chapter_number is not None and main_chapter_num > previous_chapter_number:
                    is_top_level_chapter = True
                previous_chapter_number = main_chapter_num


        # Apply page break for new top-level chapters (excluding the very first paragraph)
        if is_top_level_chapter and idx > 0:
            pPr = para._element.get_or_add_pPr()
            pB = OxmlElement('w:pageBreakBefore')
            if pPr.find(qn('w:pageBreakBefore')) is None:
                pPr.append(pB)
                print(f"INFO: Added page break before new chapter: {detected_section_number} at index {idx}.")


        # Process the previous section before moving to the next
        if current_section is not None:
            print(f"\nDEBUG: End of section {current_section} (before index {idx}). Processing section-specific checks...")
            section_end_index = idx

            # --- Section-specific checks here ---

            # AI Analysis for 1.1 and 1.2
            if current_section in ['1.1', '1.2'] and is_check_active_for_section('ai_1_1_1_2', current_section):
                 # We call this check when we detect 1.1, and pass the end of 1.2
                 # Ensure this check is only called at the start of 1.1
                 if current_section == '1.1':
                      section_1_1_index = section_start_index
                      # Find the end of section 1.2 (start of 1.3 or next numbered heading)
                      end_of_1_2_section_index = find_next_numbered_heading(section_1_1_index, doc)

                      if end_of_1_2_section_index > section_1_1_index:
                           ai_1_1_1_2_errors = analyze_section_1_1_1_2_content_in_section(doc, section_1_1_index, end_of_1_2_section_index)
                           all_errors.extend(ai_1_1_1_2_errors)
                      else:
                           print("DEBUG: Section 1.1 or 1.2 has no content for AI analysis.")


            # AI Analysis for 1.3 questions
            if current_section == '1.3' and is_check_active_for_section('ai_1_3_questions', current_section):
                 section_1_3_index = section_start_index
                 end_of_1_3_section_index = find_next_numbered_heading(section_1_3_index, doc)
                 if end_of_1_3_section_index > section_1_3_index:
                      ai_1_3_errors = analyze_section_1_3_questions_content_in_section(doc, section_1_3_index, end_of_1_3_section_index)
                      all_errors.extend(ai_1_3_errors)
                 else:
                      print("DEBUG: Section 1.3 has no content for AI analysis of questions.")

            # AI Analysis for Chapter 2
            if current_section == '2' and is_check_active_for_section('ai_chapter_2_analysis', current_section):
                 section_2_index = section_start_index
                 end_of_2_section_index = find_next_numbered_heading(section_2_index, doc)
                 if end_of_2_section_index > section_2_index:
                      ai_chapter_2_errors = analyze_chapter_2_content_in_section(doc, section_2_index, end_of_2_section_index)
                      all_errors.extend(ai_chapter_2_errors)
                 else:
                      print("DEBUG: Chapter 2 has no content for AI analysis.")


            # Sub-questions structure check (looks for "Deelvragen" heading within the section)
            if is_check_active_for_section('deelvragen_structure', current_section):
                 deelvragen_para_in_section, deelvragen_index_in_section = find_paragraph_by_text(doc, "Deelvragen")
                 if deelvragen_index_in_section != -1 and deelvragen_index_in_section >= section_start_index and deelvragen_index_in_section < section_end_index:
                      end_of_deelvragen_subsection_index = find_next_numbered_heading(deelvragen_index_in_section, doc)
                      if end_of_deelvragen_subsection_index > deelvragen_index_in_section:
                            deelvragen_errors = check_deelvragen_structure_in_section(doc, deelvragen_index_in_section, end_of_deelvragen_subsection_index)
                            for error in deelvragen_errors:
                                if isinstance(error['location'], Paragraph):
                                    error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                                    error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
                                all_errors.append(error)
                      else:
                           print(f"DEBUG: 'Deelvragen' header found at index {deelvragen_index_in_section} but no content paragraphs followed in its subsection within section {current_section}.")
                 else:
                      print(f"DEBUG: 'Deelvragen' header not found within section {current_section} boundaries ({section_start_index}-{section_end_index}).")


            # Check 1.3 intro sentences (if 1.3 is the current section)
            if current_section == '1.3' and is_check_active_for_section('section_1_3_intro', current_section):
                 section_1_3_index = section_start_index
                 section_end_index_1_3 = find_next_numbered_heading(section_1_3_index, doc)
                 if section_end_index_1_3 > section_1_3_index:
                      intro_errors = check_section_1_3_intro_in_section(doc, section_1_3_index, section_end_index_1_3)
                      for error in intro_errors:
                          if isinstance(error['location'], Paragraph):
                              error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                              error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
                          all_errors.append(error)
                 else:
                      print("DEBUG: Section 1.3 has no content for intro check.")


            # Check 1.4 content (if 1.4 is the current section)
            if current_section == '1.4' and is_check_active_for_section('section_1_4_content', current_section):
                 section_1_4_index = section_start_index
                 section_end_index_1_4 = find_next_numbered_heading(section_1_4_index, doc)
                 if section_end_index_1_4 > section_1_4_index:
                      content_1_4_errors = check_section_1_4_content_in_section(doc, section_1_4_index, section_end_index_1_4)
                      for error in content_1_4_errors:
                          if isinstance(error['location'], Paragraph):
                              error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                              error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
                          all_errors.append(error)
                 else:
                      print("DEBUG: Section 1.4 has no content for content check.")


            # General AI analysis per chunk within the section
            if is_check_active_for_section('ai_general_analysis', current_section) and model is not None:
                 print(f"\nINFO: Starting general AI analysis for section {current_section} (paragraphs {section_start_index}-{section_end_index-1})...")
                 # Filter paragraphs for this section, excluding comments
                 section_paragraphs_info = []
                 for p_idx in range(section_start_index, section_end_index):
                      para = doc.paragraphs[p_idx]
                      if not (para.text.strip().startswith("Inhoudelijke feedback:") or para.text.strip().startswith("Feedback op tekst:") or para.text.strip().startswith("Opmerking:") or para.text.strip().startswith("AI analyse voor dit stuk tekst") or para.text.strip().startswith("DC:") or para.text.strip().startswith("[opmerking]:") or para.text.strip().startswith("AI opmerking:")):
                           if hasattr(para, '_element'):
                                section_paragraphs_info.append((para, p_idx, para._element))
                           else:
                                print(f"WARNING: Paragraph {p_idx} in section {current_section} has no '_element' attribute. Cannot reliably process it for general AI chunking.")


                 if section_paragraphs_info:
                      ai_text_chunks = chunk_paragraphs_for_ai(section_paragraphs_info)
                      print(f"INFO: Section {current_section} split into {len(ai_text_chunks)} chunks for general AI analysis.")
                      for i, chunk in enumerate(ai_text_chunks):
                          if chunk['text'].strip():
                              print(f"INFO: Analyzing AI chunk {i+1}/{len(ai_text_chunks)} in section {current_section} (original paragraphs {chunk['start_para_original_idx']}-{chunk['end_para_original_idx']})...")
                              ai_response_text = analyze_chunk_with_gemini(chunk['text'])
                              if ai_response_text.strip().upper() != "OK":
                                   chunk_errors = parse_gemini_response(document=doc, response_text=ai_response_text, chunk_paragraphs_info=chunk['paragraphs_info'], chunk_start_original_idx=chunk['start_para_original_idx'], chunk_end_original_idx=chunk['end_para_original_idx'])
                                   all_errors.extend(chunk_errors)
                              else:
                                   print(f"INFO: General AI analysis for chunk {i+1} in section {current_section} was OK. No issues found.")
                 else:
                      print(f"INFO: No relevant text found in section {current_section} for general AI analysis.")


            print(f"DEBUG: Section-specific checks for section {current_section} completed.")
        # Start of new section
        current_section = detected_section_number # Use the captured number
        section_start_index = idx
        print(f"\nDEBUG: New section detected: {current_section} at index {idx}.")

    # --- Checks per paragraph ---
    # These checks are performed for EVERY paragraph, regardless of section, if active in settings
    # The section check takes place within the is_check_active_for_section function.

    # Legal text check per paragraph (with limit)
    if is_check_active_for_section('legal_citations', current_section):
         # The limit is maintained per check call here, meaning a maximum of 1 error per 50 paragraphs *within the active sections* is reported.
         # If the check is 'all', it works across the entire document.
         if idx >= last_legal_citation_error_para_idx + 50:
              legal_errors = check_legal_citations_in_paragraph(para, idx)
              if legal_errors: # Only update last_legal_citation_error_para_idx if an error was actually reported
                   last_legal_citation_error_para_idx = idx
                   # Store location info including element for sorting
                   for error in legal_errors:
                       if isinstance(error['location'], Paragraph):
                           error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                           error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
                       all_errors.append(error)


    # First 'Ik' check per paragraph (stops after the first finding)
    if is_check_active_for_section('first_ik', current_section) and not ik_already_found:
        ik_errors, ik_already_found = check_first_ik_in_paragraph(para, idx, ik_already_found)
        # Store location info including element for sorting
        for error in ik_errors:
            if isinstance(error['location'], Paragraph):
                error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
            all_errors.append(error)


    # Paragraph length check per paragraph (now with exclusion for 1.3)
    if is_check_active_for_section('paragraph_length', current_section):
        length_errors = check_paragraph_length_in_paragraph(para, idx, current_section) # Pass current_section
        # Store location info including element for sorting
        for error in length_errors:
            if isinstance(error['location'], Paragraph):
                error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
            all_errors.append(error)


# --- Process the last section after the loop ---
# This block is only executed if at least one section has been detected
if current_section is not None:
    print(f"\nDEBUG: End of last section {current_section} (document end). Processing section-specific checks...")
    section_end_index = len(doc.paragraphs)

    # --- Section-specific checks for the last section ---

    # AI Analysis for 1.1 and 1.2 (if the last section was 1.1 or 1.2)
    if current_section in ['1.1', '1.2'] and is_check_active_for_section('ai_1_1_1_2', current_section):
         if current_section == '1.1':
              section_1_1_index = section_start_index
              end_of_1_2_section_index = find_next_numbered_heading(section_1_1_index, doc)

              if end_of_1_2_section_index > section_1_1_index:
                   ai_1_1_1_2_errors = analyze_section_1_1_1_2_content_in_section(doc, section_1_1_index, end_of_1_2_section_index)
                   all_errors.extend(ai_1_1_1_2_errors)
              else:
                   print("DEBUG: Section 1.1 or 1.2 has no content for AI analysis.")


    # AI Analysis for 1.3 questions (if the last section was 1.3)
    if current_section == '1.3' and is_check_active_for_section('ai_1_3_questions', current_section):
         section_1_3_index = section_start_index
         section_end_index_1_3 = find_next_numbered_heading(section_1_3_index, doc)
         if section_end_index_1_3 > section_1_3_index:
              ai_1_3_errors = analyze_section_1_3_questions_content_in_section(doc, section_1_3_index, end_of_1_3_section_index)
              all_errors.extend(ai_1_3_errors)
         else:
              print("DEBUG: Section 1.3 has no content for AI analysis of questions.")

    # AI Analysis for Chapter 2 (if the last section was 2)
    if current_section == '2' and is_check_active_for_section('ai_chapter_2_analysis', current_section):
         section_2_index = section_start_index
         end_of_2_section_index = find_next_numbered_heading(section_2_index, doc)
         if end_of_2_section_index > section_2_index:
              ai_chapter_2_errors = analyze_chapter_2_content_in_section(doc, section_2_index, end_of_2_section_index)
              all_errors.extend(ai_chapter_2_errors)
         else:
              print("DEBUG: Chapter 2 has no content for AI analysis.")


    # Sub-questions structure check (if the last section contained the sub-questions section)
    if is_check_active_for_section('deelvragen_structure', current_section):
         deelvragen_para_in_section, deelvragen_index_in_section = find_paragraph_by_text(doc, "Deelvragen")
         if deelvragen_index_in_section != -1 and deelvragen_index_in_section >= section_start_index and deelvragen_index_in_section < section_end_index:
              end_of_deelvragen_subsection_index = find_next_numbered_heading(deelvragen_index_in_section, doc)
              if end_of_deelvragen_subsection_index > deelvragen_index_in_section:
                    deelvragen_errors = check_deelvragen_structure_in_section(doc, deelvragen_index_in_section, end_of_deelvragen_subsection_index)
                    for error in deelvragen_errors:
                        if isinstance(error['location'], Paragraph):
                            error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                            error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
                        all_errors.append(error)
              else:
                   print(f"DEBUG: 'Deelvragen' header found at index {deelvragen_index_in_section} but no content paragraphs followed in its subsection within section {current_section}.")
         else:
              print(f"DEBUG: 'Deelvragen' header not found within section {current_section} boundaries ({section_start_index}-{section_end_index}).")


    # Check 1.3 intro sentences (if the last section was 1.3)
    if current_section == '1.3' and is_check_active_for_section('section_1_3_intro', current_section):
         section_1_3_index = section_start_index
         section_end_index_1_3 = find_next_numbered_heading(section_1_3_index, doc)
         if section_end_index_1_3 > section_1_3_index:
              intro_errors = check_section_1_3_intro_in_section(doc, section_1_3_index, section_end_index_1_3)
              for error in intro_errors:
                  if isinstance(error['location'], Paragraph):
                      error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                      error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
                  all_errors.append(error)
         else:
              print("DEBUG: Section 1.3 has no content for intro check.")


    # Check 1.4 content (if the last section was 1.4)
    if current_section == '1.4' and is_check_active_for_section('section_1_4_content', current_section):
         section_1_4_index = section_start_index
         section_end_index_1_4 = find_next_numbered_heading(section_1_4_index, doc)
         if section_end_index_1_4 > section_1_4_index:
              content_1_4_errors = check_section_1_4_content_in_section(doc, section_1_4_index, section_end_index_1_4)
              for error in content_1_4_errors:
                  if isinstance(error['location'], Paragraph):
                      error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
                      error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
                  all_errors.append(error)
         else:
              print("DEBUG: Section 1.4 has no content for content check.")


    # General AI analysis per chunk within the last section
    if is_check_active_for_section('ai_general_analysis', current_section) and model is not None:
         print(f"\nINFO: Starting general AI analysis for last section {current_section} (paragraphs {section_start_index}-{section_end_index-1})...")
         # Filter paragraphs for this section, excluding comments
         section_paragraphs_info = []
         for p_idx in range(section_start_index, section_end_index):
              para = doc.paragraphs[p_idx]
              if not (para.text.strip().startswith("Inhoudelijke feedback:") or para.text.strip().startswith("Feedback op tekst:") or para.text.strip().startswith("Opmerking:") or para.text.strip().startswith("AI analyse voor dit stuk tekst") or para.text.strip().startswith("DC:") or para.text.strip().startswith("[opmerking]:") or para.text.strip().startswith("AI opmerking:")):
                   if hasattr(para, '_element'):
                        section_paragraphs_info.append((para, p_idx, para._element))
                   else:
                        print(f"WARNING: Paragraph {p_idx} in last section {current_section} has no '_element' attribute. Cannot reliably process it for general AI chunking.")


         if section_paragraphs_info:
              ai_text_chunks = chunk_paragraphs_for_ai(section_paragraphs_info)
              print(f"INFO: Last section {current_section} split into {len(ai_text_chunks)} chunks for general AI analysis.")
              for i, chunk in enumerate(ai_text_chunks):
                  if chunk['text'].strip():
                      print(f"INFO: Analyzing AI chunk {i+1}/{len(ai_text_chunks)} in last section {current_section} (original paragraphs {chunk['start_para_original_idx']}-{chunk['end_para_original_idx']})...")
                      ai_response_text = analyze_chunk_with_gemini(chunk['text'])
                      if ai_response_text.strip().upper() != "OK":
                           chunk_errors = parse_gemini_response(document=doc, response_text=ai_response_text, chunk_paragraphs_info=chunk['paragraphs_info'], chunk_start_original_idx=chunk['start_para_original_idx'], chunk_end_original_idx=chunk['end_para_original_idx'])
                           all_errors.extend(chunk_errors)
                      else:
                           print(f"INFO: General AI analysis for chunk {i+1} in last section {current_section} was OK. No issues found.")
                  else:
                      print(f"INFO: No relevant text found in last section {current_section} for general AI analysis.")


    print(f"DEBUG: Section-specific checks for last section {current_section} completed.")

# --- Check paragraph separation (document-wide) ---
# This check is performed at the end because it needs the entire list of paragraphs.
if is_check_active_for_section('paragraph_separation', 'all'):
    separation_errors = check_paragraph_separation(doc)
    # Store location info including element for sorting
    for error in separation_errors:
        if isinstance(error['location'], Paragraph):
            error['original_index'] = list(doc.paragraphs).index(error['location']) if error['location'] in doc.paragraphs else None
            error['element'] = error['location']._element if hasattr(error['location'], '_element') else None
        all_errors.append(error)


print("\nEnd of checks.")

# --- Apply error messages to the document ---
if all_errors:
    print(f"\nComments found in '{input_doc_path}'. These will be added to the document with custom formatting.")

    # Sort errors by original index in ASCENDING order for insertion AFTER the paragraph
    # The get_paragraph_index function uses the 'original_index' from the error dict.
    all_errors.sort(key=lambda x: get_paragraph_index(x, doc), reverse=False)

    for i, error in enumerate(all_errors):
        # Pass the full error dictionary which contains location, original_index, and element
        add_error_to_document(doc, error['message'], error)

else:
    print(f"\nNo comments found in '{input_doc_path}'. The document complies with the checked rules.")

try:
    doc.save(output_doc_path)
    print(f"\nChecked document with comments saved as '{output_doc_path}'.")
except Exception as e:
    print(f"\nError saving document '{output_doc_path}': {e}")

# END OF SCRIPT
