#!/usr/bin/env python3
"""
Debug script voor Word export functionaliteit.
"""

import sqlite3
import json
import os
import sys

# Voeg src directory toe aan path
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from word_export import WordFeedbackExporter

def debug_word_export():
    """Test Word export met een bestaand document."""
    
    # Verbind met de juiste database
    conn = sqlite3.connect('instance/documents.db')
    conn.row_factory = sqlite3.Row
    
    try:
        # Haal document ID 6 op (uit de logs)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT d.*, dt.name as document_type_name 
            FROM documents d 
            JOIN document_types dt ON d.document_type_id = dt.id 
            WHERE d.id = 6
        """)
        
        document = cursor.fetchone()
        
        if not document:
            print("❌ Document ID 6 niet gevonden!")
            return
        
        print(f"📄 Test document: {document['original_filename']}")
        print(f"📊 Document type: {document['document_type_name']}")
        print(f"📁 Bestandspad: {document['file_path']}")
        
        # Check of bestand bestaat
        if not os.path.exists(document['file_path']):
            print(f"❌ Document bestand niet gevonden: {document['file_path']}")
            return
        
        # Haal analysis data op
        analysis_data = json.loads(document['analysis_data'])
        print(f"📈 Analysis data keys: {list(analysis_data.keys())}")
        
        # Check feedback
        feedback_items = analysis_data.get('feedback', [])
        sections = analysis_data.get('sections', [])
        
        print(f"🔍 Aantal feedback items: {len(feedback_items)}")
        print(f"📋 Aantal secties: {len(sections)}")
        
        if not feedback_items:
            print("❌ Geen feedback items gevonden!")
            return
        
        # Toon eerste feedback items
        print("\n📋 Eerste 3 feedback items:")
        for i, item in enumerate(feedback_items[:3]):
            print(f"  {i+1}. {item.get('criteria_name', 'Onbekend')} - {item.get('status', 'unknown')}")
            print(f"     Sectie: {item.get('section_name', 'Onbekend')}")
            print(f"     Bericht: {item.get('message', 'Geen bericht')[:100]}...")
        
        # Test Word export
        print("\n🧪 Test Word export...")
        
        # Maak export filename
        base_name = os.path.splitext(document['original_filename'])[0]
        export_filename = f"{base_name}_debug_feedback.docx"
        export_path = os.path.join('instance', 'uploads', export_filename)
        
        # Bereid feedback data voor
        feedback_data = {
            'feedback_items': feedback_items,
            'sections': sections
        }
        
        print(f"📤 Export pad: {export_path}")
        print(f"📊 Feedback data structuur:")
        print(f"   - feedback_items: {len(feedback_data['feedback_items'])} items")
        print(f"   - sections: {len(feedback_data['sections'])} secties")
        
        # Exporteer naar Word
        exporter = WordFeedbackExporter()
        
        try:
            result_path = exporter.add_feedback_to_document(
                document['file_path'], 
                feedback_data, 
                export_path
            )
            print(f"✅ Export succesvol: {result_path}")
            
            # Check of bestand is aangemaakt
            if os.path.exists(result_path):
                file_size = os.path.getsize(result_path)
                print(f"📁 Export bestand grootte: {file_size} bytes")
                
                # Test het bestand openen
                try:
                    from docx import Document
                    doc = Document(result_path)
                    print(f"📄 Export document paragrafen: {len(doc.paragraphs)}")
                    
                    # Zoek naar feedback tekst
                    feedback_found = False
                    for i, para in enumerate(doc.paragraphs):
                        if 'feedback' in para.text.lower() or '❌' in para.text or '⚠️' in para.text or '✅' in para.text:
                            print(f"🔍 Feedback gevonden in paragraaf {i}: {para.text[:100]}...")
                            feedback_found = True
                    
                    if not feedback_found:
                        print("❌ Geen feedback tekst gevonden in export document!")
                        
                        # Toon alle paragrafen
                        print("\n📄 Alle paragrafen in export document:")
                        for i, para in enumerate(doc.paragraphs[:10]):  # Eerste 10
                            print(f"  {i}: {para.text[:100]}...")
                    
                except Exception as e:
                    print(f"❌ Fout bij openen export document: {e}")
            else:
                print(f"❌ Export bestand niet aangemaakt: {result_path}")
                
        except Exception as e:
            print(f"❌ Fout bij export: {e}")
            import traceback
            traceback.print_exc()
        
    except Exception as e:
        print(f"❌ Fout: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        conn.close()

if __name__ == "__main__":
    debug_word_export() 