from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

def shade_paragraph(para, hex_color="F2F2F2"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p

def h1(doc, text):
    p = doc.add_paragraph()
    shade_paragraph(p, "1A537E")
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1a, 0x53, 0x7e)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2e, 0x75, 0xb6)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def callout(doc, text, color="FFF9C4"):
    p = doc.add_paragraph()
    shade_paragraph(p, color)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = True
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def example_comment(doc, criterion, page_ref, text):
    p = doc.add_paragraph()
    shade_paragraph(p, "EBF4FF")
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    r1 = p.add_run(f"[{criterion}]")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x1a, 0x53, 0x7e)
    r1.font.size = Pt(10)
    if page_ref:
        r2 = p.add_run(f"  {page_ref}\n")
        r2.italic = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r3 = p.add_run(text)
    r3.font.size = Pt(11)
    return p

# ── COVER ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title_p.add_run("DocuCheck")
t.bold = True; t.font.size = Pt(32)
t.font.color.rgb = RGBColor(0x1a, 0x53, 0x7e)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub_p.add_run("Addendum: EU-subsidieaanvragen")
s.font.size = Pt(20); s.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Juni 2026  |  Hermes Strategic Advisory\nUitgebreide toelichting op de grant-writing markt").font.size = Pt(11)
doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()

callout(doc,
    "Kernboodschap: DocuCheck's holistische, rubric-verankerde feedbackengine is structureel uitermate "
    "geschikt voor EU-subsidieaanvragen. De core value proposition — AI die een document leest zoals "
    "een expert-evaluator en teruggeeft met actiegerichte, verankerde, criterium-gemapte feedback — "
    "adresseert het pijnlijkste en duurste probleem in Europees onderzoek: de 80–90% voorstellen die "
    "falen niet omdat de wetenschap slecht is, maar omdat het schrijven dit niet effectief communiceert.",
    "E8F5E9")

doc.add_page_break()

# ── DEEL 1: HET EU-SUBSIDIE LANDSCHAP ──────────────────────────────────────
h1(doc, "1 — Het EU-Subsidie Landschap")

h2(doc, "De belangrijkste financieringsprogramma's")

programmes = [
    ("Horizon Europe (2021–2027)",
     "€95,5 miljard totaalbudget. Drie pijlers: Excellent Science (ERC, MSCA), Global Challenges "
     "(Clusters: gezondheid, digitaal/industrie, klimaat, voeding, veiligheid), en Innovative Europe "
     "(EIC). Het primaire target voor Nederlandse universiteiten, kennisinstellingen en bedrijven."),
    ("ERC — European Research Council",
     "Frontier research via vier schema's: Starting Grant (€1,5M, vroege carrière), Consolidator "
     "Grant (€2M), Advanced Grant (€2,5M, gevestigde onderzoekers), Synergy Grant (€10M, teams van "
     "2–4 PI's). Slagingspercentage: 8–12%. Hoogste prestige in Europese wetenschap."),
    ("MSCA — Marie Skłodowska-Curie Actions",
     "Postdoctoral Fellowships, Doctoral Networks (voorheen ITN's: €3–5M per consortium, 10–15 PhD's), "
     "Staff Exchanges. Slagingspercentage Doctoral Networks: ~15–20%."),
    ("NWO — Nederlandse Onderzoeksraad",
     "Nationaal: Open Competition (€350K–€800K), VIDI (€800K), VICI (€1,5M), Gravitation (€18,75M). "
     "NWO-criteria sluiten steeds meer aan op Horizon Europe — rubrics deels overdraagbaar."),
    ("ZonMw",
     "Gezondheidszorgonderzoek: €300M+ jaarlijks. Programma's: Nationale Wetenschapsagenda (NWA), "
     "GGz, COVID-19. Aanvraagformats sluiten aan op Horizon Europe-structuur."),
    ("Overig",
     "Interreg (grensoverschrijdend: North Sea, 2 Seas, URBACT), ESF+ (sociale inclusie, onderwijs), "
     "LIFE (milieu/klimaat: €5,4B 2021–2027), EIC Accelerator (deep tech: tot €2,5M grant + equity)."),
]
for title, text in programmes:
    bullet(doc, text, title + ": ")

h2(doc, "Horizon Europe evaluatiecriteria — de volledige breakdown")
body(doc,
    "De Europese Commissie beoordeelt voorstellen op drie toplevelcriteria, elk gescoord op een "
    "0–5-schaal (drempelwaarde: 4,0). Dit zijn precies de criteria die DocuCheck per passage aangeeft.")

criteria_data = [
    ("Criterium 1: Excellence", [
        ("Soundness of the concept", "Is de wetenschappelijke/technologische basis goed onderbouwd? Is het probleem helder gedefinieerd?"),
        ("Quality and credibility of the approach", "Is de methodologie passend? Zijn risico's geïdentificeerd en gemitigeerd?"),
        ("Clarity and pertinence of objectives", "Zijn de doelstellingen SMART (Specifiek, Meetbaar, Haalbaar, Relevant, Tijdgebonden)?"),
        ("Credibility of the work plan", "Zijn mijlpalen en deliverables realistisch? Is er logische progressie van taken naar resultaten?"),
    ]),
    ("Criterium 2: Impact", [
        ("Expected outcomes and impacts", "Hoe draagt het project bij aan de verwachte resultaten van het werkprogramma? Wetenschappelijke én maatschappelijke impact?"),
        ("Measures to maximise impact", "Is er een disseminatie- en exploitatieplan? Communicatiestrategie voor niet-specialisten?"),
        ("Open science practices", "Data Management Plan aanwezig? Open access publicatie gepland? FAIR-data principes?"),
        ("Pathways to impact", "Zijn stakeholderengagement en gebruikersadoptieplannen geloofwaardig en concreet?"),
    ]),
    ("Criterium 3: Quality & Efficiency of Implementation", [
        ("Quality and effectiveness of the work plan", "Zijn werkpakketten logisch gestructureerd? Zijn afhankelijkheden tussen taken duidelijk?"),
        ("Appropriateness of management structure", "Is projectcoördinatie geloofwaardig? Zijn governance en besluitvorming beschreven?"),
        ("Complementarity of participants", "Heeft elke partner een duidelijke rol? Zijn er competentieleemtes?"),
        ("Appropriateness of total resources", "Is het budget proportioneel aan de taken? Zijn personeelskosten verantwoord?"),
    ]),
]

for crit_title, subcriteria in criteria_data:
    h3(doc, crit_title)
    for sub_title, sub_text in subcriteria:
        bullet(doc, sub_text, sub_title + ": ")

h2(doc, "Waarom 80–90% van de aanvragen worden afgewezen")
body(doc, "De meest voorkomende faalredenen, gedocumenteerd in evaluatieverslagen van de EC en post-mortems van research support offices:")

rejection_reasons = [
    ("Onvoldoende wetenschappelijke ambitie", "Het voorstel beschrijft incrementeel werk, geen doorbraak. Geen helder 'beyond the state of the art'-argument."),
    ("Zwak impact-narratief", "Uitstekende wetenschap, maar geen antwoord op: wie profiteert, hoe, en wanneer? Pathways to impact zijn vaag of afwezig."),
    ("Methodologische hiaten", "Werkplan somt taken op zonder te verklaren hóé sleutelmethoden worden uitgevoerd of wat de terugvalpositie is bij mislukking."),
    ("Budget-incoherentie", "Kosten zonder narrative verantwoording, of personeelsbudget past niet bij de ambitie van het werkplan."),
    ("Slechte consortiumbalans", "Partners hebben geen duidelijke rol, of budgetverdeling is scheefgetrokken zonder wetenschappelijke rechtvaardiging."),
    ("Mismatch met call topic", "Wetenschappelijk uitstekend, maar adresseert de specifieke expected outcomes van het werkprogrammaonderwerp niet direct."),
    ("Schrijfkwaliteit", "Zelfs uitstekend onderzoek krijgt scores onder de drempelwaarde als het voorstel slecht gestructureerd is, jargon bevat zonder uitleg, of intern inconsistent is."),
]
for title, text in rejection_reasons:
    bullet(doc, text, title + ": ")

h2(doc, "Wie schrijft de aanvragen?")
t_writers = doc.add_table(rows=1, cols=3)
t_writers.style = 'Table Grid'
for c, txt in zip(t_writers.rows[0].cells, ["Rol", "Bijdrage", "Tijdsinvestering"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
writers = [
    ("Principal Investigator (PI)", "Leidt wetenschappelijke inhoud; schrijft Excellence en deels Impact", "80–200 uur over 3–6 maanden"),
    ("Research Support Officer (RSO)", "Advies over structuur, templates, administratieve compliance; eindredactie", "40–80 uur"),
    ("Grant Consultant (extern)", "Versterkt zwakke voorstellen; dagvergoeding €1.000–€2.500", "€5.000–€20.000 per aanvraag"),
    ("Projectmanager (consortia)", "Coördineert input van meerdere partners", "Variabel"),
]
for r in writers:
    row = t_writers.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

h2(doc, "De werkelijke kosten van een afgewezen aanvraag")
cost_items = [
    "PI-tijd: 120 uur × €80/uur (volledig belaste academische kosten) = €9.600",
    "RSO-tijd: 60 uur × €60/uur = €3.600",
    "Externe consultant (indien ingeschakeld): €8.000–€15.000",
    "Totaal per aanvraag: €13.200–€28.200 in personeelskosten",
    "Bij 15% slagingspercentage: verwachte kosten per gefi­nancierd voorstel = €88.000–€188.000 in schrijfinspanning",
]
for item in cost_items:
    bullet(doc, item)

callout(doc,
    "Een afgewezen ERC-grant kan een onderzoeker een promotiecyclus kosten. Een afgewezen MSCA Doctoral "
    "Network betekent 3–4 jaar wachten tot het consortium opnieuw kan proberen. Een mislukte EIC Accelerator "
    "aanvraag kan betekenen dat een startup zijn financieringsvenster mist. De inzet is extreem hoog.",
    "FFE8E8")

doc.add_page_break()

# ── DEEL 2: HOE DOCUCHECK PAST ──────────────────────────────────────────────
h1(doc, "2 — Hoe DocuCheck Past: De Exacte Workflow")

h2(doc, "Stap-voor-stap: van draft naar gefeedback voorstel")

steps = [
    ("Week -4 voor deadline: Draft gereed",
     "De PI rondt een werkende versie af van Part B — typisch een Word-document van 35–60 pagina's "
     "met de secties Excellence, Impact en Implementation."),
    ("Rubric selecteren (5 minuten)",
     "De RSO (of PI zelf) logt in op DocuCheck, selecteert het 'Horizon Europe RIA'-rubricpakket "
     "uit de bibliotheek, en plakt optioneel de specifieke call topic expected outcomes als extra "
     "rubriccontext in een tekstveld."),
    ("Upload en verwerking (10–15 minuten)",
     "Het concept Part B wordt geüpload. DocuCheck's holistische engine leest het volledige document "
     "zonder dat de gebruiker secties hoeft te labelen. Claude bepaalt wat elke passage doet — "
     "doelstellingen definiëren, methodologie beschrijven, impact beargumenteren, tijdlijn presenteren "
     "— en mapt dit op de relevante evaluatiesubcriteria."),
    ("Output: Word-bestand met kantlijncomments",
     "DocuCheck levert een Word-bestand met tracked-change-stijl kantlijncomments. Elke comment is "
     "verankerd aan een specifieke passage en gelabeld met het criterium (bijv. '[Excellence – "
     "Soundness of Concept]'). De comment bevat specifieke, actiegerichte feedback."),
    ("Revisiecyclus",
     "PI en RSO bekijken de comments samen. De RSO gebruikt het geannoteerde document als "
     "gestructureerde checklist. Hoge-ernst-comments worden als eerste aangepakt."),
    ("Tweede pass (optioneel, Week -2)",
     "Het herziene concept wordt opnieuw geüpload om te verifiëren of kritieke leemtes zijn gedicht."),
]
for i, (title, text) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"Stap {i}: {title}\n")
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1a, 0x53, 0x7e)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

h2(doc, "DocuCheck vs. ChatGPT — het structurele verschil")

t_vs = doc.add_table(rows=1, cols=3)
t_vs.style = 'Table Grid'
for c, txt in zip(t_vs.rows[0].cells, ["Feature", "DocuCheck", "ChatGPT"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
vs_rows = [
    ("Comments verankerd aan tekstpassages", "✅ Inline kantlijncomments", "❌ Generiek op alineaniveau"),
    ("Feedback gemapped op named rubriccriteria", "✅ [Excellence – Impact]-labels", "❌ Geen rubricbewustzijn"),
    ("Output bruikbaar voor samenwerking", "✅ Word-bestand, deelbaar PI + RSO", "❌ Chat-output moet handmatig worden gekopieerd"),
    ("Holistische lezing (cross-sectie coherentie)", "✅ Detecteert tegenstrijdigheden over secties heen", "❌ Context window fragmentatie"),
    ("Audit trail voor RSO", "✅ Gedocumenteerd, reproduceerbaar", "❌ Geen registratie"),
    ("Gekalibreerd op EC-drempelscores", "✅ Rubric reflecteert 0–5 EC-schaal", "❌ Geen drempelwaardebewustzijn"),
    ("Workflow integratie", "✅ Word-bestand direct bruikbaar in review", "❌ Handmatig overkopieren vereist"),
]
for r in vs_rows:
    row = t_vs.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

callout(doc,
    "Het kritiekste praktische verschil is verankering. Wanneer een onderzoeker feedback krijgt die zegt "
    "'je impact-sectie is zwak', weten ze niet welke zin ze moeten aanpassen. DocuCheck pinnt de comment "
    "op de exacte passage — precies zoals een menselijke expert-reviewer zou doen — waardoor revisie "
    "uitvoerbaar wordt in plaats van overweldigend.",
    "E8F4FF")

h2(doc, "Wie gebruikt DocuCheck in het grantproces, en wanneer?")
users = [
    ("PI's in draft-stadium", "Uploaden hun eigen concept op het 70%-punt om structurele leemtes te identificeren vóór ze meer tijd besteden aan een fundamenteel gebrekkig argument."),
    ("RSO's in reviewstadium", "Gebruiken de DocuCheck-output als gestructureerde reviewchecklist — vervangt het informele lezen dat ze nu doen door een rubric-gemapte analyse."),
    ("Grant consultants", "Gebruiken DocuCheck als first-pass analysetool om hun eigen review te versnellen, meer voorstellen te verwerken, betere ROI per klant."),
    ("Consortiumcoördinatoren", "Voordat het concept wordt rondgestuurd naar alle partners: DocuCheck om leemtes in partnerbeschrijvingen of budgetverantwoordingen te identificeren."),
]
for title, text in users:
    bullet(doc, text, title + ": ")

h2(doc, "De meest problematische secties van een Horizon Europe-voorstel")

problem_sections = [
    ("Impact – Pathways to Impact",
     "PI's schrijven uitstekende wetenschap maar begrijpen niet wat 'pathway' betekent in EC-termen. "
     "DocuCheck signaleert wanneer een pathway wordt beschreven als een aspiratie ('we hopen beleid te "
     "beïnvloeden') in plaats van een concreet mechanisme ('we leveren een beleidsnota aan DG SANTE in "
     "maand 18 en presenteren op het European Health Forum in 2026')."),
    ("Implementation – Budget Justification",
     "Onderzoekers lijsten getallen op zonder narrative verantwoording. DocuCheck identificeert "
     "budgettabelregels die geen corresponderende taakomschrijvingen hebben."),
    ("Excellence – Beyond State of the Art",
     "De meest kritieke paragraaf in het voorstel én het meest frequente probleem. DocuCheck signaleert "
     "wanneer de SOTA-analyse beschrijvend is in plaats van kritisch — bestaand werk samenvatten zonder "
     "te identificeren welk hiaat het voorstel opvult."),
    ("Work Plan – Task Dependencies",
     "Taken die logisch afhankelijk zijn van deliverables die later gereed zijn. DocuCheck detecteert "
     "chronologische inconsistenties in het Gantt-schema."),
]
for title, text in problem_sections:
    h3(doc, title)
    body(doc, text)

doc.add_page_break()

# ── DEEL 3: BUSINESS CASE ──────────────────────────────────────────────────
h1(doc, "3 — De Business Case: Cijfers")

h2(doc, "Volume van aanvragen")

t_vol = doc.add_table(rows=1, cols=3)
t_vol.style = 'Table Grid'
for c, txt in zip(t_vol.rows[0].cells, ["Programma", "Volume", "Nederlanders"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
vol_rows = [
    ("Horizon Europe (alle instrumenten)", "45.000–55.000/jaar (EU)", "1.800–3.300/jaar (4–6%)"),
    ("ERC (alle schema's)", "~10.000/jaar (EU)", "~600–800/jaar"),
    ("NWO (alle schema's)", "6.000–8.000/jaar (NL)", "6.000–8.000/jaar"),
    ("ZonMw", "2.000–3.000/jaar (NL)", "2.000–3.000/jaar"),
    ("Totaal NL relevant", ">10.000 aanvragen/jaar", "Alle NL"),
]
for r in vol_rows:
    row = t_vol.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

h2(doc, "Prijsmodel")

t_pricing = doc.add_table(rows=1, cols=4)
t_pricing.style = 'Table Grid'
for c, txt in zip(t_pricing.rows[0].cells, ["Tier", "Prijs", "Doelgroep", "Gebruik"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
pricing_rows = [
    ("Per aanvraag", "€150–€350/run", "Individuele PI's, kleine departementen", "Concurreert met 1 uur grantconsultant-tijd"),
    ("Departementslicensie", "€3.000–€8.000/jaar", "Research department (10–30 actieve schrijvers)", "Onbeperkte runs, RSO als beheerder"),
    ("Institutionele licentie", "€15.000–€40.000/jaar", "Volledige universiteit (100+ onderzoekers)", "Custom rubrics, API-integratie, analytics"),
    ("Consortiumpakket", "€500–€1.500/consortium", "Multi-partner consortia", "Partner-specifieke feedbacksecties"),
]
for r in pricing_rows:
    row = t_pricing.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

h2(doc, "Total Addressable Market")

t_tam = doc.add_table(rows=1, cols=3)
t_tam.style = 'Table Grid'
for c, txt in zip(t_tam.rows[0].cells, ["Markt", "Aanname", "TAM"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
tam_rows = [
    ("Nederland (vol volwassenheid)", "14 universiteiten + 30 inst., 10% penetratie @€25K + individueel", "€5–8 miljoen/jaar"),
    ("Europa (3% penetratie)", "1.200 universiteiten + 500 instituten, €20K ACV gemiddeld", "~€100 miljoen/jaar"),
    ("Europa (10% penetratie)", "Zelfde basis, hogere adoptie", "~€340 miljoen/jaar"),
    ("Realistisch Jaar 1–3 (NL + BE + DE + SE)", "Early adopters, €20K ACV", "€10–30 miljoen/jaar"),
]
for r in tam_rows:
    row = t_tam.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

callout(doc,
    "Bij €25.000/jaar per instelling heeft DocuCheck slechts 12 Nederlandse universitaire klanten nodig "
    "om €300.000 ARR te genereren. Bij Europese schaal is de opportuniteit tientallen miljoenen.",
    "E8F5E9")

h2(doc, "Waarom zou een RSO betalen in plaats van ChatGPT gebruiken?")
reasons = [
    ("Aansprakelijkheid en documentatie", "Research support offices zijn professionele diensten. Ze moeten documenteren welk reviewproces is gevolgd. Een ChatGPT-gesprek laat geen audit trail; een DocuCheck Word-bestand met tijdgestempelde, rubric-gemapte comments doet dat wel."),
    ("Rubric-fideliteit", "ChatGPT kent de EC-scorerubric niet tenzij bij elk gesprek uitgebreid geprompt. DocuCheck embedt de rubric op systeemniveau — elke run is gekalibreerd op dezelfde evaluatiestandaard."),
    ("Samenwerkingsformaat", "RSO's werken met PI's via tracked changes in Word. Een ChatGPT-antwoord kan niet direct in de collaboratieve workflow worden ingevoegd. DocuCheck-output ís de workflow."),
    ("Tijdseconomie", "Een RSO die nu 4 uur nodig heeft om een 50-pagina voorstel te lezen, kan DocuCheck gebruiken als pre-processing in 15 minuten, dan 90 minuten besteden aan alleen de geflagde problemen. De tool verdient zichzelf terug in één bespaarde werkdag per aanvraag."),
    ("Institutionele geloofwaardigheid", "Een universitair research support office dat 'AI-assisted proposal review met rubric-gemapte feedback' kan aanbieden, onderscheidt zich in tevredenheidsonderzoeken en rankings."),
]
for title, text in reasons:
    bullet(doc, text, title + ": ")

doc.add_page_break()

# ── DEEL 4: CONCRETE SCENARIO'S ───────────────────────────────────────────
h1(doc, "4 — Drie Concrete Scenario's")

# Scenario 1
h2(doc, "Scenario 1: ERC Starting Grant — Cognitieve Neurowetenschappen")

body(doc,
    "Dr. Amara Osei, Assistent Professor aan de Radboud Universiteit, dient een ERC Starting Grant "
    "in op het gebied van predictive coding in auditieve perceptie. Ze heeft een sterk publicatierecord "
    "maar heeft moeite de 'groundbreaking'-ambitie van haar onderzoek duidelijk te communiceren.")

p = doc.add_paragraph()
p.add_run("Wat ze uploadt: ").bold = True
p.add_run("35-pagina Part B1 (Wetenschappelijk Voorstel) + ERC StG 'Groundbreaking Nature and Ambition'-rubric (gewicht: 40% van totaalscore)")

p2 = doc.add_paragraph()
p2.add_run("Feedback die DocuCheck genereert:").bold = True
p2.paragraph_format.space_after = Pt(4)

example_comment(doc,
    "ERC – Groundbreaking Nature",
    "Pagina 4, paragraaf 2:",
    "U beschrijft het huidige debat tussen predictive coding- en feature-detectormodellen als 'lopend'. "
    "Deze formulering onderschat de kans — reviewers willen zien dat uw aanpak dit debat zal "
    "OPLOSSEN, niet slechts bijdragen. Herschrijf deze paragraaf met een stoutmoedige, falsifieerbare "
    "claim over wat uw project definitief zal vaststellen.")

example_comment(doc,
    "ERC – Ambition",
    "Pagina 11, Taak 2:",
    "Het hier beschreven EEG-paradigma is methodologisch standaard voor het vakgebied. ERC-evaluatoren "
    "verwachten dat ERC-gefinancierd onderzoek verder gaat dan gevestigde methoden. Overweeg een zin "
    "toe te voegen die verklaart waarom geen enkel bestaand paradigma uw vraag kan beantwoorden en "
    "waarom uw nieuwe combinatie van MEG + computationele modellering het enige haalbare pad is.")

body(doc,
    "Uitkomst: Dr. Osei herziet de ambitieformulering, scherpt haar 'beyond state of the art'-claims "
    "aan, en adresseert de methodologische nieuwheid. Haar voorstel ontvangt een score van 4,5/5 op "
    "Scientific Ambition in de daadwerkelijke evaluatie. Het voorstel wordt gefinancierd.")

doc.add_paragraph()

# Scenario 2
h2(doc, "Scenario 2: Horizon Europe RIA — Duurzame Landbouw Consortium")

body(doc,
    "Het Research Support Office van Wageningen University coördineert een 12-partner Horizon Europe "
    "Research and Innovation Action op het gebied van precisie-fermentatie voor duurzame eiwitproductie. "
    "Het 58-pagina voorstel is geschreven door 12 partners uit 7 landen, 3 weken voor de deadline.")

p = doc.add_paragraph()
p.add_run("Wat ze uploaden: ").bold = True
p.add_run("58-pagina Part B (volledig voorstel) + Horizon Europe Cluster 6 (Voeding, Bieconomie) evaluatierubric + call topic expected outcomes")

p2 = doc.add_paragraph()
p2.add_run("Feedback die DocuCheck genereert:").bold = True
p2.paragraph_format.space_after = Pt(4)

example_comment(doc,
    "Implementation – Consortium Complementarity",
    "Pagina 43, Partner rollen-tabel:",
    "Partner 7 (MKB, Denemarken) is vermeld als bijdragend aan WP3 (Opschaling) en WP5 (Verspreiding) "
    "maar ontvangt 8% van het totale budget. De taken toegeschreven aan deze partner — het uitvoeren "
    "van pilot-fermentatie-experimenten — worden in WP3 beschreven als vereisend €450K aan "
    "apparatuurkosten. Verduidelijk of Partner 7 bestaande infrastructuur gebruikt (vermeld dit dan) "
    "of dat aanvullende budgetherverdeling nodig is.")

example_comment(doc,
    "Impact – Open Science",
    "Pagina 51:",
    "Het Data Management Plan stelt dat datasets 'op verzoek beschikbaar worden gesteld.' EC-evaluatoren "
    "voor Horizon Europe 2023+-calls verwachten een toezegging aan FAIR-dataprincipes en open "
    "datarepositoria. Vervang 'op verzoek' door een specifieke repository (bijv. Zenodo, PANGAEA) "
    "en een depositietijdlijn.")

example_comment(doc,
    "Excellence – Work Plan Coherence",
    "Pagina 27, WP2 Taak 2.3:",
    "Taak 2.3 verwijst naar een 'standaard inoculatieprotocol te ontwikkelen in WP1.' Echter, "
    "WP1's deliverable D1.2 (protocoldocument) is gepland voor Maand 12, en Taak 2.3 begint in "
    "Maand 10. Dit 2-maanden-hiaat blokkeert Taak 2.3. Vervroeg D1.2 naar Maand 9 of voeg een "
    "go/no-go mijlpaal toe die Taak 2.3 conditioneert op de beschikbaarheid van het protocol.")

body(doc,
    "Uitkomst: Het RSO lost drie kritieke coherentieproblemen op vóór indiening. Het voorstel "
    "ontvangt een totaalscore van 14,5/15 en wordt gefinancierd in de eerste evaluatieronde.")

doc.add_paragraph()

# Scenario 3
h2(doc, "Scenario 3: NWO VIDI — Computationele Sociale Wetenschappen")

body(doc,
    "Dr. Lena Bergström, Universitair Hoofddocent aan de Universiteit Utrecht, dient een NWO VIDI-grant "
    "(€800K) in over algoritmische bias in rechterlijke beslissingsondersteuningssystemen.")

p = doc.add_paragraph()
p.add_run("Wat ze uploadt: ").bold = True
p.add_run("25-pagina VIDI-aanvraag + NWO VIDI-evaluatiecriteria (Wetenschappelijke Kwaliteit, Originaliteit, Wetenschappelijke en Maatschappelijke Impact, Haalbaarheid)")

p2 = doc.add_paragraph()
p2.add_run("Feedback die DocuCheck genereert:").bold = True
p2.paragraph_format.space_after = Pt(4)

example_comment(doc,
    "Haalbaarheid – Datatoegang",
    "Pagina 14:",
    "U stelt voor vonnisgegevens van Nederlandse rechtbanken te analyseren (2015–2023). Deze sectie "
    "vermeldt niet of toegang tot deze data is verleend of in behandeling is. NWO-evaluatoren "
    "markeren datatoegang consequent als haalbaarheidrisico. Voeg een zin toe die bevestigt dat u "
    "een Letter of Intent heeft van de Raad voor de Rechtspraak, of beschrijf uw plan om toegang "
    "te verkrijgen binnen de eerste 6 maanden van het project.")

example_comment(doc,
    "Maatschappelijke Impact",
    "Pagina 19:",
    "De sectie maatschappelijke impact noemt 'beleidsmakers informeren' als verwachte uitkomst. Dit "
    "is onvoldoende specifiek voor VIDI-reviewers. Geef aan welke beleidsactoren u betrekt "
    "(bijv. Ministerie van Justitie, AI Act-implementatieorganen), via welk mechanisme (adviesrol, "
    "stakeholderworkshop, beleidsnota), en op welk punt in de projecttijdlijn.")

example_comment(doc,
    "Originaliteit",
    "Pagina 5:",
    "De bewering dat 'geen studie algoritmische bias in Nederlandse rechtssystemen heeft onderzocht' "
    "is sterk. Substantieer dit met een beknopte systematische literatuuroverzichtsamenvatting of "
    "verzacht de bewering. Evaluatoren zullen sceptisch zijn over een absolute afwezigheidsclaim "
    "en kunnen hiervoor penaliseren als ze verwant werk kennen.")

body(doc,
    "Uitkomst: Dr. Bergström voegt een bevestigde Letter of Intent toe van de Raad voor de "
    "Rechtspraak, concretiseert het beleidsengagementplan, en matig de originaliteitsclaim met "
    "een beknopt literatuuroverzicht. Het voorstel scoort boven de drempelwaarde op alle vier "
    "de criteria.")

doc.add_page_break()

# ── DEEL 5: WAT ER GEBOUWD MOET WORDEN ────────────────────────────────────
h1(doc, "5 — Wat Er Gebouwd Moet Worden")

h2(doc, "Verschillen ten opzichte van de scriptie-use-case")

t_diff = doc.add_table(rows=1, cols=3)
t_diff.style = 'Table Grid'
for c, txt in zip(t_diff.rows[0].cells, ["Dimensie", "Scriptie", "EU-subsidie"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
diff_rows = [
    ("Documentlengte",     "60–120 pagina's",                  "30–70 pagina's (Part B) + bijlagen"),
    ("Rubricstructuur",    "Academische beoordelingsrubric",    "EC-evaluatierubric + call-specifieke expected outcomes"),
    ("Auteurs",            "Enkelvoudig",                       "1–40 auteurs, 3–20 instellingen"),
    ("Vertrouwelijkheid",  "Studentwerk, intern",               "Commercieel gevoelig, pre-competitief"),
    ("Iteratiesnelheid",   "Dagen tot weken",                   "Uren tot dagen voor deadline"),
    ("Inzet",             "Cijfer",                            "€500K–€10M financieringsbeslissing"),
    ("Feedbackontvangers", "Student + begeleider",              "PI + RSO + consortiumpartners"),
]
for r in diff_rows:
    row = t_diff.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

h2(doc, "Te bouwen rubricpakketten (MVP-prioriteit)")

rubric_packs = [
    ("🔴 Horizon Europe RIA/IA standaardrubric",      "Hoog volume — het werkpaard-instrument"),
    ("🔴 ERC Starting Grant rubric",                  "Hoogste prestige, hoog volume"),
    ("🟡 ERC Advanced Grant rubric",                  "Gevestigde onderzoekers"),
    ("🟡 MSCA Doctoral Network rubric",               "Consortia, ~15–20% slagingspercentage"),
    ("🟡 NWO Open Competition / VIDI / VICI rubric",  "Primaire nationale funder NL"),
    ("🟢 ZonMw Programmrubric (configureerbaar)",     "€300M+ jaarlijks, gestructureerde formats"),
]
for icon_title, desc in rubric_packs:
    bullet(doc, desc, icon_title + ": ")

h2(doc, "Minimaal Viable Product voor deze markt")

mvp_items = [
    ("2 werkende rubricpakketten", "Horizon Europe RIA en ERC StG — dekt het hoogste volume en de hoogste inzet"),
    ("Call topic upload", "Een tekstveld waar de gebruiker de expected outcomes uit het werkprogramma plakt. Zonder dit kan DocuCheck de meest voorkomende afwijzingsreden (mismatch met call) niet controleren"),
    ("Ernst-labeling", "Kritiek / Belangrijk / Minder op elke comment als post-processing stap op de output van Claude"),
    ("GDPR-conforme gegevensverwerking", "Niet een technische feature maar een juridische noodzaak — een DPA-template die universiteiten kunnen tekenen"),
    ("Landingspagina in de taal van RSO's", "Referentie aan EC-evaluatiecriteria bij naam, voorbeeldfeedbackcomments, pilot-resultaat van een Nederlandse universiteit"),
]
for title, text in mvp_items:
    bullet(doc, text, title + ": ")

h2(doc, "Partnerships die marktintrede versnellen")

partnerships = [
    ("NWO / VSNU (Universiteiten van Nederland)", "Pilotpartnerschap met NWO's grant support-netwerk geeft geloofwaardigheid en toegang tot de RSO-gemeenschap. NWO organiseert al workshops over voorstellopdrachten — DocuCheck als workshoptool."),
    ("EARMA (European Association of Research Managers)", "Sterk Nederlands chapter. Een presentatie bij een EARMA-evenement met een werkende demo bereikt honderden RSO's tegelijk."),
    ("Grantconsultancybedrijven", "Firms als Grant Solutions Europe of Triumf Grant Services kunnen resellers of doorverwijspartners worden — DocuCheck integreren in hun standaard reviewworkflow."),
    ("NWO Nationale Contactpunten (NCP's)", "NCP's adviseren aanvragers regelmatig en kunnen DocuCheck aanbevelen als zelfservice-reviewtool. Gehost bij NWO en RVO."),
    ("SURF Onderwijsdagen", "De belangrijkste Nederlandse hogeronderwijsconferentie voor IT en innovatie — ideaal podium voor een working demo."),
]
for title, text in partnerships:
    bullet(doc, text, title + ": ")

h2(doc, "Conclusie: De grant-writing markt in drie zinnen")
callout(doc,
    "Nederland is de ideale eerste markt: hoog aanvraagvolume, goed georganiseerde research support "
    "infrastructuur, tech-forward academische cultuur, en nabijheid tot EU-instellingen.\n\n"
    "Het pad naar markt is helder: twee rubricpakketten, GDPR-conforme dataverwerking, een pilot "
    "met één Nederlandse universiteits-RSO, en een conferentiepresentatie bij EARMA.\n\n"
    "De business case sluit snel: bij €25.000/jaar per instelling heeft DocuCheck slechts 12 "
    "Nederlandse universitaire klanten nodig om €300.000 ARR te genereren. Bij Europese schaal "
    "is de opportuniteit tientallen miljoenen.",
    "E8F5E9")

add_hr(doc)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run(
    "DocuCheck — EU-Subsidie Addendum  |  Hermes Strategic Advisory  |  Juni 2026\n"
    "Vertrouwelijk — intern gebruik"
).font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

out = r"C:\Users\Davidcassuto\ProjectFT\GRANTS_ADDENDUM.docx"
doc.save(out)
print(f"Saved: {out}")
