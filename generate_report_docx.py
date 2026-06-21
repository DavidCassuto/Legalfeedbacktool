from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# --- Helper: shade a paragraph ---
def shade_paragraph(para, hex_color="F2F2F2"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

# --- Helper: add a horizontal rule ---
def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p

# --- Cover page ---
doc.add_paragraph()
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("DocuCheck")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1a, 0x53, 0x7e)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle_para.add_run("Strategisch Rapport")
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Opgesteld: Juni 2026  |  Hermes Strategic Advisory\nVertrouwelijk — intern gebruik").font.size = Pt(11)

doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()

# --- TOC box ---
toc_para = doc.add_paragraph()
toc_run = toc_para.add_run("INHOUDSOPGAVE")
toc_run.bold = True
toc_run.font.size = Pt(13)
shade_paragraph(toc_para, "EBF4FF")

toc_items = [
    "1.  Waarom de tool goed genoeg is voor investering",
    "2.  De vijf belangrijkste faalredenen en mitigaties",
    "3.  Concurrentieanalyse",
    "4.  Waar moet de €300K–400K seed aan worden besteed?",
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ============================================================
# DEEL 1
# ============================================================
def h1(doc, text):
    p = doc.add_paragraph()
    shade_paragraph(p, "1A537E")
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1a, 0x53, 0x7e)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def callout(doc, text, color="FFF9C4"):
    p = doc.add_paragraph()
    shade_paragraph(p, color)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = True
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    return p

h1(doc, "DEEL 1 — Waarom de Tool Goed Genoeg Is voor Investering")

h2(doc, "Het kernprobleem dat niemand anders oplost")
body(doc,
    "Het probleem dat DocuCheck oplost is niet 'studenten willen betere feedback.' Het is: "
    "de scriptiedocent is al 40 jaar de bottleneck in scriptiekwaliteit, en dat is structureel nooit aangepakt.\n\n"
    "Een HBO-rechtenstudent krijgt gemiddeld 2–4 feedbackrondes over 6–12 maanden. Elke ronde kost een "
    "begeleider 3–6 uur. Een begeleider begeleid 10–20 studenten tegelijk. De math klopt niet — dus geven "
    "begeleiders oppervlakkige feedback. Dat was niet op te lossen vóór LLMs. Nu is het dat wel.")

h2(doc, "Waarom dit GEEN ChatGPT-wrapper is")

points = [
    ("Documentparsing met structuurbehoud",
     "DocuCheck parset .docx met python-docx, behoudt koppen, alineanummers en positiemetadata. "
     "Feedback is verankerd: 'alinea 3.2, paragraaf 4 adresseert het normatieve kader onvoldoende.' "
     "Dat is bruikbaar. 'Je methodologie is zwak' is dat niet."),
    ("Rubric-mapping als structurele constraint",
     "ChatGPT geeft feedback op basis van algemene schrijfkwaliteit. DocuCheck dwingt de AI om elk "
     "criterium uit de rubric — inclusief Nederlandse juridische begrippen als 'functionele deelvraag' "
     "en 'methodologische verantwoording' — te adresseren voor elk relevant stuk tekst."),
    ("Word-commentaren als workflow-integratie",
     "De output is een .docx met ingesloten <w:comment> XML-elementen, verankerd aan specifieke "
     "tekstfragmenten — precies zoals een begeleider het handmatig zou doen."),
    ("De holistische analyse-engine",
     "De huidige prototype-versie vereist géén voorgedefinieerde secties. De AI leest het hele document "
     "en bepaalt zelf welke passages relevant zijn voor welk criterium. Dat vereist weken aan prompt "
     "engineering die een student niet in een ChatGPT-sessie kan reproduceren."),
    ("Tijdwinst die telt",
     "DIY-aanpak: 45–90 minuten. DocuCheck: 5 minuten. Dat is geen marginale tijdwinst — dat is een ander product."),
]
for title, text in points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ": ").bold = True
    p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(5)

h2(doc, "Het platform dat hieronder zit")
body(doc,
    "De kern-engine — Rubric + Document → Gestructureerde feedback in vertrouwd formaat — werkt voor elk "
    "domein met beoordelingscriteria: HBO/WO scripties, EU-subsidieaanvragen (Horizon Europe), juridische "
    "memo's bij advocatenkantoren, corporate compliance documenten, medische casusverslagen.\n\n"
    "Juridisch Nederland is het juiste startpunt: gestandaardiseerde rubrics overdraagbaar tussen de "
    "13 Nederlandse rechtenfaculteiten, homogene documentformaten, acute begeleiderspijn, en slechts "
    "13 primaire beslissers.")

h2(doc, "Waarom nu, niet 3 jaar geleden")

# Simple table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Factor"
hdr[1].text = "Situatie 2022"
hdr[2].text = "Situatie 2025"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
rows_data = [
    ("LLM-kwaliteit",    "Generiek, begeleiders zouden het afwijzen",           "Claude 3/GPT-4 geeft professioneel geloofwaardige feedback"),
    ("API-kosten",        "$2–8 per documentanalyse",                           "€0,10–0,40 per analyse"),
    ("Word-generatie",    "Specialistisch, pijnlijk",                           "python-docx comment-XML is een opgelost probleem"),
    ("EU AI Act",         "Niet van kracht",                                    "Compliance-druk richting auditbare tools — tailwind voor institutionele verkoop"),
]
for r in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

h2(doc, "De verdedigbare positie die gebouwd KAN worden")
moat_items = [
    ("Rubric-bibliotheek",                "200–500 gevalideerde Nederlandse rubrics per opleiding — dataset die jaren duurt om na te maken"),
    ("Begeleider-workflow integratie",     "Begeleiders beheren rubrics in DocuCheck → tool wordt institutionele dependency"),
    ("Institutionele relaties",            "Getekende DPA + LMS-integratie creëren switching costs los van productkwaliteit"),
    ("Uitkomstdata",                       "'Studenten die DocuCheck gebruikten leverden 23% sneller definitieve versie in' — sluit deals"),
    ("Domein-specifieke prompt engineering","Kennis van NL juridische methodologie, HBO- vs. WO-conventies — niet triviaal na te bouwen"),
]
for title, text in moat_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ": ").bold = True
    p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

h2(doc, "Wat een investeerder verwacht na 18 maanden")
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
h2r = t2.rows[0].cells
for c, txt in zip(h2r, ["KPI", "Minimumdrempel", "Sterk geval"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
kpi_rows = [
    ("Betalende instellingen",        "2 getekende contracten",  "4–5"),
    ("ARR",                           "€60.000",                 "€150.000–200.000"),
    ("Studentgebruikers",             "1.500",                   "4.000+"),
    ("Pilot-naar-betaald conversie",  ">40%",                    ">60%"),
    ("Tweede discipline",             "Ja",                      "2 extra disciplines"),
    ("Gross Margin",                  ">65%",                    ">75%"),
]
for r in kpi_rows:
    row = t2.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()
callout(doc,
    "Eerlijk: €100K ARR in 18 maanden is de ondergrens van levensvatbaarheid, niet een doel om te vieren. "
    "Met een team van 2–3 moet je €200–250K targeten.", "FFF3CD")

doc.add_page_break()

# ============================================================
# DEEL 2
# ============================================================
h1(doc, "DEEL 2 — De Vijf Belangrijkste Faalredenen en Mitigaties")

failures = [
    {
        "title": "Faalreden 1: GDPR / Universitaire inkoopwal",
        "risk": (
            "Elke tool die studentgegevens verwerkt vereist een DPA én DPIA. Anthropic is Amerikaans — "
            "data naar hun API sturen is grensoverschrijdende doorgifte onder GDPR. Eén DPO die 'nee' "
            "zegt, en het gesprek is voorbij. Eén tweet van een privacybewuste professor kan een pilot stoppen."
        ),
        "mitigation": [
            "Week 1–2: Anthropic's commerciële DPA/SCC-pakket opvragen (gratis)",
            "Maand 1–2: DPIA-template laten opstellen door privacyconsultant (€2.000–4.000)",
            "Maand 2–3: Standaard verwerkersovereenkomst opstellen",
            "Maand 3–4: EU-data residency optie implementeren (Azure OpenAI EU-regio)",
            "Doorlopend: DPO bij doelinstelling identificeren VOOR procurement begint",
        ],
        "success": "Eerste DPA getekend. DPIA goedgekeurd door een universitaire DPO.",
        "residual": "20–30% kans dat 1–2 instellingen onbereikbaar zijn door AI-moratoriums. Spreid targets over meerdere instellingen.",
    },
    {
        "title": "Faalreden 2: Microsoft Copilot maakt DocuCheck overbodig",
        "risk": (
            "M365 Copilot zit in SURF-licenties. Copilot beweegt richting 'task-complete' en kan binnen "
            "12–18 maanden rubric-intake + comment-anchoring krijgen. Als dat gebeurt vervalt de core-differentiatie."
        ),
        "mitigation": [
            "Supervisordashboard bouwen: begeleider keurt AI-feedback goed vóór levering aan student",
            "Cohortanalytics: '78 van 200 studenten faalden op criterium 3.2' — dat is een faculteitsbeheertool",
            "Audit trail voor AI Act compliance: gedocumenteerde, tijdgestempelde feedback per student",
            "Institutional Rubric Locking: begeleiders bouwen hun officiële rubric in DocuCheck → switching cost",
            "LMS-integratie (Canvas/Brightspace) vóór Copilot dat doet",
        ],
        "success": "Minimaal 2 'begeleiderslaag'-features live en geadopteerd door institutionele klanten.",
        "residual": "Als Microsoft binnen 18 maanden een rubric-feature shipt met SURF-deployment, krimpt de differentiatie. Survival hangt dan af van institutionele relaties.",
    },
    {
        "title": "Faalreden 3: TAM te klein als je bij rechten blijft",
        "risk": (
            "~4.000–5.000 NL rechtenstudenten/jaar schrijven een scriptie. Zelfs 20% penetratie = €40K ARR max. "
            "Dat is geen bedrijf; dat is een side project met facturen."
        ),
        "mitigation": [
            "Jaar 1: NL rechten — product-market fit (5.000 studenten/jaar)",
            "Jaar 1–2: Alle HBO/WO disciplines NL (50.000 studenten/jaar)",
            "Jaar 2: Vlaanderen (vrijwel nul extra lokalisatie, +20% markt)",
            "Jaar 2: EU-subsidieaanvragen Horizon Europe (hoge betalingsbereidheid, nul concurrentie)",
            "Jaar 2–3: Duitstalig — Hausarbeit/Seminararbeit (300.000+ rechtenstudenten)",
            "Jaar 3+: Law firms, corporate compliance (hoge ACV, andere sales motion)",
        ],
        "success": "Bij maand 18: NL rechten + HBO-expansie in gang, Vlaanderen pipeline actief.",
        "residual": "HBO-expansie vereist andere sales motion. Risico: je zit langer in kleine TAM dan runway toestaat.",
    },
    {
        "title": "Faalreden 4: Solo-founder bandbreedte",
        "risk": (
            "DocuCheck vereist gelijktijdige voortgang op: product engineering, prompt engineering, "
            "institutionele sales, GDPR-compliance, en marketing. Een solo-founder die alles tegelijk "
            "doet, doet niets goed genoeg."
        ),
        "mitigation": [
            "Commercial Lead aannemen in Maand 1 (€55.000–75.000/jaar of equity)",
            "Fractioneel GDPR-counsel (€1.500–2.500/maand, maand 1–6)",
            "STOPPEN met: features bouwen die niemand vroeg, generieke startup-evenementen, eigen boekhouding",
            "Automatiseren: onboarding emails, basis support (FAQ + LLM-chatbot), rubric-parsing",
        ],
        "success": "Commercial Lead/co-founder aan boord voor maand 4, duidelijke taakverdeling.",
        "residual": "Kans op juiste co-founder binnen 6 maanden: 40–60%. Fallback: B2B sales hire voor €60K.",
    },
    {
        "title": "Faalreden 5: Pricing-val (te goedkoop voor B2B, te duur voor studenten)",
        "risk": (
            "Prijs je te laag (€5/analyse) en instellingen nemen je niet serieus. "
            "Prijs je te hoog (€50/maand) en studenten onboarden niet, waardoor institutionele bewustwording wegvalt. "
            "Prijzen in niemandsland = geen van beide markten."
        ),
        "mitigation": [
            "Gratis tier: 2 analyses lifetime, publieke rubrics, watermark → studentadoptie",
            "Student Pro: €12–18/maand of €8–12/analyse → studenten wiens instelling niet gecontracteerd is",
            "Institutioneel: €3.000–8.000/jaar per faculteit (onder NL aanbestedingsdrempel van ~€5.000–10.000)",
            "Pilotcontract bevat: DPA al getekend + heldere conversietermijn + commitment tot referentie-case",
        ],
        "success": "2 instellingen geconverteerd van gratis pilot naar betaald. B2C: €1.500–3.000/maand.",
        "residual": "Bij 13 NL rechtenfaculteiten verzadigt €3.000–5.000-pricing de markt op €40–65K ARR. TAM-expansie is geen optie maar vereiste.",
    },
]

for i, f in enumerate(failures):
    h2(doc, f["title"])
    p = doc.add_paragraph()
    p.add_run("Waarom het echt gevaarlijk is: ").bold = True
    p.add_run(f["risk"]).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.add_run("Concrete mitigatiestappen:").bold = True
    for step in f["mitigation"]:
        bp = doc.add_paragraph(step, style='List Bullet')
        bp.paragraph_format.left_indent = Cm(1)
        bp.paragraph_format.space_after = Pt(2)
        bp.runs[0].font.size = Pt(11)

    p3 = doc.add_paragraph()
    p3.add_run("Wat succes eruitziet: ").bold = True
    p3.add_run(f["success"]).font.size = Pt(11)
    p3.paragraph_format.space_after = Pt(2)

    callout(doc, "Residueel risico: " + f["residual"], "FFF3CD")

doc.add_page_break()

# ============================================================
# DEEL 3
# ============================================================
h1(doc, "DEEL 3 — Concurrentieanalyse")

h2(doc, "Positioneringsmatrix")
matrix_p = doc.add_paragraph()
matrix_p.add_run(
    "DocuCheck bezit het kwadrant 'rubric-specifiek + student-facing' vrijwel alleen.\n"
    "De dreiging komt van actoren die vanuit aangrenzende posities dit kwadrant binnenrijden.\n\n"
    "  RUBRIC-SPECIFIEKE FEEDBACK\n"
    "           ▲\n"
    "           │  DocuCheck ●  ← Jouw unieke terrein\n"
    "           │\n"
    "STUDENT ◄──┼──────────────────────────► BEGELEIDER\n"
    "           │\n"
    "   ChatGPT DIY ●        Turnitin Feedback Studio ●\n"
    "   Scribbr ● (taal)\n"
    "           ▼\n"
    "       GENERIEKE FEEDBACK\n\n"
    "Microsoft Copilot (nu): generiek/student-facing\n"
    "Microsoft Copilot (18 mnd): beweegt naar jouw kwadrant ← gevaar"
).font.name = "Courier New"
matrix_p.runs[0].font.size = Pt(9)

competitors = [
    {
        "name": "Concurrent 1: Microsoft Copilot in Word",
        "kill": "Medium",
        "timeline": "18–30 maanden",
        "cant_do": [
            "Rubric als beoordelingsinput accepteren",
            "Inline kantlijncomments anchored aan rubriccriteria plaatsen",
            "Downloadbaar Word-bestand genereren (output gaat naar sidebar)",
        ],
        "advantage": "Werkt vandaag, 5 minuten, verankerde kantlijncomments, law-specifieke rubrics kant-en-klaar.",
        "build": "Supervisordashboard, LMS-integratie, audit trail voor AI Act — dingen die Microsoft niet voor niches bouwt.",
    },
    {
        "name": "Concurrent 2: ChatGPT / Claude direct",
        "kill": "Laag",
        "timeline": "Nu en doorlopend",
        "cant_do": [
            "Word-output genereren met anchored kantlijncomments",
            "Rubric consistent per criterium toepassen",
            "45–90 minuten werk reduceren tot 5 minuten",
        ],
        "advantage": "Workflow-superioriteit, consistentie, output die er professioneel uitziet voor begeleiders.",
        "build": "Vergelijkingsmodus (draft vs. draft), feedbackexplainability, snelheid onder 3 minuten.",
    },
    {
        "name": "Concurrent 3: Scribbr",
        "kill": "Laag (direct) / Medium (M&A)",
        "timeline": "12–24 maanden",
        "cant_do": [
            "Rubric-gebaseerde inhoudelijke feedback (editors zijn taalredacteuren)",
            "Instant feedback — menselijke redactie duurt 24–72 uur",
            "Juridisch-inhoudelijke beoordeling (criterium per criterium)",
        ],
        "advantage": "Rubric-specificiteit, snelheid (minuten vs. dagen), prijs (€10–20 vs. €300–500/document).",
        "build": "White-label deal of partnerschap verkennen — Scribbr heeft de distributie, jij de rubric-engine. Scribbr is ook de meest waarschijnlijke acquirer.",
    },
    {
        "name": "Concurrent 4: Turnitin / Feedback Studio",
        "kill": "Medium-High",
        "timeline": "18–36 maanden",
        "cant_do": [
            "Student-facing AI feedback (Feedback Studio is instructor-only)",
            "Rubric-gebaseerde inline Word-comments",
            "Gebruik buiten LMS-grading-interface",
        ],
        "advantage": "Word-output (Turnitin leeft in browser-grading), student-facing gebruik, werkt vandaag zonder IT-procurement.",
        "build": "Institutionele contracten tekenen VÓÓRDAT Turnitin hun feature shipt. Canvas LTI-integratie overwegen.",
    },
    {
        "name": "Concurrent 5: De zelf-bouwende professor",
        "kill": "Laag",
        "timeline": "Nu",
        "cant_do": [
            "GDPR-compliance in 6–12 maanden institutionele juridische review",
            "Maintenance voor 200 studenten/semester (reëel FTE-equivalent)",
            "Contractuele aansprakelijkheid bieden aan de instelling",
        ],
        "advantage": "Dezelfde dag live, pre-gecertificeerd voor AVG, multi-rubric support, contractuele aansprakelijkheid, voortdurende verbetering.",
        "build": "Institutionele contracten vroeg sluiten vóórdat een faculteitsprototype informele standaard wordt.",
    },
]

for c in competitors:
    h2(doc, c["name"])
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Kill probability: {c['kill']}  |  Timeline: {c['timeline']}").font.size = Pt(10)
    meta_p.runs[0].italic = True
    meta_p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph().add_run("Kan vandaag niet:").bold = True
    for item in c["cant_do"]:
        bp = doc.add_paragraph(item, style='List Bullet')
        bp.runs[0].font.size = Pt(11)
        bp.paragraph_format.space_after = Pt(2)

    p_adv = doc.add_paragraph()
    p_adv.add_run("DocuCheck's voordeel nu: ").bold = True
    p_adv.add_run(c["advantage"]).font.size = Pt(11)
    p_adv.paragraph_format.space_after = Pt(2)

    p_build = doc.add_paragraph()
    p_build.add_run("Wat te bouwen in 12 maanden: ").bold = True
    p_build.add_run(c["build"]).font.size = Pt(11)
    p_build.paragraph_format.space_after = Pt(8)

# Summary table
h2(doc, "Dreigingstabel samenvatting")
t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
for c, txt in zip(t3.rows[0].cells, ["Concurrent", "Kill Probability", "Timeline", "Primaire dreiging"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
threat_rows = [
    ("Microsoft Copilot",    "Medium",      "18–30 maanden", "SURF-deployment + rubric feature"),
    ("ChatGPT/Claude DIY",   "Laag",        "Nu",            "UX-pariteit"),
    ("Scribbr",              "Laag / Medium","12–24 maanden", "Acquisitie of nieuw AI-product"),
    ("Turnitin",             "Medium-High", "18–36 maanden", "Institutionele LMS-feature"),
    ("Faculteit DIY",        "Laag",        "Nu",            "Prototype-adoptie vóór contract"),
]
for r in threat_rows:
    row = t3.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

doc.add_page_break()

# ============================================================
# DEEL 4
# ============================================================
h1(doc, "DEEL 4 — Waar moet de €300K–400K seed aan worden besteed?")

callout(doc,
    "In één zin: 48% op mensen, 12% op sales, 12% in reserve — en niets op features die niemand gevraagd heeft.",
    "E8F5E9")

h2(doc, "Budget breakdown (gebaseerd op €350K midpoint)")
t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
for c, txt in zip(t4.rows[0].cells, ["Categorie", "Bedrag", "%"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
budget_rows = [
    ("🧑‍💼 Team & Hiring",       "€168.000", "48%"),
    ("📣 Marketing & Sales",    "€42.000",  "12%"),
    ("🛡️ Runway Buffer",        "€42.000",  "12%"),
    ("🏛️ GDPR & Compliance",    "€28.000",  "8%"),
    ("🖥️ Infrastructuur & Tech", "€21.000",  "6%"),
    ("🔧 Operationeel",         "€14.000",  "4%"),
    ("TOTAAL",                  "€315.000", "100%"),
]
for r in budget_rows:
    row = t4.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
        if r[0] == "TOTAAL":
            row[i].paragraphs[0].runs[0].bold = True
doc.add_paragraph()

h2(doc, "Team breakdown (de grootste post)")
t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
for c, txt in zip(t5.rows[0].cells, ["Rol", "Kosten", "Timing"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
team_rows = [
    ("Commercial Lead (meest kritische hire)", "€112.500 (18 mnd)", "Maand 1"),
    ("Fractioneel GDPR-counsel",              "€18.000 (18 mnd)",  "Maand 1–2"),
    ("Part-time UX Designer",                 "€15.000 (6 mnd)",   "Maand 2–4"),
    ("Founder salaris (levensonderhoud)",      "€22.500 (€1.250/mnd)", "Doorlopend"),
]
for r in team_rows:
    row = t5.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

h2(doc, "Gefaseerde uitrol")
phases = [
    ("Maand 1–6: Fundament (€130.000)",
     ["Commercial Lead aannemen — elke week zonder is verloren outreach",
      "GDPR-counsel en DPO inschakelen voor maand 2",
      "SQLite → PostgreSQL migreren vóór maand 3 (hard deadline voor institutionele demo's)",
      "3–5 institutionele pilots tekenen (gratis of gereduceerd)",
      "Doel: GDPR-documentatie klaar voor procurement review"]),
    ("Maand 7–12: Tractie (€110.000)",
     ["Pilots omzetten naar betalende contracten",
      "TAM-expansie: Belgische rechtenfaculteiten, MBA-programma's verkennen",
      "Content/SEO-programma starten",
      "Security audit commissionen (maand 8–9)",
      "Doel: €15.000–40.000 ARR, 2–3 betalende instellingen"]),
    ("Maand 13–18: Schalen & Series A prep (€75.000)",
     ["Verdubbelen op wat werkt",
      "1–2 aangrenzende verticalen (medische school, engineering theses)",
      "Series A investor outreach starten bij maand 14–15",
      "Doel: €80.000–150.000 ARR, 12+ institutionele klanten"]),
]
for phase_title, phase_items in phases:
    p = doc.add_paragraph()
    p.add_run(phase_title).bold = True
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    for item in phase_items:
        bp = doc.add_paragraph(item, style='List Bullet')
        bp.runs[0].font.size = Pt(11)
        bp.paragraph_format.left_indent = Cm(1)
        bp.paragraph_format.space_after = Pt(2)

h2(doc, "Wat je NIET moet uitgeven")
dont_items = [
    "Duur kantoor — remote-first is normaal en geaccepteerd in NL",
    "Full agency rebrand — één goede UX-contractor > tien branding agencies",
    "Enterprise sales tools voor je enterprise sales hebt (Salesforce etc.) — HubSpot gratis is genoeg tot €200K ARR",
    "Extra developers aannemen — jij bent de developer; leverage zit commercieel",
    "Conferentie-sponsorships — bijwonen: €500; sponsoren: €10.000+",
    "Outsourced cold outreach agencies — spam naar NL onderwijs = permanente blacklisting",
]
for item in dont_items:
    bp = doc.add_paragraph("❌  " + item, style='List Bullet')
    bp.runs[0].font.size = Pt(11)
    bp.paragraph_format.space_after = Pt(3)

h2(doc, "De meest kritische investering")
callout(doc,
    "Commercial Lead aannemen in Maand 1.\n\n"
    "De bottleneck is niet het product — het is distributie in een complex, langzaam institutioneel markt "
    "waar relaties en procurement-navigatie alles bepalen. Als deze hire fout gaat (te junior, verkeerd netwerk), "
    "valt het hele 18-maanden plan in duigen. Bied equity boven cash aan om iemand senior aan te trekken.",
    "E8F4FF")

h2(doc, "Runway en Series A")
t6 = doc.add_table(rows=1, cols=2)
t6.style = 'Table Grid'
for c, txt in zip(t6.rows[0].cells, ["Metric", "Getal"]):
    c.text = txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
runway_rows = [
    ("Maandelijkse burn",           "~€18.000/maand"),
    ("Runway zonder omzet",         "~18–19 maanden"),
    ("Bij €100K ARR: netto-burn",   "~€10.000/mnd → +3–4 mnd extra"),
    ("Start Series A gesprekken",   "Bij €50.000–70.000 ARR"),
    ("Sluit Series A",              "Bij €100.000–120.000 ARR, ≥6 mnd runway over"),
    ("Target raise Series A",       "€1,5M–2,5M"),
]
for r in runway_rows:
    row = t6.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

callout(doc,
    "Het Series A-verhaal:\n\n"
    "\"DocuCheck startte als AI-feedbacktool voor Nederlandse rechtenstudenten. In 18 maanden tekenden we "
    "contracten met [X] universiteiten, bereikten [X]K ARR, en ontdekten dat ons product even goed werkt "
    "voor medische, engineering en business scripties — elk programma met gestructureerde beoordelingscriteria. "
    "Europa heeft 4.000+ HEI's. We hebben bewezen dat institutionele procurement te kraken is in Nederland. "
    "Nu nemen we het playbook naar Duitsland, België en het VK.\"",
    "E8F5E9")

add_hr(doc)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run(
    "DocuCheck Strategic Report  |  Hermes Strategic Advisory  |  Juni 2026\n"
    "Vertrouwelijk — intern gebruik"
).font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

out_path = r"C:\Users\Davidcassuto\ProjectFT\STRATEGIC_REPORT.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
