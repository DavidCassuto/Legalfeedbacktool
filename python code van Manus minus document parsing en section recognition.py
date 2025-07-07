# 1. CHECK_SMART_FORMULATION FUNCTIE

def check_smart_formulation(criterion, section):
    """
    Controleert of tekst SMART geformuleerd is
    SMART = Specifiek, Meetbaar, Acceptabel, Realistisch, Tijdgebonden
    
    Args:
        criterion: Criterium dictionary
        section: Sectie dictionary met content
    
    Returns:
        Feedback dictionary of None als criterium wordt gehaald
    """
    content = section.get('content', '').lower()
    
    if not content.strip():
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'violation',
            'message': f"De {section['name'].lower()} is leeg",
            'suggestion': "Voeg inhoud toe aan deze sectie",
            'confidence': 1.0,
            'color': criterion['color']
        }
    
    # SMART aspecten met indicatoren
    smart_indicators = {
        'specifiek': {
            'keywords': [
                'specifiek', 'concreet', 'duidelijk', 'precies', 'exact',
                'bepaald', 'vastgesteld', 'gedefinieerd', 'omschreven',
                'wie', 'wat', 'waar', 'wanneer', 'waarom', 'hoe'
            ],
            'patterns': [
                r'\b(wie|wat|waar|wanneer|waarom|hoe)\b',
                r'\b(specifiek|concreet|duidelijk)\b',
                r'\b(precies|exact|bepaald)\b'
            ],
            'weight': 1.0
        },
        'meetbaar': {
            'keywords': [
                'meetbaar', 'kwantificeerbaar', 'cijfers', 'percentage', 'aantal',
                'procent', 'euro', 'meter', 'kilogram', 'uur', 'dag', 'week',
                'maand', 'jaar', 'stuks', 'keer', 'maal', 'x', '%'
            ],
            'patterns': [
                r'\d+\s*(procent|%|euro|€|meter|m|kilogram|kg|uur|dag|dagen|week|weken|maand|maanden|jaar|jaren)',
                r'\d+\s*(stuks|keer|maal|x)',
                r'\b\d+[.,]?\d*\b',  # Algemene getallen
                r'\b(hoeveel|aantal|cijfer|getal|waarde|score)\b'
            ],
            'weight': 1.2  # Meetbaarheid is extra belangrijk
        },
        'acceptabel': {
            'keywords': [
                'acceptabel', 'haalbaar', 'realistisch', 'mogelijk', 'bereikbaar',
                'uitvoerbaar', 'redelijk', 'passend', 'geschikt', 'relevant',
                'belangrijk', 'waardevol', 'zinvol', 'nuttig'
            ],
            'patterns': [
                r'\b(acceptabel|haalbaar|realistisch)\b',
                r'\b(mogelijk|bereikbaar|uitvoerbaar)\b',
                r'\b(relevant|belangrijk|zinvol)\b'
            ],
            'weight': 0.8
        },
        'realistisch': {
            'keywords': [
                'realistisch', 'haalbaar', 'mogelijk', 'bereikbaar', 'uitvoerbaar',
                'redelijk', 'praktisch', 'werkbaar', 'doenlijk', 'feasible',
                'capaciteit', 'middelen', 'resources', 'budget', 'tijd'
            ],
            'patterns': [
                r'\b(realistisch|haalbaar|mogelijk)\b',
                r'\b(praktisch|werkbaar|doenlijk)\b',
                r'\b(capaciteit|middelen|resources|budget)\b'
            ],
            'weight': 1.0
        },
        'tijdgebonden': {
            'keywords': [
                'deadline', 'datum', 'week', 'maand', 'jaar', 'tijd', 'periode',
                'termijn', 'planning', 'schema', 'tijdschema', 'tijdslijn',
                'januari', 'februari', 'maart', 'april', 'mei', 'juni',
                'juli', 'augustus', 'september', 'oktober', 'november', 'december',
                'maandag', 'dinsdag', 'woensdag', 'donderdag', 'vrijdag', 'zaterdag', 'zondag'
            ],
            'patterns': [
                r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b',  # Datums: 01-01-2024
                r'\b(januari|februari|maart|april|mei|juni|juli|augustus|september|oktober|november|december)\b',
                r'\b(week|weken|maand|maanden|jaar|jaren)\s*\d+\b',
                r'\b\d+\s*(week|weken|maand|maanden|jaar|jaren)\b',
                r'\b(deadline|datum|termijn|planning)\b'
            ],
            'weight': 1.1
        }
    }
    
    # Analyseer elk SMART aspect
    found_aspects = {}
    missing_aspects = []
    weak_aspects = []
    
    for aspect, indicators in smart_indicators.items():
        score = calculate_smart_aspect_score(content, indicators)
        found_aspects[aspect] = score
        
        if score['total_score'] < 0.3:  # Zeer zwak
            missing_aspects.append(aspect)
        elif score['total_score'] < 0.6:  # Zwak aanwezig
            weak_aspects.append(aspect)
    
    # Bepaal overall SMART score
    total_score = sum(aspect['total_score'] * smart_indicators[name]['weight'] 
                     for name, aspect in found_aspects.items())
    max_possible_score = sum(smart_indicators[name]['weight'] for name in smart_indicators.keys())
    smart_percentage = (total_score / max_possible_score) * 100
    
    # Genereer feedback op basis van score
    if smart_percentage < 40:  # Zeer slecht
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'violation',
            'message': f"De {section['name'].lower()} is niet SMART geformuleerd (score: {smart_percentage:.0f}%)",
            'suggestion': generate_smart_suggestions(missing_aspects, weak_aspects, found_aspects),
            'confidence': 0.9,
            'color': criterion['color'],
            'details': {
                'smart_score': smart_percentage,
                'missing_aspects': missing_aspects,
                'weak_aspects': weak_aspects,
                'aspect_scores': found_aspects
            }
        }
    
    elif smart_percentage < 70:  # Matig
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'warning',
            'message': f"De {section['name'].lower()} kan SMART-er geformuleerd worden (score: {smart_percentage:.0f}%)",
            'suggestion': generate_smart_suggestions(missing_aspects, weak_aspects, found_aspects),
            'confidence': 0.7,
            'color': '#F9C74F',
            'details': {
                'smart_score': smart_percentage,
                'missing_aspects': missing_aspects,
                'weak_aspects': weak_aspects,
                'aspect_scores': found_aspects
            }
        }
    
    # Score >= 70%: Goed geformuleerd, geen feedback nodig
    return None

def calculate_smart_aspect_score(content, indicators):
    """
    Berekent score voor een specifiek SMART aspect
    
    Returns:
        {
            'keyword_score': float,
            'pattern_score': float,
            'total_score': float,
            'found_keywords': list,
            'found_patterns': list
        }
    """
    import re
    
    # Keyword matching
    found_keywords = []
    for keyword in indicators['keywords']:
        if keyword in content:
            found_keywords.append(keyword)
    
    keyword_score = min(len(found_keywords) / 3, 1.0)  # Max score bij 3+ keywords
    
    # Pattern matching
    found_patterns = []
    for pattern in indicators['patterns']:
        matches = re.findall(pattern, content, re.IGNORECASE)
        if matches:
            found_patterns.extend(matches)
    
    pattern_score = min(len(found_patterns) / 2, 1.0)  # Max score bij 2+ patterns
    
    # Combineer scores (keywords 60%, patterns 40%)
    total_score = (keyword_score * 0.6) + (pattern_score * 0.4)
    
    return {
        'keyword_score': keyword_score,
        'pattern_score': pattern_score,
        'total_score': total_score,
        'found_keywords': found_keywords,
        'found_patterns': found_patterns
    }

def generate_smart_suggestions(missing_aspects, weak_aspects, found_aspects):
    """
    Genereert concrete suggesties voor SMART verbetering
    """
    suggestions = []
    
    # Suggesties voor ontbrekende aspecten
    smart_suggestions = {
        'specifiek': [
            "Maak duidelijk WIE, WAT, WAAR en WANNEER",
            "Gebruik concrete en precieze bewoordingen",
            "Vermijd vage termen zoals 'beter', 'meer', 'goed'"
        ],
        'meetbaar': [
            "Voeg cijfers, percentages of aantallen toe",
            "Definieer hoe succes gemeten wordt",
            "Gebruik kwantificeerbare indicatoren"
        ],
        'acceptabel': [
            "Leg uit waarom dit doel relevant en belangrijk is",
            "Toon aan dat het doel acceptabel is voor betrokkenen",
            "Motiveer de waarde van het doel"
        ],
        'realistisch': [
            "Beschrijf beschikbare middelen en capaciteit",
            "Toon aan dat het doel haalbaar is",
            "Noem eventuele beperkingen of uitdagingen"
        ],
        'tijdgebonden': [
            "Voeg een concrete deadline toe",
            "Specificeer de tijdsperiode",
            "Gebruik datums of tijdsindicaties"
        ]
    }
    
    # Prioriteer ontbrekende aspecten
    for aspect in missing_aspects:
        suggestions.append(f"**{aspect.title()}**: {smart_suggestions[aspect][0]}")
    
    # Voeg suggesties toe voor zwakke aspecten
    for aspect in weak_aspects:
        if len(suggestions) < 3:  # Beperk aantal suggesties
            suggestions.append(f"**{aspect.title()}** (verbeteren): {smart_suggestions[aspect][1]}")
    
    # Als geen specifieke suggesties, geef algemene tip
    if not suggestions:
        suggestions.append("Controleer of alle SMART aspecten duidelijk aanwezig zijn")
    
    return " | ".join(suggestions[:3])  # Max 3 suggesties

def analyze_smart_examples(content):
    """
    Analyseert voorbeelden van goede/slechte SMART formuleringen in de tekst
    """
    examples = {
        'good_examples': [],
        'bad_examples': []
    }
    
    # Split content in zinnen
    sentences = content.split('.')
    
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 20:  # Te kort voor analyse
            continue
            
        # Eenvoudige heuristiek voor goede/slechte voorbeelden
        has_numbers = bool(re.search(r'\d+', sentence))
        has_time = bool(re.search(r'(week|maand|jaar|datum|deadline)', sentence.lower()))
        has_specific = bool(re.search(r'(wie|wat|waar|specifiek|concreet)', sentence.lower()))
        
        score = sum([has_numbers, has_time, has_specific])
        
        if score >= 2:
            examples['good_examples'].append(sentence)
        elif score == 0:
            examples['bad_examples'].append(sentence)
    
    return examples

def get_smart_improvement_priority(found_aspects):
    """
    Bepaalt welke SMART aspecten prioriteit hebben voor verbetering
    """
    # Sorteer aspecten op score (laagste eerst)
    sorted_aspects = sorted(found_aspects.items(), key=lambda x: x[1]['total_score'])
    
    priority_list = []
    for aspect, score_data in sorted_aspects:
        if score_data['total_score'] < 0.6:
            priority_list.append({
                'aspect': aspect,
                'score': score_data['total_score'],
                'priority': 'high' if score_data['total_score'] < 0.3 else 'medium'
            })
    
    return priority_list

# Voorbeeld van gebruik:
def test_smart_formulation():
    """Test functie om SMART formulation te demonstreren"""
    
    # Test criterium
    test_criterion = {
        'id': 1,
        'name': 'SMART doelstelling',
        'color': '#F94144'
    }
    
    # Test secties
    test_sections = [
        {
            'id': 'good_example',
            'name': 'Goede Doelstelling',
            'content': '''Het doel is om binnen 6 maanden de klanttevredenheid te verhogen van 75% naar 85% door het implementeren van een nieuw klantenservice systeem. Dit wordt gemeten via maandelijkse klanttevredenheidsonderzoeken met minimaal 100 respondenten per maand. Het project heeft een budget van €50.000 en wordt uitgevoerd door het customer service team van 5 personen. De deadline is 31 december 2024.'''
        },
        {
            'id': 'bad_example',
            'name': 'Slechte Doelstelling',
            'content': '''We willen de klantenservice verbeteren. Dit is belangrijk voor het bedrijf en zal helpen om betere resultaten te behalen. We gaan ons best doen om dit te realiseren.'''
        }
    ]
    
    # Test beide voorbeelden
    for section in test_sections:
        result = check_smart_formulation(test_criterion, section)
        print(f"\n=== {section['name']} ===")
        if result:
            print(f"Status: {result['status']}")
            print(f"Message: {result['message']}")
            print(f"Suggestion: {result['suggestion']}")
            if 'details' in result:
                print(f"SMART Score: {result['details']['smart_score']:.1f}%")
        else:
            print("Status: PASSED - Goed SMART geformuleerd!")

# Uncomment om te testen:
# test_smart_formulation()

# 2. AANVULLENDE SMART UTILITIES

def create_smart_checklist():
    """Creëert een checklist voor SMART doelstellingen"""
    return {
        'specifiek': [
            "Is duidelijk WIE betrokken is?",
            "Is duidelijk WAT er bereikt moet worden?",
            "Is duidelijk WAAR dit plaatsvindt?",
            "Zijn vage termen vermeden?"
        ],
        'meetbaar': [
            "Zijn er concrete cijfers genoemd?",
            "Is duidelijk HOE succes gemeten wordt?",
            "Zijn er kwantificeerbare indicatoren?",
            "Kunnen voortgang en resultaat objectief beoordeeld worden?"
        ],
        'acceptabel': [
            "Is het doel relevant voor de organisatie?",
            "Zijn stakeholders het eens met het doel?",
            "Is de waarde van het doel duidelijk?",
            "Sluit het doel aan bij bredere doelstellingen?"
        ],
        'realistisch': [
            "Zijn de benodigde middelen beschikbaar?",
            "Is er voldoende tijd en capaciteit?",
            "Zijn mogelijke obstakels in kaart gebracht?",
            "Is het doel uitdagend maar haalbaar?"
        ],
        'tijdgebonden': [
            "Is er een concrete deadline?",
            "Zijn er tussentijdse mijlpalen?",
            "Is de tijdslijn realistisch?",
            "Is duidelijk wanneer het doel behaald moet zijn?"
        ]
    }

def generate_smart_template():
    """Genereert een template voor SMART doelstellingen"""
    return """
SMART Doelstelling Template:

**Specifiek**: [WIE] wil [WAT] bereiken door [HOE]
**Meetbaar**: Dit wordt gemeten aan de hand van [INDICATOR] met als streefwaarde [CIJFER/PERCENTAGE]
**Acceptabel**: Dit doel is relevant omdat [REDEN] en draagt bij aan [BREDERE DOELSTELLING]
**Realistisch**: We hebben [MIDDELEN/CAPACITEIT] beschikbaar en rekening gehouden met [BEPERKINGEN]
**Tijdgebonden**: Het doel moet behaald zijn op [DATUM] met tussentijdse evaluatie op [MIJLPALEN]

Voorbeeld:
Het marketing team (WIE) wil de online verkoop verhogen met 25% (WAT + MEETBAAR) door implementatie van een nieuwe e-commerce strategie (HOE). Dit wordt gemeten via maandelijkse verkooprapportages (INDICATOR). Het doel is relevant omdat online verkoop cruciaal is voor bedrijfsgroei (ACCEPTABEL). We hebben een budget van €75.000 en 3 FTE beschikbaar (REALISTISCH). Deadline is 30 juni 2024 met evaluatie elke maand (TIJDGEBONDEN).
    """
#. 3  SIMILARITY BEREKENING

def calculate_similarity(text1, text2):
    """
    Berekent similarity tussen twee teksten (0.0 - 1.0)
    Gebruikt Levenshtein distance en gemeenschappelijke woorden
    """
    import difflib
    
    # Normaliseer teksten
    text1 = text1.lower().strip()
    text2 = text2.lower().strip()
    
    if text1 == text2:
        return 1.0
    
    # Sequence matcher voor basis similarity
    seq_similarity = difflib.SequenceMatcher(None, text1, text2).ratio()
    
    # Woord-gebaseerde similarity
    words1 = set(text1.split())
    words2 = set(text2.split())
    
    if not words1 and not words2:
        return 1.0
    if not words1 or not words2:
        return 0.0
    
    # Jaccard similarity (gemeenschappelijke woorden / alle woorden)
    intersection = words1.intersection(words2)
    union = words1.union(words2)
    word_similarity = len(intersection) / len(union) if union else 0.0
    
    # Combineer beide methoden (gewogen gemiddelde)
    combined_similarity = (seq_similarity * 0.4) + (word_similarity * 0.6)
    
    return round(combined_similarity, 3)

def levenshtein_distance(s1, s2):
    """Berekent Levenshtein distance tussen twee strings"""
    if len(s1) < len(s2):
        return levenshtein_distance(s2, s1)

    if len(s2) == 0:
        return len(s1)

    previous_row = list(range(len(s2) + 1))
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row
    
    return previous_row[-1]

def fuzzy_match_score(text1, text2):
    """Alternatieve fuzzy matching met meer geavanceerde algoritmes"""
    # Normaliseer
    text1 = text1.lower().strip()
    text2 = text2.lower().strip()
    
    # Exact match
    if text1 == text2:
        return 1.0
    
    # Substring match
    if text1 in text2 or text2 in text1:
        shorter = min(len(text1), len(text2))
        longer = max(len(text1), len(text2))
        return shorter / longer * 0.9  # Penalty voor niet-exacte match
    
    # Levenshtein-gebaseerde similarity
    max_len = max(len(text1), len(text2))
    if max_len == 0:
        return 1.0
    
    distance = levenshtein_distance(text1, text2)
    similarity = 1 - (distance / max_len)
    
    return max(0.0, similarity)

# 4. SECTIE CONTENT EXTRACTIE

def extract_section_content(paragraphs, start_index, end_index):
    """
    Extraheert de volledige content van een sectie
    
    Args:
        paragraphs: List van paragraph dictionaries
        start_index: Start paragraph index (inclusief)
        end_index: End paragraph index (inclusief)
    
    Returns:
        {
            'full_text': str,
            'word_count': int,
            'paragraph_count': int,
            'sentences': list,
            'headings': list,
            'content_paragraphs': list  # Alleen content, geen headings
        }
    """
    if start_index < 0 or end_index >= len(paragraphs) or start_index > end_index:
        return {
            'full_text': '',
            'word_count': 0,
            'paragraph_count': 0,
            'sentences': [],
            'headings': [],
            'content_paragraphs': []
        }
    
    section_paragraphs = paragraphs[start_index:end_index + 1]
    
    full_text = ""
    content_paragraphs = []
    headings = []
    
    for para in section_paragraphs:
        text = para['text'].strip()
        if text:
            full_text += text + "\n"
            
            if para['is_heading']:
                headings.append(text)
            else:
                content_paragraphs.append(text)
    
    # Splits in zinnen (eenvoudige methode)
    sentences = split_into_sentences(full_text)
    
    # Tel woorden (exclusief headings voor nauwkeurigere telling)
    content_text = "\n".join(content_paragraphs)
    word_count = len(content_text.split()) if content_text.strip() else 0
    
    return {
        'full_text': full_text.strip(),
        'word_count': word_count,
        'paragraph_count': len(content_paragraphs),
        'sentences': sentences,
        'headings': headings,
        'content_paragraphs': content_paragraphs
    }

def split_into_sentences(text):
    """Split tekst in zinnen"""
    import re
    
    # Eenvoudige sentence splitting
    # Splits op . ! ? gevolgd door spatie en hoofdletter
    sentence_pattern = r'[.!?]+\s+(?=[A-Z])'
    sentences = re.split(sentence_pattern, text)
    
    # Clean up en filter lege zinnen
    sentences = [s.strip() for s in sentences if s.strip()]
    
    return sentences

def extract_keywords(text, min_length=3):
    """Extraheert keywords uit tekst"""
    import re
    
    # Verwijder leestekens en splits in woorden
    words = re.findall(r'\b[a-zA-Z]{' + str(min_length) + ',}\b', text.lower())
    
    # Nederlandse stopwoorden
    stopwords = {
        'het', 'de', 'een', 'van', 'en', 'in', 'op', 'met', 'voor', 'door',
        'aan', 'bij', 'uit', 'over', 'naar', 'als', 'dat', 'dit', 'zijn',
        'hebben', 'worden', 'kan', 'moet', 'zal', 'zou', 'ook', 'nog',
        'maar', 'omdat', 'dus', 'echter', 'daarom', 'hierdoor', 'waarbij'
    }
    
    # Filter stopwoorden
    keywords = [word for word in words if word not in stopwords]
    
    # Tel frequentie
    from collections import Counter
    keyword_freq = Counter(keywords)
    
    return keyword_freq.most_common(10)  # Top 10 keywords

# 5. ONTBREKENDE SECTIES VINDEN

def find_missing_sections(found_sections, section_templates):
    """
    Vindt ontbrekende verplichte secties
    
    Args:
        found_sections: List van gevonden secties
        section_templates: List van sectie templates
    
    Returns:
        List van ontbrekende sectie objecten
    """
    found_ids = {section['id'] for section in found_sections}
    missing_sections = []
    
    for template in section_templates:
        if template.get('required', False) and template['id'] not in found_ids:
            missing_section = {
                'id': template['id'],
                'name': template['name'],
                'template_name': template['name'],
                'content': '',
                'found': False,
                'confidence': 0.0,
                'level': 1,
                'start_paragraph': None,
                'end_paragraph': None,
                'missing_reason': determine_missing_reason(template, found_sections),
                'suggestions': generate_missing_section_suggestions(template)
            }
            missing_sections.append(missing_section)
    
    return missing_sections

def determine_missing_reason(template, found_sections):
    """Bepaalt waarom een sectie ontbreekt"""
    # Kijk of er vergelijkbare secties zijn gevonden
    similar_sections = []
    
    for section in found_sections:
        similarity = calculate_similarity(section['name'], template['name'])
        if similarity > 0.3:  # Enige gelijkenis
            similar_sections.append((section, similarity))
    
    if similar_sections:
        best_match = max(similar_sections, key=lambda x: x[1])
        return f"Mogelijk vergelijkbaar met '{best_match[0]['name']}' (similarity: {best_match[1]:.2f})"
    
    return "Sectie niet gevonden in document"

def generate_missing_section_suggestions(template):
    """Genereert suggesties voor ontbrekende secties"""
    suggestions = []
    
    # Basis suggestie
    suggestions.append(f"Voeg een sectie toe met de titel '{template['name']}'")
    
    # Alternatieve namen
    if template.get('alternative_names'):
        alt_names = template['alternative_names']
        if isinstance(alt_names, str):
            import json
            try:
                alt_names = json.loads(alt_names)
            except:
                alt_names = [alt_names]
        
        if alt_names:
            suggestions.append(f"Alternatieve titels: {', '.join(alt_names)}")
    
    # Positie suggestie
    if template.get('order'):
        if template['order'] == 1:
            suggestions.append("Deze sectie hoort meestal aan het begin van het document")
        elif template['order'] <= 3:
            suggestions.append("Deze sectie hoort meestal in het eerste deel van het document")
        else:
            suggestions.append("Deze sectie hoort meestal later in het document")
    
    return suggestions

# 6.  TOEPASSELIJKE SECTIES BEPALEN

def get_applicable_sections(criterion, sections, document_type):
    """
    Bepaalt op welke secties een criterium van toepassing is
    
    Args:
        criterion: Criterium dictionary
        sections: List van gevonden secties
        document_type: Type document
    
    Returns:
        List van secties waar criterium op toegepast moet worden
    """
    applicable_sections = []
    
    # Haal criteria-sectie mappings op uit database
    mappings = get_criterion_section_mappings(criterion['id'], document_type)
    
    if criterion['application_scope'] == 'all':
        # Criterium geldt voor alle secties
        applicable_sections = [s for s in sections if s['found']]
        
    elif criterion['application_scope'] == 'specific_sections':
        # Criterium geldt alleen voor specifieke secties
        if mappings:
            mapped_section_ids = {m['section_id'] for m in mappings if not m['is_excluded']}
            applicable_sections = [s for s in sections if s['id'] in mapped_section_ids and s['found']]
        else:
            # Fallback: gebruik standaard secties voor criterium type
            default_sections = get_default_sections_for_criterion(criterion)
            applicable_sections = [s for s in sections if s['id'] in default_sections and s['found']]
            
    elif criterion['application_scope'] == 'exclude_sections':
        # Criterium geldt voor alle secties behalve uitgesloten
        excluded_section_ids = {m['section_id'] for m in mappings if m['is_excluded']}
        applicable_sections = [s for s in sections if s['id'] not in excluded_section_ids and s['found']]
    
    return applicable_sections

def get_criterion_section_mappings(criterion_id, document_type):
    """Haalt mappings op uit database (mock implementatie)"""
    # In echte implementatie: database query
    # Hier: mock data voor demonstratie
    
    mock_mappings = {
        1: [  # SMART doelstelling
            {'section_id': 'doelstelling', 'is_excluded': False},
            {'section_id': 'onderzoeksvragen', 'is_excluded': False}
        ],
        2: [  # Hoofdvraag aansluiting
            {'section_id': 'onderzoeksvragen', 'is_excluded': False},
            {'section_id': 'probleemstelling', 'is_excluded': False}
        ],
        3: [  # APA bronvermelding
            {'section_id': 'literatuur', 'is_excluded': False}
        ],
        4: [  # Minimaal 500 woorden
            {'section_id': 'inleiding', 'is_excluded': False},
            {'section_id': 'methode', 'is_excluded': False}
        ]
    }
    
    return mock_mappings.get(criterion_id, [])

def get_default_sections_for_criterion(criterion):
    """Bepaalt standaard secties voor een criterium type"""
    defaults = {
        'tekstueel': ['inleiding', 'probleemstelling', 'doelstelling', 'methode'],
        'structureel': ['inleiding', 'probleemstelling', 'doelstelling', 'onderzoeksvragen', 'methode', 'planning'],
        'inhoudelijk': ['probleemstelling', 'doelstelling', 'onderzoeksvragen', 'methode'],
        'referenties': ['literatuur']
    }
    
    return defaults.get(criterion.get('rule_type', 'tekstueel'), [])

# 7. STRUCTURELE CRITERIA CONTROLE

def check_structural_criterion(criterion, section):
    """
    Controleert structurele criteria (lengte, opbouw, volgorde)
    """
    criterion_name = criterion['name'].lower()
    
    # Woordtelling criteria
    if 'woorden' in criterion_name or 'word' in criterion_name:
        return check_word_count(criterion, section)
    
    # Paragraaf criteria
    if 'paragraaf' in criterion_name or 'alinea' in criterion_name:
        return check_paragraph_structure(criterion, section)
    
    # Kopjes/structuur criteria
    if 'kopje' in criterion_name or 'structuur' in criterion_name:
        return check_heading_structure(criterion, section)
    
    # Volgorde criteria
    if 'volgorde' in criterion_name or 'order' in criterion_name:
        return check_section_order(criterion, section)
    
    return None

def check_paragraph_structure(criterion, section):
    """Controleert paragraaf structuur"""
    content = section.get('content', '')
    paragraphs = content.split('\n\n')  # Eenvoudige paragraaf scheiding
    paragraph_count = len([p for p in paragraphs if p.strip()])
    
    # Extract minimum uit criterium naam
    import re
    min_match = re.search(r'(\d+)', criterion['name'])
    min_paragraphs = int(min_match.group(1)) if min_match else 3
    
    if paragraph_count < min_paragraphs:
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'violation',
            'message': f"Sectie heeft {paragraph_count} paragrafen, minimaal {min_paragraphs} vereist",
            'suggestion': f"Voeg {min_paragraphs - paragraph_count} paragrafen toe voor betere structuur",
            'confidence': 0.9,
            'color': criterion['color']
        }
    
    return None

def check_heading_structure(criterion, section):
    """Controleert kopjes structuur"""
    headings = section.get('headings', [])
    
    if not headings:
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'warning',
            'message': "Sectie heeft geen subkopjes",
            'suggestion': "Overweeg subkopjes toe te voegen voor betere structuur",
            'confidence': 0.7,
            'color': '#F9C74F'
        }
    
    return None

def check_section_order(criterion, section):
    """Controleert sectie volgorde (vereist context van hele document)"""
    # Deze functie zou de hele document context nodig hebben
    # Voor nu: placeholder implementatie
    return None

# 8. INHOUDELIJKE CRITERIA CONTROLE

def check_content_criterion(criterion, section):
    """
    Controleert inhoudelijke criteria (concepten, argumentatie, logica)
    """
    criterion_name = criterion['name'].lower()
    content = section.get('content', '').lower()
    
    # SMART criteria (al geïmplementeerd in eerdere code)
    if 'smart' in criterion_name:
        return check_smart_formulation(criterion, section)
    
    # Aansluiting/coherentie criteria
    if 'aansluiting' in criterion_name or 'coherent' in criterion_name:
        return check_coherence(criterion, section)
    
    # Argumentatie criteria
    if 'argument' in criterion_name or 'onderbouw' in criterion_name:
        return check_argumentation(criterion, section)
    
    # Brongebruik criteria
    if 'bron' in criterion_name or 'literatuur' in criterion_name:
        return check_source_usage(criterion, section)
    
    # Duidelijkheid criteria
    if 'duidelijk' in criterion_name or 'helder' in criterion_name:
        return check_clarity(criterion, section)
    
    return None

def check_coherence(criterion, section):
    """Controleert coherentie en samenhang"""
    content = section.get('content', '')
    sentences = section.get('sentences', [])
    
    if len(sentences) < 3:
        return None  # Te kort om coherentie te beoordelen
    
    # Eenvoudige coherentie check: zoek naar verbindingswoorden
    transition_words = [
        'daarom', 'dus', 'echter', 'bovendien', 'tevens', 'voorts',
        'ten eerste', 'ten tweede', 'vervolgens', 'tenslotte',
        'omdat', 'doordat', 'waardoor', 'hierdoor', 'daardoor'
    ]
    
    transition_count = sum(1 for word in transition_words if word in content.lower())
    transition_ratio = transition_count / len(sentences)
    
    if transition_ratio < 0.1:  # Minder dan 10% van zinnen heeft verbinding
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'warning',
            'message': "Sectie mist verbindingswoorden voor betere samenhang",
            'suggestion': "Gebruik woorden zoals 'daarom', 'echter', 'bovendien' om zinnen te verbinden",
            'confidence': 0.6,
            'color': '#F9C74F'
        }
    
    return None

def check_argumentation(criterion, section):
    """Controleert kwaliteit van argumentatie"""
    content = section.get('content', '').lower()
    
    # Zoek naar argumentatie indicatoren
    argument_indicators = [
        'omdat', 'doordat', 'aangezien', 'gezien', 'vanwege',
        'ten eerste', 'ten tweede', 'bovendien', 'daarnaast',
        'bijvoorbeeld', 'namelijk', 'immers', 'blijkbaar'
    ]
    
    evidence_indicators = [
        'onderzoek toont', 'studies wijzen', 'bewijs', 'data',
        'statistiek', 'cijfers', 'resultaten', 'bevindingen'
    ]
    
    arg_count = sum(1 for indicator in argument_indicators if indicator in content)
    evidence_count = sum(1 for indicator in evidence_indicators if indicator in content)
    
    word_count = section.get('word_count', 0)
    
    if word_count > 100 and arg_count == 0:
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'warning',
            'message': "Sectie mist duidelijke argumentatie structuur",
            'suggestion': "Gebruik woorden zoals 'omdat', 'ten eerste', 'bijvoorbeeld' om argumenten te structureren",
            'confidence': 0.7,
            'color': '#F9C74F'
        }
    
    return None

def check_source_usage(criterion, section):
    """Controleert brongebruik en referenties"""
    content = section.get('content', '')
    
    # Zoek naar referentie patronen
    import re
    
    # APA style referenties: (Auteur, jaar)
    apa_pattern = r'\([A-Za-z]+,?\s*\d{4}\)'
    apa_refs = re.findall(apa_pattern, content)
    
    # Voetnoten: [1], [2], etc.
    footnote_pattern = r'\[\d+\]'
    footnotes = re.findall(footnote_pattern, content)
    
    # Algemene bron indicatoren
    source_indicators = ['volgens', 'aldus', 'zoals', 'onderzoek van', 'studie van']
    source_mentions = sum(1 for indicator in source_indicators if indicator in content.lower())
    
    total_refs = len(apa_refs) + len(footnotes) + source_mentions
    word_count = section.get('word_count', 0)
    
    # Verwacht minimaal 1 referentie per 200 woorden
    expected_refs = max(1, word_count // 200)
    
    if total_refs < expected_refs:
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'violation',
            'message': f"Sectie heeft {total_refs} referenties, minimaal {expected_refs} verwacht",
            'suggestion': "Voeg meer bronverwijzingen toe om stellingen te onderbouwen",
            'confidence': 0.8,
            'color': criterion['color']
        }
    
    return None

def check_clarity(criterion, section):
    """Controleert duidelijkheid van tekst"""
    content = section.get('content', '')
    sentences = section.get('sentences', [])
    
    if not sentences:
        return None
    
    # Gemiddelde zinslengte
    total_words = sum(len(sentence.split()) for sentence in sentences)
    avg_sentence_length = total_words / len(sentences)
    
    # Te lange zinnen kunnen onduidelijk zijn
    if avg_sentence_length > 25:
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'warning',
            'message': f"Gemiddelde zinslengte is {avg_sentence_length:.1f} woorden (lang)",
            'suggestion': "Overweeg kortere zinnen voor betere leesbaarheid",
            'confidence': 0.6,
            'color': '#F9C74F'
        }
    
    return None

# 9.  WOORDTELLING CONTROLE

def check_word_count(criterion, section):
    """
    Controleert woordtelling criteria
    
    Herkent patronen zoals:
    - "Minimaal 500 woorden"
    - "Tussen 200-400 woorden"
    - "Maximaal 1000 woorden"
    """
    import re
    
    criterion_name = criterion['name'].lower()
    word_count = section.get('word_count', 0)
    
    # Extract getallen uit criterium naam
    numbers = re.findall(r'\d+', criterion_name)
    
    if not numbers:
        return None
    
    # Bepaal type controle
    if 'minimaal' in criterion_name or 'minimum' in criterion_name:
        min_words = int(numbers[0])
        if word_count < min_words:
            return {
                'criterion_id': criterion['id'],
                'criterion_name': criterion['name'],
                'section_id': section['id'],
                'section_name': section['name'],
                'status': 'violation',
                'message': f"Sectie heeft {word_count} woorden, minimaal {min_words} vereist",
                'suggestion': f"Voeg {min_words - word_count} woorden toe",
                'confidence': 1.0,
                'color': criterion['color']
            }
    
    elif 'maximaal' in criterion_name or 'maximum' in criterion_name:
        max_words = int(numbers[0])
        if word_count > max_words:
            return {
                'criterion_id': criterion['id'],
                'criterion_name': criterion['name'],
                'section_id': section['id'],
                'section_name': section['name'],
                'status': 'violation',
                'message': f"Sectie heeft {word_count} woorden, maximaal {max_words} toegestaan",
                'suggestion': f"Verwijder {word_count - max_words} woorden",
                'confidence': 1.0,
                'color': criterion['color']
            }
    
    elif 'tussen' in criterion_name and len(numbers) >= 2:
        min_words = int(numbers[0])
        max_words = int(numbers[1])
        
        if word_count < min_words:
            return {
                'criterion_id': criterion['id'],
                'criterion_name': criterion['name'],
                'section_id': section['id'],
                'section_name': section['name'],
                'status': 'violation',
                'message': f"Sectie heeft {word_count} woorden, {min_words}-{max_words} vereist",
                'suggestion': f"Voeg {min_words - word_count} woorden toe",
                'confidence': 1.0,
                'color': criterion['color']
            }
        elif word_count > max_words:
            return {
                'criterion_id': criterion['id'],
                'criterion_name': criterion['name'],
                'section_id': section['id'],
                'section_name': section['name'],
                'status': 'violation',
                'message': f"Sectie heeft {word_count} woorden, {min_words}-{max_words} vereist",
                'suggestion': f"Verwijder {word_count - max_words} woorden",
                'confidence': 1.0,
                'color': criterion['color']
            }
    
    # Geen overtreding gevonden
    return None

def get_word_count_statistics(sections):
    """Geeft statistieken over woordtelling per sectie"""
    stats = {
        'total_words': 0,
        'sections': [],
        'average_words_per_section': 0,
        'longest_section': None,
        'shortest_section': None
    }
    
    section_word_counts = []
    
    for section in sections:
        if section.get('found', False):
            word_count = section.get('word_count', 0)
            section_word_counts.append(word_count)
            stats['total_words'] += word_count
            
            section_stat = {
                'name': section['name'],
                'word_count': word_count,
                'percentage': 0  # Wordt later berekend
            }
            stats['sections'].append(section_stat)
    
    if section_word_counts:
        stats['average_words_per_section'] = stats['total_words'] / len(section_word_counts)
        
        # Vind langste en kortste sectie
        max_words = max(section_word_counts)
        min_words = min(section_word_counts)
        
        for section_stat in stats['sections']:
            section_stat['percentage'] = (section_stat['word_count'] / stats['total_words']) * 100 if stats['total_words'] > 0 else 0
            
            if section_stat['word_count'] == max_words:
                stats['longest_section'] = section_stat
            if section_stat['word_count'] == min_words:
                stats['shortest_section'] = section_stat
    
    return stats

# 10. DOCX UPLOAD EN TEKST EXTRACTIE

def extract_document_content(filepath):
    """
    Extraheert tekst en structuur uit een .docx bestand
    
    Returns:
    {
        'paragraphs': [{'text': str, 'style': str, 'level': int}],
        'full_text': str,
        'metadata': {'title': str, 'author': str, 'created': datetime}
    }
    """
    try:
        doc = docx.Document(filepath)
        
        paragraphs = []
        full_text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                para_data = {
                    'text': paragraph.text.strip(),
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'level': get_heading_level(paragraph),
                    'is_heading': is_heading_style(paragraph.style.name if paragraph.style else '')
                }
                paragraphs.append(para_data)
                full_text += paragraph.text + "\n"
        
        # Metadata extractie
        props = doc.core_properties
        metadata = {
            'title': props.title or 'Onbekend',
            'author': props.author or 'Onbekend',
            'created': props.created or datetime.now(),
            'word_count': len(full_text.split()),
            'paragraph_count': len(paragraphs)
        }
        
        return {
            'success': True,
            'paragraphs': paragraphs,
            'full_text': full_text,
            'metadata': metadata
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def get_heading_level(paragraph):
    """Bepaal heading level (1-6) op basis van style"""
    style_name = paragraph.style.name if paragraph.style else ''
    
    if 'Heading 1' in style_name or 'Kop 1' in style_name:
        return 1
    elif 'Heading 2' in style_name or 'Kop 2' in style_name:
        return 2
    elif 'Heading 3' in style_name or 'Kop 3' in style_name:
        return 3
    elif 'Heading 4' in style_name or 'Kop 4' in style_name:
        return 4
    elif 'Heading 5' in style_name or 'Kop 5' in style_name:
        return 5
    elif 'Heading 6' in style_name or 'Kop 6' in style_name:
        return 6
    else:
        return 0  # Geen heading

def is_heading_style(style_name):
    """Check of een style een heading is"""
    heading_indicators = ['Heading', 'Kop', 'Title', 'Titel']
    return any(indicator in style_name for indicator in heading_indicators)
    
# 11. SECTIEHERKENNING SYSTEEM

def recognize_sections(paragraphs, document_type='plan_van_aanpak'):
    """
    Herkent secties in het document op basis van headings en patronen
    
    Input: List van paragraph dictionaries
    Output: List van section dictionaries
    
    Returns:
    [
        {
            'id': 'inleiding',
            'name': 'Inleiding', 
            'content': 'Volledige tekst van de sectie...',
            'start_paragraph': 0,
            'end_paragraph': 5,
            'level': 1,
            'subsections': [...],
            'found': True,
            'confidence': 0.95
        }
    ]
    """
    
    # Laad sectie templates voor document type
    section_templates = get_section_templates(document_type)
    
    recognized_sections = []
    current_section = None
    
    for i, paragraph in enumerate(paragraphs):
        if paragraph['is_heading']:
            # Probeer sectie te matchen
            matched_section = match_section_template(paragraph['text'], section_templates)
            
            if matched_section:
                # Sluit vorige sectie af
                if current_section:
                    current_section['end_paragraph'] = i - 1
                    current_section['content'] = extract_section_content(
                        paragraphs, 
                        current_section['start_paragraph'], 
                        current_section['end_paragraph']
                    )
                    recognized_sections.append(current_section)
                
                # Start nieuwe sectie
                current_section = {
                    'id': matched_section['id'],
                    'name': paragraph['text'],
                    'template_name': matched_section['name'],
                    'start_paragraph': i,
                    'end_paragraph': None,
                    'level': paragraph['level'],
                    'found': True,
                    'confidence': matched_section['confidence'],
                    'subsections': []
                }
    
    # Sluit laatste sectie af
    if current_section:
        current_section['end_paragraph'] = len(paragraphs) - 1
        current_section['content'] = extract_section_content(
            paragraphs, 
            current_section['start_paragraph'], 
            current_section['end_paragraph']
        )
        recognized_sections.append(current_section)
    
    # Controleer ontbrekende verplichte secties
    missing_sections = find_missing_sections(recognized_sections, section_templates)
    recognized_sections.extend(missing_sections)
    
    return recognized_sections

def match_section_template(heading_text, templates):
    """Match een heading tegen sectie templates"""
    heading_lower = heading_text.lower().strip()
    
    best_match = None
    best_score = 0
    
    for template in templates:
        # Exacte match
        if heading_lower == template['name'].lower():
            return {'id': template['id'], 'name': template['name'], 'confidence': 1.0}
        
        # Alternative names
        for alt_name in template.get('alternative_names', []):
            if heading_lower == alt_name.lower():
                return {'id': template['id'], 'name': template['name'], 'confidence': 0.9}
        
        # Pattern matching
        if template.get('pattern'):
            import re
            if re.search(template['pattern'], heading_lower):
                confidence = 0.8
                if confidence > best_score:
                    best_match = {'id': template['id'], 'name': template['name'], 'confidence': confidence}
                    best_score = confidence
        
        # Fuzzy matching (eenvoudig)
        similarity = calculate_similarity(heading_lower, template['name'].lower())
        if similarity > 0.7 and similarity > best_score:
            best_match = {'id': template['id'], 'name': template['name'], 'confidence': similarity}
            best_score = similarity
    
    return best_match if best_score > 0.6 else None

def get_section_templates(document_type):
    """Haal sectie templates op voor document type"""
    templates = {
        'plan_van_aanpak': [
            {
                'id': 'inleiding',
                'name': 'Inleiding',
                'alternative_names': ['introductie', 'intro', 'voorwoord'],
                'pattern': r'(inleid|intro)',
                'required': True,
                'order': 1
            },
            {
                'id': 'probleemstelling',
                'name': 'Probleemstelling',
                'alternative_names': ['probleem', 'aanleiding', 'probleemanalyse'],
                'pattern': r'(probleem|aanleiding)',
                'required': True,
                'order': 2
            },
            {
                'id': 'doelstelling',
                'name': 'Doelstelling',
                'alternative_names': ['doel', 'doelen', 'doelstellingen'],
                'pattern': r'(doel|doelstell)',
                'required': True,
                'order': 3
            },
            {
                'id': 'onderzoeksvragen',
                'name': 'Onderzoeksvragen',
                'alternative_names': ['hoofdvraag', 'deelvragen', 'onderzoeksvraag'],
                'pattern': r'(onderzoeks?vra|hoofdvra|deelvra)',
                'required': True,
                'order': 4
            },
            {
                'id': 'methode',
                'name': 'Methode',
                'alternative_names': ['methodologie', 'onderzoeksmethode', 'aanpak'],
                'pattern': r'(methode|methodolog|aanpak)',
                'required': True,
                'order': 5
            },
            {
                'id': 'planning',
                'name': 'Planning',
                'alternative_names': ['tijdsplanning', 'schema', 'tijdschema'],
                'pattern': r'(planning|tijds|schema)',
                'required': True,
                'order': 6
            },
            {
                'id': 'literatuur',
                'name': 'Literatuurlijst',
                'alternative_names': ['bronnen', 'referenties', 'bibliografie'],
                'pattern': r'(literatuur|bronnen|referent|bibliograf)',
                'required': True,
                'order': 7
            }
        ],
        'onderzoeksrapport': [
            # Andere templates voor onderzoeksrapport
        ]
    }
    
    return templates.get(document_type, [])

    # 12.  FEEDBACK GENERATIE SYSTEEM

    def generate_feedback(sections, criteria_list, document_type):
    """
    Genereert feedback op basis van secties en criteria
    
    Returns:
    [
        {
            'criterion_id': 1,
            'criterion_name': 'SMART doelstelling',
            'section_id': 'doelstelling',
            'section_name': 'Doelstelling',
            'status': 'violation',  # 'ok', 'warning', 'violation'
            'message': 'De doelstelling is niet SMART geformuleerd',
            'suggestion': 'Zorg dat de doelstelling Specifiek, Meetbaar, Acceptabel, Realistisch en Tijdgebonden is',
            'location': 'Paragraaf 3, regel 2-4',
            'confidence': 0.8,
            'color': '#F94144'
        }
    ]
    """
    
    feedback_items = []
    
    for criterion in criteria_list:
        if not criterion['is_enabled']:
            continue
            
        # Bepaal welke secties dit criterium moet controleren
        applicable_sections = get_applicable_sections(criterion, sections, document_type)
        
        for section in applicable_sections:
            feedback_item = apply_criterion_to_section(criterion, section)
            if feedback_item:
                feedback_items.append(feedback_item)
    
    # Sorteer feedback op ernst en sectie
    feedback_items.sort(key=lambda x: (x['status'], x['section_id']))
    
    return feedback_items

def apply_criterion_to_section(criterion, section):
    """Past een criterium toe op een specifieke sectie"""
    
    # Verschillende types criteria
    if criterion['rule_type'] == 'tekstueel':
        return check_textual_criterion(criterion, section)
    elif criterion['rule_type'] == 'structureel':
        return check_structural_criterion(criterion, section)
    elif criterion['rule_type'] == 'inhoudelijk':
        return check_content_criterion(criterion, section)
    
    return None

def check_textual_criterion(criterion, section):
    """Controleert tekstuele criteria (woordgebruik, zinsbouw, etc.)"""
    
    # Voorbeeld: SMART criteria check
    if 'smart' in criterion['name'].lower():
        return check_smart_formulation(criterion, section)
    
    # Voorbeeld: Woordtelling
    if 'woorden' in criterion['name'].lower():
        return check_word_count(criterion, section)
    
    return None

def check_smart_formulation(criterion, section):
    """Controleert of tekst SMART geformuleerd is"""
    content = section['content'].lower()
    
    smart_indicators = {
        'specifiek': ['specifiek', 'concreet', 'duidelijk'],
        'meetbaar': ['meetbaar', 'kwantificeerbaar', 'cijfers', 'percentage'],
        'acceptabel': ['acceptabel', 'haalbaar', 'realistisch'],
        'realistisch': ['realistisch', 'haalbaar', 'mogelijk'],
        'tijdgebonden': ['deadline', 'datum', 'week', 'maand', 'jaar', 'tijd']
    }
    
    found_aspects = []
    missing_aspects = []
    
    for aspect, indicators in smart_indicators.items():
        if any(indicator in content for indicator in indicators):
            found_aspects.append(aspect)
        else:
            missing_aspects.append(aspect)
    
    if len(missing_aspects) > 2:
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'violation',
            'message': f"De {section['name'].lower()} mist SMART aspecten: {', '.join(missing_aspects)}",
            'suggestion': f"Voeg de volgende aspecten toe: {', '.join(missing_aspects)}",
            'confidence': 0.7,
            'color': criterion['color']
        }
    elif len(missing_aspects) > 0:
        return {
            'criterion_id': criterion['id'],
            'criterion_name': criterion['name'],
            'section_id': section['id'],
            'section_name': section['name'],
            'status': 'warning',
            'message': f"De {section['name'].lower()} zou kunnen verbeteren op: {', '.join(missing_aspects)}",
            'suggestion': f"Overweeg om deze aspecten toe te voegen: {', '.join(missing_aspects)}",
            'confidence': 0.6,
            'color': '#F9C74F'
        }
    
    return None  # Geen feedback = criterium is OK

# 13. DATABASE SCHEMA

-- Document types (Plan van Aanpak, Onderzoeksrapport, etc.)
CREATE TABLE document_types (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL UNIQUE,
    description TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Secties per document type (hiërarchisch)
CREATE TABLE sections (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    document_type_id INTEGER,
    name TEXT NOT NULL,
    identifier TEXT,  -- 'inleiding', 'probleemstelling', etc.
    parent_id INTEGER,  -- Voor subsecties
    level INTEGER DEFAULT 1,  -- 1=hoofdsectie, 2=subsectie, etc.
    order_index INTEGER,  -- Volgorde in document
    alternative_names TEXT,  -- JSON array van alternatieve namen
    pattern TEXT,  -- Regex pattern voor herkenning
    is_required BOOLEAN DEFAULT 1,
    description TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (document_type_id) REFERENCES document_types (id),
    FOREIGN KEY (parent_id) REFERENCES sections (id)
);

-- Criteria (beoordelingsregels)
CREATE TABLE criteria (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL,
    description TEXT,
    error_message TEXT,  -- Bericht bij overtreding
    instruction_link TEXT,  -- Link naar uitleg
    category TEXT DEFAULT 'Algemeen',
    color TEXT DEFAULT '#F94144',
    rule_type TEXT DEFAULT 'tekstueel',  -- 'tekstueel', 'structureel', 'inhoudelijk'
    application_scope TEXT DEFAULT 'specific_sections',  -- 'all', 'specific_sections', 'exclude_sections'
    max_occurrences_per TEXT DEFAULT 'paragraph',  -- 'document', 'section', 'paragraph'
    severity TEXT DEFAULT 'error',  -- 'error', 'warning', 'info'
    is_enabled BOOLEAN DEFAULT 1,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Koppeling criteria aan secties
CREATE TABLE criteria_section_mappings (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    criteria_id INTEGER,
    section_id INTEGER,
    document_type_id INTEGER,
    is_excluded BOOLEAN DEFAULT 0,  -- True = criterium NIET toepassen op deze sectie
    weight REAL DEFAULT 1.0,  -- Gewicht van criterium in deze sectie
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (criteria_id) REFERENCES criteria (id),
    FOREIGN KEY (section_id) REFERENCES sections (id),
    FOREIGN KEY (document_type_id) REFERENCES document_types (id)
);

-- Geüploade documenten
CREATE TABLE documents (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    filename TEXT NOT NULL,
    original_filename TEXT,
    document_type_id INTEGER,
    file_path TEXT,
    file_size INTEGER,
    upload_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    analysis_status TEXT DEFAULT 'pending',  -- 'pending', 'completed', 'failed'
    analysis_data TEXT,  -- JSON met analyse resultaten
    FOREIGN KEY (document_type_id) REFERENCES document_types (id)
);

-- Feedback items per document
CREATE TABLE feedback_items (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id INTEGER,
    criteria_id INTEGER,
    section_id INTEGER,
    status TEXT,  -- 'ok', 'warning', 'violation'
    message TEXT,
    suggestion TEXT,
    location TEXT,  -- Waar in document (paragraaf, regel)
    confidence REAL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (document_id) REFERENCES documents (id),
    FOREIGN KEY (criteria_id) REFERENCES criteria (id),
    FOREIGN KEY (section_id) REFERENCES sections (id)
);

# 14. Voorbeelddata:

def populate_example_data():
    """Vult database met voorbeelddata"""
    
    # Document types
    document_types = [
        ('Plan van Aanpak', 'Onderzoeksplan voor HBO-opdrachten'),
        ('Onderzoeksrapport', 'Eindrapport van uitgevoerd onderzoek'),
        ('Thesis', 'Afstudeerscriptie'),
        ('Essay', 'Korte academische tekst')
    ]
    
    # Secties voor Plan van Aanpak
    pva_sections = [
        ('Inleiding', 'inleiding', 1, 1, '["introductie", "intro", "voorwoord"]', r'(inleid|intro)'),
        ('Probleemstelling', 'probleemstelling', 1, 2, '["probleem", "aanleiding"]', r'(probleem|aanleiding)'),
        ('Doelstelling', 'doelstelling', 1, 3, '["doel", "doelen"]', r'(doel|doelstell)'),
        ('Onderzoeksvragen', 'onderzoeksvragen', 1, 4, '["hoofdvraag", "deelvragen"]', r'(onderzoeks?vra)'),
        ('Methode', 'methode', 1, 5, '["methodologie", "aanpak"]', r'(methode|aanpak)'),
        ('Planning', 'planning', 1, 6, '["tijdsplanning", "schema"]', r'(planning|tijds)'),
        ('Literatuurlijst', 'literatuur', 1, 7, '["bronnen", "referenties"]', r'(literatuur|bronnen)')
    ]
    
    # Criteria voorbeelden
    criteria_examples = [
        ('SMART doelstelling', 'Doelstelling moet SMART geformuleerd zijn', 'tekstueel', 'Inhoudelijk'),
        ('Hoofdvraag aansluiting', 'Hoofdvraag sluit aan bij probleemstelling', 'inhoudelijk', 'Inhoudelijk'),
        ('APA bronvermelding', 'Bronnen volgens APA-stijl', 'tekstueel', 'Referenties'),
        ('Minimaal 500 woorden', 'Sectie moet minimaal 500 woorden bevatten', 'structureel', 'Structuur')
    ]

# 15. UI CODE VOOR ANALYSE WEERGAVE

ANALYSIS_TEMPLATE = '''
{% block content %}
    <div class="header">
        <h1>📊 Document Analyse</h1>
        <p>{{ document.original_filename }}</p>
        <div class="nav-breadcrumb">
            <a href="/">Home</a> > <a href="/documents">Documenten</a> > Analyse
        </div>
    </div>
    
    <!-- Document Info Card -->
    <div class="card">
        <h3>📄 Document Informatie</h3>
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px;">
            <div>
                <strong>Bestand:</strong><br>{{ document.original_filename }}
            </div>
            <div>
                <strong>Type:</strong><br>{{ document_type.name }}
            </div>
            <div>
                <strong>Upload:</strong><br>{{ document.upload_time.strftime('%d-%m-%Y %H:%M') }}
            </div>
            <div>
                <strong>Status:</strong><br>
                {% if document.analysis_status == 'completed' %}
                    <span class="badge badge-success">✅ Geanalyseerd</span>
                {% elif document.analysis_status == 'failed' %}
                    <span class="badge badge-warning">❌ Fout</span>
                {% else %}
                    <span class="badge" style="background: #F9C74F; color: #2B2D42;">⏳ Bezig</span>
                {% endif %}
            </div>
        </div>
    </div>
    
    <!-- Sectie Overzicht -->
    <div class="card">
        <h3>🔍 Gevonden Secties</h3>
        <table class="table">
            <thead>
                <tr>
                    <th>Sectie</th>
                    <th>Status</th>
                    <th>Woorden</th>
                    <th>Confidence</th>
                    <th>Acties</th>
                </tr>
            </thead>
            <tbody>
                {% for section in sections %}
                <tr>
                    <td>
                        <strong>{{ section.name }}</strong>
                        {% if section.level > 1 %}
                            <small style="color: #6C757D;">(Niveau {{ section.level }})</small>
                        {% endif %}
                    </td>
                    <td>
                        {% if section.found %}
                            <span class="badge badge-success">✅ Gevonden</span>
                        {% else %}
                            <span class="badge badge-warning">❌ Ontbreekt</span>
                        {% endif %}
                    </td>
                    <td>{{ section.word_count or 0 }}</td>
                    <td>
                        {% if section.confidence %}
                            <div style="background: #E9ECEF; border-radius: 10px; height: 20px; width: 100px; position: relative;">
                                <div style="background: #84A98C; height: 100%; width: {{ (section.confidence * 100)|round }}%; border-radius: 10px;"></div>
                                <span style="position: absolute; top: 0; left: 50%; transform: translateX(-50%); font-size: 12px; line-height: 20px;">{{ (section.confidence * 100)|round }}%</span>
                            </div>
                        {% endif %}
                    </td>
                    <td>
                        <button class="btn" style="padding: 4px 8px; font-size: 12px;" onclick="showSectionContent('{{ section.id }}')">👁️ Bekijk</button>
                    </td>
                </tr>
                {% endfor %}
            </tbody>
        </table>
    </div>
    
    <!-- Feedback Overzicht -->
    <div class="card">
        <h3>💬 Feedback & Suggesties</h3>
        
        <!-- Feedback statistieken -->
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 15px; margin-bottom: 20px;">
            <div style="text-align: center; padding: 15px; background: #f8f9fa; border-radius: 8px;">
                <div style="font-size: 24px; font-weight: bold; color: #F94144;">{{ feedback_stats.violations }}</div>
                <div style="font-size: 14px; color: #6C757D;">Overtredingen</div>
            </div>
            <div style="text-align: center; padding: 15px; background: #f8f9fa; border-radius: 8px;">
                <div style="font-size: 24px; font-weight: bold; color: #F9C74F;">{{ feedback_stats.warnings }}</div>
                <div style="font-size: 14px; color: #6C757D;">Waarschuwingen</div>
            </div>
            <div style="text-align: center; padding: 15px; background: #f8f9fa; border-radius: 8px;">
                <div style="font-size: 24px; font-weight: bold; color: #84A98C;">{{ feedback_stats.passed }}</div>
                <div style="font-size: 14px; color: #6C757D;">Correct</div>
            </div>
        </div>
        
        <!-- Feedback items -->
        {% for feedback in feedback_items %}
        <div style="background: #f8f9fa; border-left: 4px solid {{ feedback.color }}; padding: 15px; margin: 10px 0; border-radius: 0 8px 8px 0;">
            <div style="display: flex; justify-content: space-between; align-items: start;">
                <div style="flex: 1;">
                    <h4 style="color: {{ feedback.color }}; margin-bottom: 5px;">
                        {{ feedback.criterion_name }}
                        <small style="color: #6C757D; font-weight: normal;">({{ feedback.section_name }})</small>
                    </h4>
                    <p style="margin-bottom: 10px;">{{ feedback.message }}</p>
                    {% if feedback.suggestion %}
                    <p style="background: rgba(255,255,255,0.7); padding: 10px; border-radius: 5px; margin-bottom: 10px;">
                        <strong>💡 Suggestie:</strong> {{ feedback.suggestion }}
                    </p>
                    {% endif %}
                    {% if feedback.location %}
                    <small style="color: #6C757D;">📍 Locatie: {{ feedback.location }}</small>
                    {% endif %}
                </div>
                <div style="margin-left: 15px;">
                    {% if feedback.status == 'violation' %}
                        <span class="badge" style="background: #F94144; color: white;">❌ Fout</span>
                    {% elif feedback.status == 'warning' %}
                        <span class="badge" style="background: #F9C74F; color: #2B2D42;">⚠️ Waarschuwing</span>
                    {% else %}
                        <span class="badge badge-success">✅ OK</span>
                    {% endif %}
                </div>
            </div>
        </div>
        {% endfor %}
    </div>
    
    <!-- Acties -->
    <div class="card">
        <h3>🔄 Acties</h3>
        <a href="/documents/{{ document.id }}/export" class="btn">📄 Export PDF Rapport</a>
        <a href="/documents/{{ document.id }}/reanalyze" class="btn btn-secondary">🔄 Opnieuw Analyseren</a>
        <a href="/documents/upload" class="btn">📤 Nieuw Document</a>
        <a href="/documents" class="btn btn-secondary">📋 Terug naar Overzicht</a>
    </div>
    
    <!-- JavaScript voor interactiviteit -->
    <script>
    function showSectionContent(sectionId) {
        // Toon sectie inhoud in modal of uitklapbaar gebied
        alert('Sectie inhoud voor: ' + sectionId);
    }
    </script>
{% endblock %}
'''
# 16. DOCUMENT UPLOAD FUNCTIONALITEIT - TOEVOEGING AAN FEEDBACK TOOL

# EXTRA IMPORTS TOEVOEGEN (bovenaan het bestand)

from werkzeug.utils import secure_filename
import docx  # pip install python-docx
import re
import json

# CONFIGURATIE TOEVOEGEN (na app.config['DATABASE'])

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'docx', 'doc'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# DOCUMENT UPLOAD ROUTES TOEVOEGEN

@app.route('/documents')
def documents_list():
    """Overzicht van geüploade documenten"""
    return render_template_string(BASE_TEMPLATE + '''
    {% block title %}Document Beheer - Feedback Tool{% endblock %}
    {% block content %}
        <div class="header">
            <h1>📄 Document Beheer</h1>
            <p>Upload en analyseer documenten voor feedback</p>
            <div class="nav-breadcrumb">
                <a href="/">Home</a> > Document Beheer
            </div>
        </div>
        
        <div class="card">
            <h3>Document Upload</h3>
            <p>Upload een Word-document (.docx) voor automatische analyse en feedback.</p>
            <a href="/documents/upload" class="btn">📤 Document Uploaden</a>
        </div>
        
        <div class="card">
            <h3>Recente Uploads</h3>
            <p>Hier komen de recent geüploade documenten te staan.</p>
            <div class="empty-state">
                <h3>Nog geen documenten geüpload</h3>
                <p>Upload uw eerste document om te beginnen met analyse.</p>
            </div>
        </div>
    {% endblock %}
    ''')

@app.route('/documents/upload', methods=['GET', 'POST'])
def document_upload():
    """Document upload formulier en verwerking"""
    if request.method == 'POST':
        # Controleer of er een bestand is geüpload
        if 'document' not in request.files:
            flash('Geen bestand geselecteerd.', 'error')
            return redirect(request.url)
        
        file = request.files['document']
        document_type = request.form.get('document_type', 'plan_van_aanpak')
        
        if file.filename == '':
            flash('Geen bestand geselecteerd.', 'error')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            try:
                # Veilige bestandsnaam maken
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"{timestamp}_{filename}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                
                # Bestand opslaan
                file.save(filepath)
                
                # Document analyseren
                analysis_result = analyze_document(filepath, document_type)
                
                if analysis_result['success']:
                    flash(f'Document succesvol geüpload en geanalyseerd! {analysis_result["sections_found"]} secties gevonden.', 'success')
                    return redirect(url_for('document_view', filename=filename))
                else:
                    flash(f'Document geüpload maar analyse mislukt: {analysis_result["error"]}', 'error')
                    
            except Exception as e:
                flash(f'Fout bij uploaden: {str(e)}', 'error')
        else:
            flash('Alleen .docx bestanden zijn toegestaan.', 'error')
    
    return render_template_string(BASE_TEMPLATE + '''
    {% block title %}Document Uploaden - Feedback Tool{% endblock %}
    {% block content %}
        <div class="header">
            <h1>📤 Document Uploaden</h1>
            <p>Upload een Word-document voor automatische analyse</p>
            <div class="nav-breadcrumb">
                <a href="/">Home</a> > <a href="/documents">Documenten</a> > Upload
            </div>
        </div>
        
        <div class="card">
            {% with messages = get_flashed_messages(with_categories=true) %}
                {% if messages %}
                    <div class="flash-messages">
                        {% for category, message in messages %}
                            <div class="flash-{{ category }}">{{ message }}</div>
                        {% endfor %}
                    </div>
                {% endif %}
            {% endwith %}
            
            <form method="POST" enctype="multipart/form-data">
                <div class="form-group">
                    <label for="document">Document selecteren *</label>
                    <input type="file" id="document" name="document" accept=".docx,.doc" required>
                    <small style="color: #6C757D;">Alleen Word-documenten (.docx) zijn toegestaan. Maximaal 16MB.</small>
                </div>
                
                <div class="form-group">
                    <label for="document_type">Document Type</label>
                    <select id="document_type" name="document_type">
                        <option value="plan_van_aanpak" selected>Plan van Aanpak</option>
                        <option value="onderzoeksrapport">Onderzoeksrapport</option>
                        <option value="thesis">Thesis</option>
                        <option value="essay">Essay</option>
                    </select>
                </div>
                
                <div style="background: #e8f5e8; border: 1px solid #4CAF50; padding: 15px; border-radius: 8px; margin: 20px 0;">
                    <h4 style="margin-bottom: 10px;">📋 Wat gebeurt er bij upload?</h4>
                    <ul style="margin-left: 20px;">
                        <li>Document wordt geanalyseerd op structuur</li>
                        <li>Secties worden automatisch herkend</li>
                        <li>Criteria worden toegepast per sectie</li>
                        <li>Feedback rapport wordt gegenereerd</li>
                    </ul>
                </div>
                
                <div style="margin-top: 30px;">
                    <button type="submit" class="btn btn-success">📤 Document Uploaden & Analyseren</button>
                    <a href="/documents" class="btn btn-secondary">❌ Annuleren</a>
                </div>
            </form>
        </div>
    {% endblock %}
    ''')

@app.route('/documents/view/<filename>')
def document_view(filename):
    """Bekijk geanalyseerd document"""
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        flash('Document niet gevonden.', 'error')
        return redirect(url_for('documents_list'))
    
    # Laad analyse resultaten (zou in database moeten staan in productie)
    analysis_file = filepath.replace('.docx', '_analysis.json')
    analysis_data = {}
    
    if os.path.exists(analysis_file):
        try:
            with open(analysis_file, 'r', encoding='utf-8') as f:
                analysis_data = json.load(f)
        except:
            analysis_data = {'error': 'Kon analyse niet laden'}
    
    return render_template_string(BASE_TEMPLATE + '''
    {% block title %}Document Analyse - {{ filename }}{% endblock %}
    {% block content %}
        <div class="header">
            <h1>📊 Document Analyse</h1>
            <p>Resultaten voor: {{ filename }}</p>
            <div class="nav-breadcrumb">
                <a href="/">Home</a> > <a href="/documents">Documenten</a> > Analyse
            </div>
        </div>
        
        <div class="card">
            <h3>📄 Document Informatie</h3>
            <table class="table">
                <tr><td><strong>Bestandsnaam:</strong></td><td>{{ filename }}</td></tr>
                <tr><td><strong>Upload tijd:</strong></td><td>{{ analysis_data.get('upload_time', 'Onbekend') }}</td></tr>
                <tr><td><strong>Document type:</strong></td><td>{{ analysis_data.get('document_type', 'Onbekend') }}</td></tr>
                <tr><td><strong>Aantal secties:</strong></td><td>{{ analysis_data.get('sections_found', 0) }}</td></tr>
            </table>
        </div>
        
        {% if analysis_data.get('sections') %}
        <div class="card">
            <h3>🔍 Gevonden Secties</h3>
            <table class="table">
                <thead>
                    <tr>
                        <th>Sectie</th>
                        <th>Type</th>
                        <th>Inhoud Preview</th>
                        <th>Status</th>
                    </tr>
                </thead>
                <tbody>
                    {% for section in analysis_data.sections %}
                    <tr>
                        <td><strong>{{ section.name }}</strong></td>
                        <td><span class="badge" style="background-color: {{ section.color or '#4D908E' }}; color: white;">{{ section.type or 'Algemeen' }}</span></td>
                        <td>{{ (section.content or '')[:100] }}{% if (section.content or '')|length > 100 %}...{% endif %}</td>
                        <td>
                            {% if section.found %}
                                <span class="badge badge-success">✅ Gevonden</span>
                            {% else %}
                                <span class="badge badge-warning">⚠️ Ontbreekt</span>
                            {% endif %}
                        </td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
        {% endif %}
        
        {% if analysis_data.get('feedback') %}
        <div class="card">
            <h3>💬 Feedback & Suggesties</h3>
            {% for feedback_item in analysis_data.feedback %}
            <div style="background: #f8f9fa; border-left: 4px solid {{ feedback_item.color or '#F94144' }}; padding: 15px; margin: 10px 0; border-radius: 0 8px 8px 0;">
                <h4 style="color: {{ feedback_item.color or '#F94144' }};">{{ feedback_item.title }}</h4>
                <p>{{ feedback_item.message }}</p>
                {% if feedback_item.suggestion %}
                <p><strong>Suggestie:</strong> {{ feedback_item.suggestion }}</p>
                {% endif %}
            </div>
            {% endfor %}
        </div>
        {% endif %}
        
        <div class="card">
            <h3>🔄 Acties</h3>
            <a href="/documents/upload" class="btn">📤 Nieuw Document Uploaden</a>
            <a href="/documents" class="btn btn-secondary">📋 Terug naar Overzicht</a>
        </div>
    {% endblock %}
    ''', filename=filename, analysis_data=analysis_data)

# === DOCUMENT ANALYSE FUNCTIE ===
def analyze_document(filepath, document_type):
    """Analyseer een Word-document en genereer feedback"""
    try:
        # Laad het Word-document
        doc = docx.Document(filepath)
        
        # Extract tekst en structuur
        sections_found = []
        full_text = ""
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text += text + "\n"
                
                # Eenvoudige sectieherkenning
                if is_section_header(text):
                    sections_found.append({
                        'name': text,
                        'type': classify_section(text),
                        'content': text,
                        'found': True,
                        'color': get_section_color(text)
                    })
        
        # Genereer feedback
        feedback = generate_feedback(sections_found, full_text, document_type)
        
        # Sla analyse op
        analysis_data = {
            'upload_time': datetime.now().isoformat(),
            'document_type': document_type,
            'sections_found': len(sections_found),
            'sections': sections_found,
            'feedback': feedback,
            'full_text_length': len(full_text)
        }
        
        # Sla analyse resultaat op
        analysis_file = filepath.replace('.docx', '_analysis.json')
        with open(analysis_file, 'w', encoding='utf-8') as f:
            json.dump(analysis_data, f, ensure_ascii=False, indent=2)
        
        return {
            'success': True,
            'sections_found': len(sections_found),
            'analysis_data': analysis_data
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def is_section_header(text):
    """Bepaal of een tekst een sectie header is"""
    # Eenvoudige heuristieken
    if len(text) > 100:  # Te lang voor een header
        return False
    
    # Patronen voor headers
    header_patterns = [
        r'^\d+\.?\s+[A-Z]',  # "1. Inleiding" of "1 Inleiding"
        r'^[A-Z][a-z]+\s*$',  # "Inleiding"
        r'^[A-Z\s]+$',  # "INLEIDING"
        r'^\d+\.\d+',  # "1.1 Subkop"
    ]
    
    for pattern in header_patterns:
        if re.match(pattern, text):
            return True
    
    return False

def classify_section(text):
    """Classificeer het type sectie"""
    text_lower = text.lower()
    
    if any(word in text_lower for word in ['inleiding', 'introductie']):
        return 'Inleiding'
    elif any(word in text_lower for word in ['probleem', 'aanleiding']):
        return 'Probleemstelling'
    elif any(word in text_lower for word in ['onderzoek', 'methode']):
        return 'Onderzoek'
    elif any(word in text_lower for word in ['conclusie', 'afsluiting']):
        return 'Conclusie'
    elif any(word in text_lower for word in ['literatuur', 'bronnen']):
        return 'Literatuur'
    else:
        return 'Algemeen'

def get_section_color(text):
    """Geef kleur op basis van sectie type"""
    section_type = classify_section(text)
    colors = {
        'Inleiding': '#4D908E',
        'Probleemstelling': '#F94144',
        'Onderzoek': '#84A98C',
        'Conclusie': '#F6BD60',
        'Literatuur': '#52796F',
        'Algemeen': '#6C757D'
    }
    return colors.get(section_type, '#6C757D')

def generate_feedback(sections, full_text, document_type):
    """Genereer feedback op basis van gevonden secties"""
    feedback = []
    
    # Controleer verplichte secties voor Plan van Aanpak
    if document_type == 'plan_van_aanpak':
        required_sections = ['inleiding', 'probleem', 'onderzoek', 'planning']
        found_sections = [s['name'].lower() for s in sections]
        
        for required in required_sections:
            if not any(required in found for found in found_sections):
                feedback.append({
                    'title': f'Ontbrekende sectie: {required.title()}',
                    'message': f'De sectie "{required.title()}" is niet gevonden in het document.',
                    'suggestion': f'Voeg een sectie toe met de titel "{required.title()}" of een vergelijkbare naam.',
                    'color': '#F94144'
                })
    
    # Algemene feedback
    if len(full_text) < 500:
        feedback.append({
            'title': 'Document te kort',
            'message': 'Het document lijkt erg kort te zijn.',
            'suggestion': 'Zorg voor voldoende uitwerking van alle onderdelen.',
            'color': '#F9C74F'
        })
    
    if len(sections) < 3:
        feedback.append({
            'title': 'Weinig structuur',
            'message': 'Er zijn weinig duidelijke secties gevonden.',
            'suggestion': 'Gebruik duidelijke kopjes om de structuur te verbeteren.',
            'color': '#F9C74F'
        })
    
    # Positieve feedback
    if len(sections) >= 5:
        feedback.append({
            'title': 'Goede structuur',
            'message': 'Het document heeft een duidelijke structuur met meerdere secties.',
            'suggestion': '',
            'color': '#84A98C'
        })
    
    return feedback

# NAVIGATIE TOEVOEGEN AAN HOOFDPAGINA

# In de index() route, voeg toe aan de content:

        <div class="card">
            <h3>📄 Document Beheer</h3>
            <p>Upload en analyseer documenten voor automatische feedback.</p>
            <a href="/documents" class="btn">Beheer Documenten</a>
        </div>


