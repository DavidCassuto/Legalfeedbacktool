"""
Inline Word comments module.

Voegt comments toe op specifieke locaties in een Word document bij criteria-afwijkingen.

Regels:
- Alles behalve status 'ok' krijgt een comment (ook 'info', 'warning', 'violation', 'error').
- Als de afwijking te herleiden is tot een sub-paragraaf (H2/H3), wordt het comment
  ALLEEN daar geplaatst, NIET nogmaals bij de hogere sectie (H1).
- Elke afwijking krijgt een eigen comment, verankerd aan de kortst mogelijke tekst.

Technische aanpak (vermijdt 'unreadable content' in Word):
- Alle originele bestanden worden ONGEWIJZIGD overgenomen uit het originele docx.
- Alleen word/document.xml wordt via lxml minimaal aangepast (comment markers).
- word/comments.xml wordt toegevoegd.
- [Content_Types].xml en _rels worden bijgewerkt via string-insertie (geen herserializatie).
"""

import os
import re
import shutil
import zipfile
import logging
from copy import deepcopy
from datetime import datetime, timezone
from typing import Optional, List, Dict, Any, Tuple

from docx import Document
from lxml import etree

logger = logging.getLogger(__name__)

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
COMMENTS_REL_TYPE = (
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
)
COMMENTS_CT = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
)


# ---------------------------------------------------------------------------
# Paragraaf-structuur helpers  (python-docx, alleen lezen)
# ---------------------------------------------------------------------------

def _heading_level(paragraph) -> int:
    """Geeft het heading-niveau van een paragraaf (0 = geen heading).
    Werkt voor zowel Engelse ('Heading 1') als Nederlandse ('Kop 1') stijlnamen.
    """
    style = paragraph.style
    if style is None:
        return 0
    # style_id is altijd Engels in OOXML: "Heading1", "Heading2", ongeacht documenttaal
    style_id = getattr(style, 'style_id', '') or ''
    if style_id.startswith('Heading'):
        try:
            return int(style_id[len('Heading'):])
        except ValueError:
            return 1
    # Fallback: controleer ook de stijlnaam (Engels of gelokaliseerd)
    name = style.name or ''
    if name.startswith('Heading'):
        try:
            return int(name.split()[-1])
        except (ValueError, IndexError):
            return 1
    # Nederlandse heading-namen
    name_lower = name.lower()
    for nl, level in [('kop 1', 1), ('kop 2', 2), ('kop 3', 3), ('kop 4', 4),
                      ('kop1', 1), ('kop2', 2), ('kop3', 3), ('kop4', 4)]:
        if name_lower == nl:
            return level
    return 0


def _build_para_structure(doc: Document) -> List[Dict]:
    """
    Bouw een platte structuurlijst van alle paragrafen.
    Gebruikt python-docx ALLEEN voor lezen (geen save).
    """
    return [
        {
            'idx':        i,
            'para':       p,
            'level':      _heading_level(p),
            'text':       p.text.strip(),
            'is_heading': _heading_level(p) > 0,
        }
        for i, p in enumerate(doc.paragraphs)
    ]


# ---------------------------------------------------------------------------
# Plaatsbepalingslogica (hierarchy-regel)
# ---------------------------------------------------------------------------

def _is_toc_line(text: str) -> bool:
    """Herken een inhoudsopgave-regel: punt-leiders (....) of een tab/spaties gevolgd
    door een paginanummer. Zulke regels herhalen hoofdstuktitels en veroorzaken
    foute comment-plaatsingen — daarom uitsluiten bij het plaatsen."""
    t = (text or '').strip()
    if not t:
        return False
    if re.search(r'\.{4,}', t) or '…' in t or '\t' in t:
        return True
    # "Iets ......  78"  of  "Hoofdstuk 3 ... 33" zonder expliciete leaders maar wel eindigend op paginanr
    if re.search(r'[ .·]{6,}\d{1,4}$', t):
        return True
    return False


def _find_target_paragraph_idx(
    para_structure: List[Dict],
    section_name: str,
    offending_snippet: Optional[str],
    section_heading_texts: Optional[Dict[str, str]] = None,
) -> Tuple[int, str]:
    """
    Bepaal de meest specifieke paragraaf voor een comment.

    Algoritme:
      1. Zoek de sectie heading (H1/H2) op naam → bepaal sectiegrenzen.
         1a. Direct: sectienaam zit in heading-tekst.
         1b. Fallback: gebruik de opgeslagen heading_text uit sectie-herkenning
             (voor alias/fuzzy-matches: bijv. 'Doelstelling' → '1.4 Doel van het onderzoek').
      2. Als er een offending_snippet is:
           a. Zoek de paragraaf BINNEN de sectie die de snippet bevat.
           b. Zoek de dichtstbijzijnde sub-heading (H2/H3) vóór die paragraaf.
           c. Gebruik die sub-heading → comment staat bij H2/H3, NIET bij H1.
      3. Geen snippet of geen match: gebruik de sectie heading zelf.

    Returns: (para_idx, location_description)
    """
    name_lower = section_name.lower()

    # Stap 1a: directe match – sectienaam zit in heading-tekst
    # Prefer exact match (na verwijdering nummering) boven substring match.
    # Bijv. 'deelvragen' moet matchen op '1.3.2 Deelvragen' (exact)
    # en NIET op '1.3 Hoofd- en deelvragen' (substring).
    heading_idx = None
    substring_match_idx = None
    for item in para_structure:
        if not item['is_heading']:
            continue
        htxt = item['text'].lower()
        if name_lower not in htxt:
            continue
        # Verwijder nummeringsprefix (bijv. "1.3.2 ") voor exacte vergelijking
        cleaned = re.sub(r'^[\d.]+\s*', '', htxt).strip()
        if cleaned == name_lower:
            heading_idx = item['idx']
            break  # exacte match – direct stoppen
        elif substring_match_idx is None:
            substring_match_idx = item['idx']
    if heading_idx is None and substring_match_idx is not None:
        heading_idx = substring_match_idx

    # Stap 1b: fallback via opgeslagen heading_text (alias/fuzzy-matches)
    if heading_idx is None and section_heading_texts:
        raw_heading = section_heading_texts.get(section_name, '')
        if raw_heading:
            raw_lower = raw_heading.lower()
            for item in para_structure:
                if item['is_heading'] and raw_lower in item['text'].lower():
                    heading_idx = item['idx']
                    logger.debug("heading_text-fallback: '%s' → heading '%s' → para %d",
                                 section_name, raw_heading, heading_idx)
                    break

    # Stap 1c: prefix-overlap fallback (werkt ook zonder opgeslagen heading_text)
    # Bijv. 'Doelstelling' → 'doel' in '1.4 Doel van het onderzoek' (4-char prefix)
    #       'Risicoanalyse' → 'risico' in '2.6 ... risicobenadering' (6-char prefix)
    if heading_idx is None:
        def _cpfx(a: str, b: str) -> int:
            n = 0
            for x, y in zip(a, b):
                if x != y:
                    return n
                n += 1
            return n

        def _sig_words(text: str) -> List[str]:
            return [w for w in re.sub(r'[^\w]', ' ', text.lower()).split()
                    if len(w) >= 4 and not w.isdigit()]

        name_words = _sig_words(name_lower)
        best_score, best_idx = 0, None
        for item in para_structure:
            if not item['is_heading']:
                continue
            h_words = _sig_words(item['text'])
            score = 0
            for nw in name_words:
                for hw in h_words:
                    plen = _cpfx(nw, hw)
                    if plen >= 4:
                        score += plen
            if score > best_score:
                best_score = score
                best_idx   = item['idx']
        if best_score >= 4 and best_idx is not None:
            heading_idx = best_idx
            logger.debug("prefix-fallback: '%s' → para %d (score=%d)",
                         section_name, heading_idx, best_score)

    if heading_idx is None:
        # Geen sectie-heading gevonden (bijv. document_only criterium).
        # Als er een snippet is: zoek door het hele document.
        # Anders: val terug op paragraaf 0.
        if offending_snippet and len(offending_snippet.strip()) >= 5:
            section_items = para_structure  # heel document doorzoeken
        else:
            return 0, 'document'
    else:
        section_level = para_structure[heading_idx]['level']
        section_end = len(para_structure)
        for item in para_structure:
            if (item['idx'] > heading_idx
                    and item['is_heading']
                    and item['level'] <= section_level):
                section_end = item['idx']
                break

        section_items = [
            it for it in para_structure
            if heading_idx <= it['idx'] < section_end
        ]

    if offending_snippet and len(offending_snippet.strip()) >= 5:
        snippet_lower = re.sub(r'\s+', ' ', offending_snippet.lower().strip())
        # Inhoudsopgave-regels uitsluiten: die herhalen hoofdstuktitels en geven foute matches
        section_items = [it for it in section_items if not _is_toc_line(it.get('text', ''))]

        # Poging 1: exacte substring-match (snippet volledig in één paragraaf)
        match_idx = None
        for item in section_items:
            if item['is_heading']:
                continue
            item_text = re.sub(r'\s+', ' ', item['text'].lower())
            if snippet_lower in item_text:
                match_idx = item['idx']
                break

        # Poging 2: prefix-match (eerste 60 tekens van snippet)
        # Bruikbaar als de citaat over een paragraafgrens loopt of door whitespace
        # of speciale tekens net iets afwijkt van p.text.
        if match_idx is None:
            prefix = snippet_lower[:60].strip()
            if len(prefix) >= 15:
                for item in section_items:
                    if item['is_heading']:
                        continue
                    item_text = re.sub(r'\s+', ' ', item['text'].lower())
                    if prefix in item_text:
                        match_idx = item['idx']
                        logger.debug(
                            "Prefix-match voor snippet '%s...' → para %d",
                            prefix[:30], item['idx']
                        )
                        break

        # Poging 3: woordoverlap-match (voor AI-citaten die enigszins afwijken van p.text)
        # Pak de significante woorden (≥5 tekens) uit het snippet en zoek de paragraaf
        # met de meeste overeenkomsten.
        # Regel: sectie-lokale matches hebben ABSOLUTE prioriteit over document-brede matches.
        # Alleen als er GEEN sectie-lokale match is (< 0.35), wordt document-breed geprobeerd.
        # Dit voorkomt dat een comment van sectie A in sectie B (bijv. Hoofd- en Deelvragen)
        # belandt omdat die sectie toevallig meer woorden deelt met het snippet.
        if match_idx is None:
            def _sig_words_snippet(text: str) -> List[str]:
                return [w for w in re.sub(r'[^\w]', ' ', text.lower()).split()
                        if len(w) >= 5 and not w.isdigit()]

            snip_words = _sig_words_snippet(snippet_lower)
            if len(snip_words) >= 3:

                def _best_overlap(items):
                    best_score, best_idx = 0, None
                    for item in items:
                        if item['is_heading']:
                            continue
                        item_words = set(_sig_words_snippet(item['text']))
                        score = sum(1 for w in snip_words if w in item_words)
                        norm_score = score / len(snip_words)
                        if norm_score > best_score and score >= 3:
                            best_score = norm_score
                            best_idx = item['idx']
                    return best_score, best_idx

                # Eerst sectie-lokaal
                sec_score, sec_idx = _best_overlap(section_items)
                if sec_idx is not None and sec_score >= 0.35:
                    match_idx = sec_idx
                    logger.debug(
                        "Sectie woord-overlap match (score=%.0f%%) voor snippet '%s...' → para %d",
                        sec_score * 100, snippet_lower[:30], sec_idx
                    )
                else:
                    # Fallback: document-breed (alleen als sectie-lokaal niets vond)
                    outside_items = [it for it in para_structure
                                     if it not in section_items and not _is_toc_line(it.get('text', ''))]
                    doc_score, doc_idx = _best_overlap(outside_items)
                    if doc_idx is not None and doc_score >= 0.35:
                        match_idx = doc_idx
                        logger.debug(
                            "Document-breed woord-overlap match (score=%.0f%%) voor snippet '%s...' → para %d",
                            doc_score * 100, snippet_lower[:30], doc_idx
                        )

        if match_idx is not None:
            matched = para_structure[match_idx]
            if not matched['is_heading']:
                # Snippet gevonden in een specifieke alinea → altijd direct plaatsen.
                # Niet doorsturen naar een sub-heading: de comment hoort bij de alinea zelf.
                return matched['idx'], f"paragraaf in '{section_name}'"

        # Titel-/kopfeedback: als de snippet een KOPTEKST aanwijst (en er geen
        # sectiecontext is), plaats de comment op die kop zelf i.p.v. bovenaan.
        if heading_idx is None:
            sl = re.sub(r'\s+', ' ', offending_snippet.lower().strip())
            for item in para_structure:
                if not item['is_heading'] or _is_toc_line(item['text']):
                    continue
                ht = re.sub(r'\s+', ' ', item['text'].lower()).strip()
                ht_clean = re.sub(r'^[\d.]+\s*', '', ht)
                if len(ht_clean) >= 6 and (ht_clean in sl or sl in ht):
                    return item['idx'], 'koptekst'

    return heading_idx if heading_idx is not None else 0, f"sectie '{section_name}'"


# ---------------------------------------------------------------------------
# Snippet-plaatsing: markers rond specifieke tekst in een paragraaf
# ---------------------------------------------------------------------------

def _place_markers_around_snippet(
    p_el,          # lxml w:p element
    id_str: str,   # comment id als string
    snippet: str,  # de tekst om te markeren
) -> bool:
    """
    Zoek `snippet` in de tekst van de paragraaf-runs en plaats
    commentRangeStart/End alleen rond de runs die de snippet bevatten.

    Returns True als de snippet is gevonden en markers zijn geplaatst.
    """
    # Bouw een mapping: [(run_element, run_text, start_pos, end_pos), ...]
    runs = p_el.findall(f'{{{W_NS}}}r')
    run_map: list = []  # (run_element, cumulative_start)
    full_text = ''
    for r in runs:
        t_el = r.find(f'{{{W_NS}}}t')
        if t_el is not None and t_el.text:
            run_map.append((r, len(full_text)))
            full_text += t_el.text
        else:
            # Runs zonder tekst (bijv. tab, break) overslaan maar positie bijhouden
            run_map.append((r, len(full_text)))

    if not full_text:
        return False

    # Zoek snippet (case-insensitive) in de paragraaftekst
    idx = full_text.lower().find(snippet.lower())
    if idx < 0:
        # Probeer met genormaliseerde spaties (Word kan extra spaties/tabs hebben)
        norm_full = re.sub(r'\s+', ' ', full_text)
        norm_snip = re.sub(r'\s+', ' ', snippet)
        idx = norm_full.lower().find(norm_snip.lower())
        if idx < 0:
            return False

    snip_start = idx
    snip_end   = idx + len(snippet)

    # Vind de eerste run die overlapt met snippet-start
    first_run = None
    last_run  = None
    for i, (r, cum_start) in enumerate(run_map):
        t_el = r.find(f'{{{W_NS}}}t')
        run_len = len(t_el.text) if (t_el is not None and t_el.text) else 0
        cum_end = cum_start + run_len
        if run_len == 0:
            continue
        # Overlapt deze run met het snippet-bereik?
        if cum_end > snip_start and cum_start < snip_end:
            if first_run is None:
                first_run = (i, r)
            last_run = (i, r)

    if first_run is None:
        return False

    # Plaats commentRangeStart VOOR de eerste overlappende run
    crs = etree.Element(f'{{{W_NS}}}commentRangeStart')
    crs.set(f'{{{W_NS}}}id', id_str)
    first_run[1].addprevious(crs)

    # Plaats commentRangeEnd NA de laatste overlappende run
    cre = etree.Element(f'{{{W_NS}}}commentRangeEnd')
    cre.set(f'{{{W_NS}}}id', id_str)
    last_run[1].addnext(cre)

    return True


# ---------------------------------------------------------------------------
# Lichte MARKERING (arcering / onderstreping) i.p.v. comments
# Voor taalfouten: de student ziet de plek, zonder zware comment-ballon.
# ---------------------------------------------------------------------------

_XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


def _ensure_rpr(run):
    rPr = run.find(f'{{{W_NS}}}rPr')
    if rPr is None:
        rPr = etree.Element(f'{{{W_NS}}}rPr')
        run.insert(0, rPr)
    return rPr


def _apply_mark_to_rpr(rPr, color: Optional[str], underline: bool):
    if color:
        for ex in rPr.findall(f'{{{W_NS}}}highlight'):
            rPr.remove(ex)
        hl = etree.SubElement(rPr, f'{{{W_NS}}}highlight')
        hl.set(f'{{{W_NS}}}val', color)
    if underline:
        for ex in rPr.findall(f'{{{W_NS}}}u'):
            rPr.remove(ex)
        u = etree.SubElement(rPr, f'{{{W_NS}}}u')
        u.set(f'{{{W_NS}}}val', 'single')


def _run_with_text(src_run, text: str, mark: bool, color: Optional[str], underline: bool):
    new = deepcopy(src_run)
    t = new.find(f'{{{W_NS}}}t')
    if t is None:
        t = etree.SubElement(new, f'{{{W_NS}}}t')
    t.text = text
    t.set(_XML_SPACE, 'preserve')
    if mark:
        _apply_mark_to_rpr(_ensure_rpr(new), color, underline)
    return new


def _highlight_snippet_in_para(p_el, snippet: str, color: Optional[str], underline: bool,
                               mark_text: str = '') -> bool:
    """Markeer (arcering/onderstreping) de runs die `snippet` bevatten; splitst
    runs zodat alleen de exacte tekst gemarkeerd wordt. Als `mark_text` is opgegeven
    en binnen de gevonden passage voorkomt, wordt ALLEEN dat deel gemarkeerd."""
    runs = p_el.findall(f'{{{W_NS}}}r')
    infos = []
    full = ''
    for r in runs:
        t_el = r.find(f'{{{W_NS}}}t')
        txt = t_el.text if (t_el is not None and t_el.text is not None) else ''
        infos.append((r, t_el, txt, len(full), len(full) + len(txt)))
        full += txt
    if not full:
        return False

    def _flex(text: str):
        toks = [re.escape(t) for t in text.split() if t]
        return re.compile(r'\s+'.join(toks), re.IGNORECASE) if toks else None

    pat = _flex(snippet.strip())
    m = pat.search(full) if pat else None
    if not m:
        head = snippet.strip()[:40]
        pat = _flex(head) if len(head) >= 8 else None
        m = pat.search(full) if pat else None
        if not m:
            return False
    s, e = m.start(), m.end()

    # Versmal tot ALLEEN het foutieve stukje, als dat binnen de passage te vinden is.
    if mark_text and mark_text.strip():
        mpat = _flex(mark_text.strip())
        mm = mpat.search(full, s, e) if mpat else None
        if mm is None and mpat is not None:      # ook net buiten de quote-grens toestaan
            mm = mpat.search(full)
        if mm is not None:
            s, e = mm.start(), mm.end()

    for (r, t_el, txt, rs, re_) in infos:
        if t_el is None or not txt or re_ <= s or rs >= e:
            continue
        os_ = max(s, rs) - rs
        oe_ = min(e, re_) - rs
        before, match, after = txt[:os_], txt[os_:oe_], txt[oe_:]
        if not before and not after:
            _apply_mark_to_rpr(_ensure_rpr(r), color, underline)
            continue
        pieces = []
        if before:
            pieces.append(_run_with_text(r, before, False, color, underline))
        pieces.append(_run_with_text(r, match, True, color, underline))
        if after:
            pieces.append(_run_with_text(r, after, False, color, underline))
        for nr in pieces:
            r.addprevious(nr)
        parent = r.getparent()
        if parent is not None:
            parent.remove(r)
    return True


def add_highlights(
    original_docx_path: str,
    highlight_items: List[Any],
    output_path: str,
    color: Optional[str] = 'yellow',
    underline: bool = False,
) -> Tuple[str, int]:
    """
    Markeer (geel/onderstreept) de aangegeven tekstfragmenten in het document,
    ZONDER comments. `highlight_items` is een lijst van snippets (str) of dicts
    met 'offending_snippet'. Returns (output_path, aantal_gemarkeerd).
    """
    snippets = []   # (locate_text, mark_text)
    for it in highlight_items:
        if isinstance(it, str):
            loc, mark = it, ''
        else:
            loc = it.get('offending_snippet') or it.get('quote') or ''
            mark = it.get('fout') or it.get('mark') or ''
        loc, mark = (loc or '').strip(), (mark or '').strip()
        # quote mag kort zijn als er een fout-deel is; anders min. 4 tekens om te plaatsen
        if len(loc) >= 4 or len(mark) >= 3:
            snippets.append((loc or mark, mark))

    if not snippets:
        if os.path.abspath(output_path) != os.path.abspath(original_docx_path):
            shutil.copy2(original_docx_path, output_path)
        return output_path, 0

    with zipfile.ZipFile(original_docx_path, 'r') as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    parser = etree.XMLParser(remove_blank_text=False, resolve_entities=False)
    tree = etree.fromstring(files['word/document.xml'], parser)
    body = tree.find(f'{{{W_NS}}}body')
    paras = body.findall(f'{{{W_NS}}}p') if body is not None else []

    placed = 0
    for snip, mark in snippets:
        norm = re.sub(r'\s+', ' ', snip.lower().strip())
        for p in paras:
            ptext = ''.join(t.text or '' for t in p.findall(f'.//{{{W_NS}}}t'))
            if not ptext or _is_toc_line(ptext):
                continue
            np = re.sub(r'\s+', ' ', ptext.lower())
            if norm in np or (len(norm) >= 15 and norm[:60] in np):
                if _highlight_snippet_in_para(p, snip, color, underline, mark_text=mark):
                    placed += 1
                break

    files['word/document.xml'] = etree.tostring(
        tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for n, d in files.items():
            zout.writestr(n, d)

    logger.info("%d tekstfragment(en) gemarkeerd -> %s", placed, output_path)
    return output_path, placed


# ---------------------------------------------------------------------------
# Comment markers toevoegen aan document.xml via lxml
# (GEEN python-docx save → originele bestanden blijven intact)
# ---------------------------------------------------------------------------

def _shortest_anchor(snippet: str, min_len: int = 20, max_len: int = 80) -> str:
    """
    Geeft de kortst mogelijke maar nog herkenbare ankertekst uit het snippet:
    - Eerste zin (eindigend op . ? ! gevolgd door spatie of einde), minimaal min_len tekens
    - Als er geen zinseinde is: eerste max_len tekens
    """
    if not snippet:
        return snippet
    s = snippet.strip()
    # Eerste zin die lang genoeg is
    m = re.search(r'^.{' + str(min_len) + r',}?[.?!](?=\s|$)', s)
    if m:
        return m.group(0)[:max_len]
    return s[:max_len]


def _add_markers_to_doc_xml(
    doc_xml_bytes: bytes,
    items_list: List[Tuple[int, Dict]],
) -> Tuple[bytes, List[Tuple[int, str, str]]]:
    """
    Parseer document.xml met lxml, voeg comment markers toe en serialiseer terug.
    Elk feedbackitem krijgt zijn eigen aparte comment, verankerd aan zijn eigen snippet.

    Gebruikt body.findall('{w}p') zodat de indices overeenkomen met
    python-docx doc.paragraphs (beide zijn directe kinderen van w:body).

    Returns:
        (gewijzigde_document_xml_bytes, comments_data)
        comments_data: lijst van (comment_id, tekst, author)
    """
    parser = etree.XMLParser(remove_blank_text=False, resolve_entities=False)
    doc_tree = etree.fromstring(doc_xml_bytes, parser)

    body = doc_tree.find(f'{{{W_NS}}}body')
    if body is None:
        raise ValueError("Geen w:body gevonden in document.xml")

    # Directe w:p kinderen van w:body – zelfde volgorde als doc.paragraphs
    body_paras = body.findall(f'{{{W_NS}}}p')

    comments_data: List[Tuple[int, str, str]] = []

    # Vind de hoogste bestaande w:id in het hele document (bookmarks, comments, etc.
    # delen dezelfde ID-ruimte in OOXML). Start onze comment-IDs daarboven
    # om ID-conflicten te voorkomen met bijv. Table of Contents bookmarks.
    nsmap_xpath = {'w': W_NS}
    all_ids = doc_tree.xpath('//*/@w:id', namespaces=nsmap_xpath)
    max_existing_id = max((int(v) for v in all_ids if v.isdigit()), default=-1)
    comment_id = max_existing_id + 1
    logger.debug("Hoogste bestaande w:id = %d, start comment-IDs bij %d",
                 max_existing_id, comment_id)

    # ── Verwijder ALLE bestaande comment markers uit document.xml ──
    # Het originele document kan al comments bevatten (met andere IDs).
    # Als we comments.xml overschrijven maar de oude markers in document.xml
    # laten staan, ontstaan orphaned references (commentRangeStart/End die
    # verwijzen naar IDs die niet meer in comments.xml bestaan).
    # Word detecteert dit als corruptie → opent het bestand als "Document 1"
    # (herstelde versie) → commentaren gaan verloren.
    _to_remove = []
    for el in doc_tree.iter():
        if el.tag in (f'{{{W_NS}}}commentRangeStart', f'{{{W_NS}}}commentRangeEnd'):
            _to_remove.append(el)
        elif el.tag == f'{{{W_NS}}}commentReference':
            # Verwijder de hele w:r parent (bevat alleen de commentReference + rStyle)
            r_el = el.getparent()
            if r_el is not None and r_el.tag == f'{{{W_NS}}}r':
                _to_remove.append(r_el)
    for el in _to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    logger.debug("%d bestaande comment marker(s) verwijderd uit document.xml", len(_to_remove))

    # Groepeer items op ANKERTEKST (genormaliseerd op 60 tekens):
    # - Meerdere items die DEZELFDE zin/tekst aanwijzen → één comment (samenvoegen)
    # - Items die VERSCHILLENDE zinnen aanwijzen → elk een eigen comment
    # De onderscheidende factor is de ankertekst, NIET de paragraaf.
    # Zo krijgen bijv. drie deelvragen elk hun eigen comment (elk een andere zin),
    # maar twee opmerkingen over dezelfde korte hoofdvraag-zin worden samengevoegd.

    # Stap A: bereken per item de anchor en normaliseer als groeperingssleutel
    from collections import OrderedDict
    keyed: List[Tuple[str, int, str, Dict]] = []  # (anchor_key, para_idx, anchor, fi)
    for para_idx, fi in items_list:
        raw = fi.get('offending_snippet', '')
        anchor = _shortest_anchor(raw) if raw else ''
        # Normaliseer: lowercase, witruimte samenvoegen, eerste 60 tekens
        key = re.sub(r'\s+', ' ', anchor.lower().strip())[:60] if anchor else f'__para_{para_idx}__'
        keyed.append((key, para_idx, anchor, fi))
        logger.debug(
            "GROEP-KEY: criteria=%s | para=%d | key=%r | anchor=%r",
            fi.get('criteria_name', '?'), para_idx, key, anchor[:50] if anchor else ''
        )

    # Stap B: groepeer op (anchor_key, para_idx) — zelfde zin op zelfde paragraaf = samen.
    # Extra: als twee items op dezelfde para_idx staan en hun anchor_key een prefix-overlap
    # hebben van ≥30 tekens, worden ze ook samengevoegd (vangt licht verschillende citaten op).
    groups: OrderedDict = OrderedDict()

    def _find_existing_group(anchor_key: str, para_idx: int) -> Optional[tuple]:
        """Zoek een bestaande groep op dezelfde para_idx met voldoende ankeroverlap."""
        prefix = anchor_key[:30]
        if len(prefix) < 15:
            return None
        for gk in groups:
            if gk[1] == para_idx and gk[0].startswith(prefix):
                return gk
        return None

    for anchor_key, para_idx, anchor, fi in keyed:
        existing = _find_existing_group(anchor_key, para_idx)
        if existing:
            groups[existing]['items'].append(fi)
        else:
            group_key = (anchor_key, para_idx)
            if group_key not in groups:
                groups[group_key] = {'para_idx': para_idx, 'anchor': anchor, 'items': []}
            groups[group_key]['items'].append(fi)

    logger.debug("TOTAAL GROEPEN: %d (uit %d items)", len(groups), len(keyed))

    # Stap C: één comment per groep
    for group_key, grp in groups.items():
        para_idx = grp['para_idx']
        anchor   = grp['anchor']
        group    = grp['items']

        if para_idx >= len(body_paras):
            continue

        p_el   = body_paras[para_idx]
        id_str = str(comment_id)

        placed = False
        if anchor and len(anchor.strip()) >= 5:
            placed = _place_markers_around_snippet(p_el, id_str, anchor.strip())

        if not placed:
            # Fallback: markers rond de hele paragraaf
            crs = etree.Element(f'{{{W_NS}}}commentRangeStart')
            crs.set(f'{{{W_NS}}}id', id_str)
            pPr = p_el.find(f'{{{W_NS}}}pPr')
            if pPr is not None:
                pPr.addnext(crs)
            else:
                p_el.insert(0, crs)
            cre = etree.Element(f'{{{W_NS}}}commentRangeEnd')
            cre.set(f'{{{W_NS}}}id', id_str)
            p_el.append(cre)

        # -- w:r met w:commentReference: altijd aan einde paragraaf --
        ref_run  = etree.SubElement(p_el, f'{{{W_NS}}}r')
        ref_rPr  = etree.SubElement(ref_run, f'{{{W_NS}}}rPr')
        ref_rSt  = etree.SubElement(ref_rPr, f'{{{W_NS}}}rStyle')
        ref_rSt.set(f'{{{W_NS}}}val', 'CommentReference')
        ref      = etree.SubElement(ref_run, f'{{{W_NS}}}commentReference')
        ref.set(f'{{{W_NS}}}id', id_str)

        # -- Commenttekst: items samenvoegen als meerdere, anders enkelvoudig --
        lines: List[str] = []
        for fi in group:
            msg = fi.get('message', '')
            sug = fi.get('suggestion', '')
            check_type = fi.get('check_type', '')
            has_snippet = bool((fi.get('offending_snippet') or '').strip())
            if msg:
                # Voeg "(AI - geen citaat)" toe voor LLM-feedback zonder verifieerbaar citaat
                if check_type in ('llm_review', 'holistic') and not has_snippet:
                    lines.append(f"{msg}\n(AI - geen citaat)")
                else:
                    lines.append(msg)
            if sug:
                lines.append(f"Suggestie: {sug}")

        comments_data.append((comment_id, '\n\n'.join(lines), 'DocuCheck (AI)'))
        comment_id += 1

    # Herserializeer ALLEEN document.xml (alle andere bestanden blijven ongewijzigd)
    new_xml = etree.tostring(
        doc_tree,
        xml_declaration=True,
        encoding='UTF-8',
        standalone=True,
    )
    return new_xml, comments_data


# ---------------------------------------------------------------------------
# comments.xml opbouwen met correcte OOXML-structuur
# ---------------------------------------------------------------------------

def _safe_text(text: str) -> str:
    """Verwijder tekens die Word's XML-parser niet aankan."""
    replacements = {
        '\u274c': '[FOUT]',
        '\u26a0': '[WAARSCH.]',
        '\ufe0f': '',
        '\U0001f4a1': '[TIP]',
    }
    result = []
    for ch in text:
        if ch in replacements:
            result.append(replacements[ch])
        elif ord(ch) > 0xFFFF:
            result.append('')
        else:
            result.append(ch)
    return ''.join(result)


def _build_comments_xml(comments_data: List[Tuple[int, str, str]]) -> bytes:
    """
    Bouw comments.xml met correcte OOXML-structuur:
    - w:pStyle val="CommentText" in de paragraaf-eigenschappen
    - w:annotationRef in de eerste run (vereist door Word voor weergave)
    - Elke nieuwe regel wordt een aparte w:p binnen de comment
    """
    now    = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')
    nsmap  = {
        'w': W_NS,
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    root = etree.Element(f'{{{W_NS}}}comments', nsmap=nsmap)

    for cid, ctext, author in comments_data:
        comment_el = etree.SubElement(root, f'{{{W_NS}}}comment')
        comment_el.set(f'{{{W_NS}}}id',       str(cid))
        comment_el.set(f'{{{W_NS}}}author',   author)
        comment_el.set(f'{{{W_NS}}}date',     now)
        comment_el.set(f'{{{W_NS}}}initials', 'FT')

        lines = _safe_text(ctext).split('\n')

        for line_idx, line in enumerate(lines):
            cp     = etree.SubElement(comment_el, f'{{{W_NS}}}p')
            pPr    = etree.SubElement(cp,  f'{{{W_NS}}}pPr')
            pStyle = etree.SubElement(pPr, f'{{{W_NS}}}pStyle')
            pStyle.set(f'{{{W_NS}}}val', 'CommentText')

            # Eerste paragraaf: annotationRef (ankerpunt voor Word)
            if line_idx == 0:
                r_ref  = etree.SubElement(cp,    f'{{{W_NS}}}r')
                rPr    = etree.SubElement(r_ref, f'{{{W_NS}}}rPr')
                rStyle = etree.SubElement(rPr,   f'{{{W_NS}}}rStyle')
                rStyle.set(f'{{{W_NS}}}val', 'CommentReference')
                etree.SubElement(r_ref, f'{{{W_NS}}}annotationRef')

            r_text = etree.SubElement(cp, f'{{{W_NS}}}r')
            t      = etree.SubElement(r_text, f'{{{W_NS}}}t')
            t.text = line.strip() if line.strip() else ' '
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


# ---------------------------------------------------------------------------
# Publieke hoofdfunctie
# ---------------------------------------------------------------------------

def add_inline_comments(
    original_docx_path: str,
    feedback_items: List[Dict[str, Any]],
    recognized_sections: List[Dict[str, Any]],
    output_path: Optional[str] = None,
) -> str:
    """
    Voeg inline Word comments toe bij criteria-afwijkingen.

    Werkwijze (vermijdt 'unreadable content'):
    1. Alle originele bestanden worden ONGEWIJZIGD uit het originele docx gehaald.
    2. Alleen word/document.xml wordt minimaal aangepast (comment markers via lxml).
    3. word/comments.xml wordt toegevoegd.
    4. [Content_Types].xml en _rels worden bijgewerkt via string-insertie.
    5. Alles wordt in één keer naar het uitvoerbestand geschreven.
    """
    if output_path is None:
        base, ext   = os.path.splitext(original_docx_path)
        output_path = f"{base}_gecommentarieerd{ext}"

    # Filter: 'violation', 'warning' en 'info' krijgen een Word-comment.
    # Uitgesloten:
    #   - 'ok' items (geslaagd)
    #   - tool-fouten (confidence == 0.0): dit zijn interne foutmeldingen over
    #     mislukte LLM-calls, geen pedagogische feedback voor de student.
    active = [
        fi for fi in feedback_items
        if fi.get('status') in ('violation', 'warning', 'info', 'error')
        and fi.get('confidence', 1.0) > 0.0
    ]


    if not active:
        logger.info("Geen afwijkingen - document ongewijzigd gekopieerd.")
        shutil.copy2(original_docx_path, output_path)
        return output_path

    # Stap 1: lees paragraafstructuur via python-docx (alleen lezen, geen save)
    doc            = Document(original_docx_path)
    para_structure = _build_para_structure(doc)

    # Bouw aanvullende mapping sectienaam → originele heading-tekst
    # (opgeslagen tijdens sectie-herkenning voor alias/fuzzy-matches)
    section_heading_texts: Dict[str, str] = {
        s['name']: s['heading_text']
        for s in recognized_sections
        if s.get('heading_text') and s.get('name')
    }

    # Stap 2: bepaal per feedbackitem de doelplaats (paragraaf-index)
    # Elk item krijgt zijn eigen comment → geen groepering meer.
    # Dedupliceer op (criteria_id, snippet): hetzelfde criterium + snippet via
    # zowel parent- als sub-sectie geeft anders twee identieke comments.
    items_list: List[Tuple[int, Dict]] = []
    seen_snippet_keys: set = set()
    for fi in active:
        snippet = fi.get('offending_snippet')
        if snippet:
            crit_id   = fi.get('criteria_id') or fi.get('criteria_name', '')
            dedup_key = (crit_id, snippet.strip()[:80])
            if dedup_key in seen_snippet_keys:
                logger.debug("Dedup: %s met snippet '%s...' al aanwezig — overgeslagen",
                             crit_id, snippet[:30])
                continue
            seen_snippet_keys.add(dedup_key)

        section_name = fi.get('section_name', '')
        para_idx, loc = _find_target_paragraph_idx(
            para_structure, section_name, snippet, section_heading_texts
        )
        logger.debug("Route: %s → para %d (%s)",
                     fi.get('criteria_name') or fi.get('criterion_name', '?'),
                     para_idx, loc)
        items_list.append((para_idx, fi))

    if not items_list:
        logger.warning("Geen doelrefs gevonden voor feedback items.")
        shutil.copy2(original_docx_path, output_path)
        return output_path

    # Stap 3: lees ALLE bestanden ongewijzigd uit het originele docx
    with zipfile.ZipFile(original_docx_path, 'r') as zin:
        all_files = {name: zin.read(name) for name in zin.namelist()}

    # Stap 3b: verwijder modern-comment bestanden + hun rels/content-type entries.
    #
    # REDEN: het AANWEZIG ZIJN van de relationship-entries voor commentsExtended.xml,
    # commentsIds.xml en commentsExtensible.xml dwingt Word 365 om de *moderne*
    # threaded-commentaar-renderer te gebruiken (i.p.v. de legacy renderer).
    # Die moderne renderer vereist geldige w16cid:commentId-entries (paraId-koppelingen)
    # in commentsIds.xml om de comment-tekst te tonen. Zonder die koppelingen toont Word
    # lege ballonnen — ook als comments.xml de juiste tekst bevat.
    #
    # Oplossing: verwijder de drie bestanden én hun rels-entries volledig.
    # Zonder die entries gebruikt Word automatisch de legacy renderer,
    # die comments.xml rechtstreeks leest en de tekst correct toont.
    _EXT_TARGETS = ['commentsExtended.xml', 'commentsIds.xml', 'commentsExtensible.xml']

    # 3b-1: verwijder de bestanden zelf
    for target in _EXT_TARGETS:
        fpath = f'word/{target}'
        if fpath in all_files:
            del all_files[fpath]
            logger.debug("%s verwijderd uit ZIP", fpath)

    # 3b-2: verwijder hun Relationship-entries uit document.xml.rels
    _rels_key = 'word/_rels/document.xml.rels'
    if _rels_key in all_files:
        _rels_str = all_files[_rels_key].decode('utf-8', errors='replace')
        for target in _EXT_TARGETS:
            _rels_str = re.sub(
                r'<Relationship\b[^>]*\bTarget="' + re.escape(target) + r'"[^>]*/>\s*',
                '',
                _rels_str,
            )
        all_files[_rels_key] = _rels_str.encode('utf-8')

    # 3b-3: verwijder hun Override-entries uit [Content_Types].xml
    _ct_key = '[Content_Types].xml'
    if _ct_key in all_files:
        _ct_str = all_files[_ct_key].decode('utf-8', errors='replace')
        for target in _EXT_TARGETS:
            _ct_str = re.sub(
                r'<Override\b[^>]*\bPartName="[^"]*' + re.escape(target) + r'"[^>]*/>\s*',
                '',
                _ct_str,
            )
        all_files[_ct_key] = _ct_str.encode('utf-8')

    # Stap 4: pas document.xml aan via lxml (alle andere bestanden blijven intact)
    new_doc_xml, comments_data = _add_markers_to_doc_xml(
        all_files['word/document.xml'], items_list
    )
    all_files['word/document.xml'] = new_doc_xml

    # Stap 5: voeg comments.xml toe
    comments_xml = _build_comments_xml(comments_data)
    all_files['word/comments.xml'] = comments_xml
    logger.debug("%d comment(s) aangemaakt", len(comments_data))

    # Stap 6: update [Content_Types].xml via string-insertie (geen XML-herserializatie)
    ct_str = all_files['[Content_Types].xml'].decode('utf-8', errors='replace')
    if 'comments.xml' not in ct_str:
        ct_str = ct_str.replace(
            '</Types>',
            f'<Override PartName="/word/comments.xml" ContentType="{COMMENTS_CT}"/>'
            '</Types>',
        )
        all_files['[Content_Types].xml'] = ct_str.encode('utf-8')

    # Stap 7: update word/_rels/document.xml.rels via string-insertie
    rels_key = 'word/_rels/document.xml.rels'
    if rels_key in all_files:
        rels_str = all_files[rels_key].decode('utf-8', errors='replace')
        if 'comments.xml' not in rels_str:
            nums    = [int(m) for m in re.findall(r'Id="rId(\d+)"', rels_str)]
            new_num = max(nums, default=0) + 1
            rels_str = rels_str.replace(
                '</Relationships>',
                f'<Relationship Id="rId{new_num}" '
                f'Type="{COMMENTS_REL_TYPE}" '
                f'Target="comments.xml"/>'
                '</Relationships>',
            )
            all_files[rels_key] = rels_str.encode('utf-8')

    # Stap 8: schrijf definitief uitvoerbestand
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in all_files.items():
            zout.writestr(name, data)

    logger.info(
        f"{len(comments_data)} inline comment(s) toegevoegd -> {output_path}"
    )
    return output_path
