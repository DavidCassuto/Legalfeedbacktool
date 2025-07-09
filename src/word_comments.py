"""
Word Comments Module - Geavanceerde Word comments implementatie
Gebruikt python-docx en XML manipulatie voor betrouwbare comment toevoeging bij specifieke secties.
"""

import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt
from docx.text.paragraph import Paragraph
import re

logger = logging.getLogger(__name__)

class AdvancedWordCommentManager:
    """Beheert geavanceerde Word comments via python-docx en XML manipulatie."""
    
    def __init__(self):
        """Initialiseer de comment manager."""
        self.comment_id_counter = 1
        
    def add_section_specific_comments(self, doc_path: str, feedback_data: Dict[str, Any], 
                                    sections_data: List[Dict], output_path: Optional[str] = None) -> str:
        """
        Voeg Word comments toe aan specifieke secties in het document.
        
        Args:
            doc_path: Pad naar het originele Word document
            feedback_data: Dictionary met feedback per sectie
            sections_data: Lijst van herkende secties met karakterposities
            output_path: Pad voor het geëxporteerde document
            
        Returns:
            Pad naar het geëxporteerde document met comments
        """
        if output_path is None:
            base_name = os.path.splitext(doc_path)[0]
            output_path = f"{base_name}_met_sectie_comments.docx"
            
        logger.info("Start toevoegen sectie-specifieke Word comments...")
        
        try:
            # Kopieer het originele document
            doc = Document(doc_path)
            
            # Groepeer feedback per sectie
            feedback_per_section = self._group_feedback_by_section(feedback_data, sections_data)
            
            # Voeg comments toe voor elke sectie
            for section_info in sections_data:
                section_id = section_info.get('db_id')
                if section_id and section_id in feedback_per_section:
                    section_feedback = feedback_per_section[section_id]
                    self._add_comments_to_section(doc, section_info, section_feedback)
            
            # Sla het document op
            doc.save(output_path)
            logger.info(f"Sectie-specifieke Word comments toegevoegd: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Fout bij toevoegen sectie-specifieke comments: {e}")
            raise
    
    def _group_feedback_by_section(self, feedback_data: Dict[str, Any], 
                                  sections_data: List[Dict]) -> Dict[int, List[Dict]]:
        """Groepeer feedback per sectie ID."""
        feedback_per_section = {}
        
        for feedback_item in feedback_data.get('feedback', []):
            section_id = feedback_item.get('section_id')
            if section_id:
                if section_id not in feedback_per_section:
                    feedback_per_section[section_id] = []
                feedback_per_section[section_id].append(feedback_item)
        
        return feedback_per_section
    
    def _add_comments_to_section(self, doc: Document, section_info: Dict, 
                                section_feedback: List[Dict]) -> None:
        """Voeg comments toe aan een specifieke sectie."""
        section_name = section_info.get('name', 'Onbekende sectie')
        logger.info(f"Voeg comments toe aan sectie: {section_name}")
        
        # Zoek de paragraaf die het dichtst bij het begin van de sectie ligt
        target_paragraph = self._find_section_start_paragraph(doc, section_info)
        
        if target_paragraph is None:
            logger.warning(f"Kon geen geschikte paragraaf vinden voor sectie: {section_name}")
            return
        
        # Voeg een comment toe met alle feedback voor deze sectie
        comment_text = self._format_section_feedback(section_name, section_feedback)
        self._add_comment_after_paragraph(doc, target_paragraph, comment_text)
    
    def _find_section_start_paragraph(self, doc: Document, section_info: Dict) -> Optional[Paragraph]:
        """Vind de paragraaf die het dichtst bij het begin van de sectie ligt."""
        section_start_char = section_info.get('start_char', 0)
        section_end_char = section_info.get('end_char', 0)
        
        if section_start_char == 0 and section_end_char == 0:
            # Fallback: zoek naar de sectienaam in paragrafen
            return self._find_paragraph_by_text(doc, section_info.get('name', ''))
        
        # Zoek de paragraaf die het dichtst bij section_start_char ligt
        current_char_pos = 0
        closest_paragraph = None
        min_distance = float('inf')
        
        for para in doc.paragraphs:
            para_text = para.text
            para_start = current_char_pos
            para_end = current_char_pos + len(para_text)
            
            # Bereken afstand tot sectie begin
            distance = abs(para_start - section_start_char)
            if distance < min_distance:
                min_distance = distance
                closest_paragraph = para
            
            current_char_pos = para_end + 1  # +1 voor newline
        
        return closest_paragraph
    
    def _find_paragraph_by_text(self, doc: Document, search_text: str) -> Optional[Paragraph]:
        """Vind een paragraaf die de zoektekst bevat."""
        search_text_lower = search_text.lower()
        
        for para in doc.paragraphs:
            if search_text_lower in para.text.lower():
                return para
        
        return None
    
    def _format_section_feedback(self, section_name: str, feedback_list: List[Dict]) -> str:
        """Format feedback voor een sectie als comment tekst."""
        if not feedback_list:
            return f"Geen feedback voor sectie: {section_name}"
        
        comment_lines = [f"Feedback voor sectie: {section_name}"]
        comment_lines.append("=" * 50)
        
        for i, feedback in enumerate(feedback_list, 1):
            criteria_name = feedback.get('criteria_name', 'Onbekend criterium')
            status = feedback.get('status', 'UNKNOWN')
            message = feedback.get('message', 'Geen bericht')
            
            comment_lines.append(f"{i}. {criteria_name} ({status})")
            comment_lines.append(f"   {message}")
            comment_lines.append("")
        
        return "\n".join(comment_lines)
    
    def _add_comment_after_paragraph(self, doc: Document, target_paragraph: Paragraph, 
                                   comment_text: str) -> None:
        """Voeg een comment toe na een specifieke paragraaf."""
        try:
            # Maak een nieuwe paragraaf voor de comment
            new_paragraph = doc.add_paragraph()
            
            # Format de comment met kleur en stijl
            self._format_comment_paragraph(new_paragraph, comment_text)
            
            # Verplaats de paragraaf naar de juiste positie via XML manipulatie
            self._insert_paragraph_after(doc, target_paragraph, new_paragraph)
            
            logger.info(f"Comment toegevoegd na paragraaf: {target_paragraph.text[:50]}...")
            
        except Exception as e:
            logger.error(f"Fout bij toevoegen comment: {e}")
    
    def _format_comment_paragraph(self, paragraph: Paragraph, comment_text: str) -> None:
        """Format een paragraaf als comment met kleur en stijl."""
        # Definieer stijlen
        blue_color = RGBColor(0x00, 0x00, 0xFF)  # Blauw
        font_size_pt = Pt(10)
        
        # Voeg de comment tekst toe met formatting
        run = paragraph.add_run(comment_text)
        run.font.color.rgb = blue_color
        run.font.bold = True
        run.font.size = font_size_pt
        run.font.italic = False
    
    def _insert_paragraph_after(self, doc: Document, target_paragraph: Paragraph, 
                              new_paragraph: Paragraph) -> None:
        """Voeg een paragraaf toe na een andere paragraaf via XML manipulatie."""
        try:
            # Haal de XML elementen op
            target_element = target_paragraph._element
            new_element = new_paragraph._element
            
            # Haal de parent op (meestal de body)
            parent = target_element.getparent()
            
            # Vind de index van de target paragraaf
            index = parent.index(target_element)
            
            # Voeg de nieuwe paragraaf toe na de target paragraaf
            parent.insert(index + 1, new_element)
            
        except Exception as e:
            logger.error(f"Fout bij XML manipulatie: {e}")
            # Fallback: voeg toe aan het einde van het document
            pass

# Backward compatibility
class WordCommentManager(AdvancedWordCommentManager):
    """Backward compatibility wrapper."""
    
    def add_real_comments_to_document(self, doc_path: str, feedback_data: Dict[str, Any], 
                                    output_path: Optional[str] = None) -> str:
        """Backward compatibility method."""
        # Voor backward compatibility, gebruik een eenvoudige sectie mapping
        sections_data = [{'db_id': 1, 'name': 'Document', 'start_char': 0, 'end_char': 0}]
        return self.add_section_specific_comments(doc_path, feedback_data, sections_data, output_path) 