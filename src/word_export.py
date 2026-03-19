"""
Word document export module.
Voegt feedback toe aan het einde van elke sectie in het Word document.
"""

import os
import copy
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import logging

logger = logging.getLogger(__name__)

class WordFeedbackExporter:
    """Exporteert feedback naar Word document aan het einde van elke sectie."""
    
    def __init__(self):
        """Initialiseer de Word export module."""
        pass
    
    def add_feedback_to_document(self, original_file_path: str, feedback_data: Dict[str, Any], 
                                output_file_path: Optional[str] = None) -> str:
        """
        Voeg feedback toe aan het einde van elke sectie in het Word document.
        
        Args:
            original_file_path: Pad naar het originele Word document
            feedback_data: Dictionary met feedback data
            output_file_path: Pad voor het output document (optioneel)
            
        Returns:
            Pad naar het aangepaste document
        """
        try:
            # Open het originele document
            doc = Document(original_file_path)
            
            # Voeg feedback toe per sectie
            sections = feedback_data.get('sections', [])
            feedback_items = feedback_data.get('feedback_items', [])
            
            # Groepeer feedback per sectie
            feedback_by_section = {}
            for item in feedback_items:
                section_name = item.get('section_name', 'Algemeen')
                if section_name not in feedback_by_section:
                    feedback_by_section[section_name] = []
                feedback_by_section[section_name].append(item)
            
            # Voeg feedback toe aan het einde van elke sectie
            for section in sections:
                if section.get('found', False):
                    section_name = section.get('name', '')
                    section_feedback = feedback_by_section.get(section_name, [])
                    if section_feedback:
                        self._add_section_feedback(doc, section, section_feedback)
            
            # Voeg algemene feedback toe aan het einde van het document
            general_feedback = feedback_by_section.get('Algemeen', [])
            if general_feedback:
                self._add_general_feedback(doc, general_feedback)
            
            # Bepaal output pad
            if not output_file_path:
                base_name = os.path.splitext(original_file_path)[0]
                output_file_path = f"{base_name}_feedback.docx"
            
            # Sla het aangepaste document op
            doc.save(output_file_path)
            logger.info(f"Feedback toegevoegd aan document: {output_file_path}")
            
            return output_file_path
            
        except Exception as e:
            logger.error(f"Fout bij toevoegen feedback aan Word document: {e}")
            raise
    
    def _add_section_feedback(self, doc: Document, section: Dict[str, Any], feedback_items: List[Dict[str, Any]]):
        """Voeg feedback toe aan het einde van een sectie."""
        section_name = section.get('name', 'Onbekende sectie')
        
        # Zoek de sectie heading in het document
        section_heading = self._find_section_heading(doc, section_name)
        
        if section_heading:
            # Voeg feedback toe na de sectie heading
            feedback_text = self._format_section_feedback(section_name, feedback_items)
            self._add_feedback_paragraph(doc, section_heading, feedback_text)
        else:
            # Als sectie niet gevonden, voeg toe aan einde van document
            feedback_text = self._format_section_feedback(section_name, feedback_items)
            self._add_feedback_paragraph(doc, None, feedback_text)
    
    def _add_general_feedback(self, doc: Document, feedback_items: List[Dict[str, Any]]):
        """Voeg algemene feedback toe aan het einde van het document."""
        feedback_text = self._format_general_feedback(feedback_items)
        self._add_feedback_paragraph(doc, None, feedback_text)
    
    def _find_section_heading(self, doc: Document, section_name: str):
        """Zoek een sectie heading in het document."""
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                heading_text = paragraph.text.lower()
                if section_name.lower() in heading_text:
                    return paragraph
        return None
    
    def _add_feedback_paragraph(self, doc: Document, after_paragraph, feedback_text: str):
        """Voeg een feedback paragraaf toe na een specifieke paragraaf of aan het einde."""
        try:
            if after_paragraph:
                # Voeg feedback toe na de sectie heading
                feedback_para = doc.add_paragraph()
                feedback_para._element.addprevious(after_paragraph._element)
            else:
                # Voeg feedback toe aan het einde van het document
                feedback_para = doc.add_paragraph()
            
            # Voeg feedback tekst toe met formatting
            feedback_run = feedback_para.add_run(feedback_text)
            feedback_run.font.color.rgb = RGBColor(128, 128, 128)  # Grijs
            feedback_run.italic = True
            feedback_run.font.size = Inches(0.1)  # Kleine tekst
            
        except Exception as e:
            logger.warning(f"Kon feedback paragraaf niet toevoegen: {e}")
            # Fallback: voeg feedback toe als gewone tekst
            if after_paragraph:
                feedback_para = doc.add_paragraph()
                feedback_para._element.addprevious(after_paragraph._element)
            else:
                feedback_para = doc.add_paragraph()
            feedback_para.add_run(feedback_text)
    
    def _format_section_feedback(self, section_name: str, feedback_items: List[Dict[str, Any]]) -> str:
        """Format feedback items voor een sectie."""
        if not feedback_items:
            return f"Geen feedback voor sectie '{section_name}'"
        
        feedback_lines = [f"Feedback voor sectie '{section_name}':"]
        
        for item in feedback_items:
            status = item.get('status', 'unknown')
            message = item.get('message', 'Geen bericht')
            suggestion = item.get('suggestion', '')
            criterion_name = item.get('criterion_name', 'Onbekend criterium')
            
            # Voeg status icon toe
            if status == 'error' or status == 'violation':
                feedback_lines.append(f"❌ {criterion_name}: {message}")
            elif status == 'warning':
                feedback_lines.append(f"⚠️ {criterion_name}: {message}")
            elif status == 'ok':
                feedback_lines.append(f"✅ {criterion_name}: {message}")
            else:
                feedback_lines.append(f"ℹ️ {criterion_name}: {message}")
            
            if suggestion:
                feedback_lines.append(f"   💡 Suggestie: {suggestion}")
        
        return "\n".join(feedback_lines)
    
    def _format_general_feedback(self, feedback_items: List[Dict[str, Any]]) -> str:
        """Format algemene feedback."""
        if not feedback_items:
            return "Geen algemene feedback"
        
        feedback_lines = ["Algemene feedback:"]
        
        for item in feedback_items:
            status = item.get('status', 'unknown')
            message = item.get('message', 'Geen bericht')
            suggestion = item.get('suggestion', '')
            criterion_name = item.get('criterion_name', 'Onbekend criterium')
            
            # Voeg status icon toe
            if status == 'error' or status == 'violation':
                feedback_lines.append(f"❌ {criterion_name}: {message}")
            elif status == 'warning':
                feedback_lines.append(f"⚠️ {criterion_name}: {message}")
            elif status == 'ok':
                feedback_lines.append(f"✅ {criterion_name}: {message}")
            else:
                feedback_lines.append(f"ℹ️ {criterion_name}: {message}")
            
            if suggestion:
                feedback_lines.append(f"   💡 Suggestie: {suggestion}")
        
        return "\n".join(feedback_lines)
    
    def create_feedback_summary_document(self, feedback_data: Dict[str, Any], 
                                       output_file_path: str) -> str:
        """
        Maak een apart document met een samenvatting van alle feedback.
        
        Args:
            feedback_data: Dictionary met feedback data
            output_file_path: Pad voor het output document
            
        Returns:
            Pad naar het samenvatting document
        """
        try:
            # Maak een nieuw document
            doc = Document()
            
            # Voeg titel toe
            title = doc.add_heading('Feedback Samenvatting', 0)
            
            # Voeg feedback secties toe
            self._add_feedback_section(doc, feedback_data)
            
            # Sla het document op
            doc.save(output_file_path)
            logger.info(f"Feedback samenvatting opgeslagen: {output_file_path}")
            
            return output_file_path
            
        except Exception as e:
            logger.error(f"Fout bij maken feedback samenvatting: {e}")
            raise
    
    def _add_feedback_section(self, doc: Document, feedback_data: Dict[str, Any]):
        """Voeg feedback secties toe aan het samenvatting document."""
        sections = feedback_data.get('sections', [])
        feedback_items = feedback_data.get('feedback_items', [])
        
        # Groepeer feedback per sectie
        feedback_by_section = {}
        for item in feedback_items:
            section_name = item.get('section_name', 'Algemeen')
            if section_name not in feedback_by_section:
                feedback_by_section[section_name] = []
            feedback_by_section[section_name].append(item)
        
        # Voeg sectie feedback toe
        for section in sections:
            if section.get('found', False):
                section_name = section.get('name', '')
                section_feedback = feedback_by_section.get(section_name, [])
                if section_feedback:
                    self._add_section_info(doc, section)
                    self._add_feedback_items(doc, section_feedback)
        
        # Voeg algemene feedback toe
        general_feedback = feedback_by_section.get('Algemeen', [])
        if general_feedback:
            doc.add_heading('Algemene Feedback', level=1)
            self._add_feedback_items(doc, general_feedback)
    
    def _add_feedback_items(self, doc: Document, feedback_items: List[Dict[str, Any]]):
        """Voeg feedback items toe aan het document."""
        for item in feedback_items:
            status = item.get('status', 'unknown')
            message = item.get('message', 'Geen bericht')
            suggestion = item.get('suggestion', '')
            criterion_name = item.get('criterion_name', 'Onbekend criterium')
            
            # Bepaal kleur op basis van status
            if status == 'error' or status == 'violation':
                color = 'red'
                icon = '❌'
            elif status == 'warning':
                color = 'orange'
                icon = '⚠️'
            elif status == 'ok':
                color = 'green'
                icon = '✅'
            else:
                color = 'blue'
                icon = 'ℹ️'
            
            # Voeg feedback item toe
            self._add_feedback_item(doc, item, color, icon)
    
    def _add_feedback_item(self, doc: Document, item: Dict[str, Any], color: str, icon: str):
        """Voeg een individueel feedback item toe."""
        status = item.get('status', 'unknown')
        message = item.get('message', 'Geen bericht')
        suggestion = item.get('suggestion', '')
        criterion_name = item.get('criterion_name', 'Onbekend criterium')
        
        # Voeg criterium naam toe
        para = doc.add_paragraph()
        run = para.add_run(f"{icon} {criterion_name}")
        run.bold = True
        
        # Voeg bericht toe
        if message:
            doc.add_paragraph(f"Bericht: {message}")
        
        # Voeg suggestie toe
        if suggestion:
            doc.add_paragraph(f"Suggestie: {suggestion}")
        
        # Voeg lege regel toe
        doc.add_paragraph()
    
    def _add_section_info(self, doc: Document, section: Dict[str, Any]):
        """Voeg sectie informatie toe."""
        section_name = section.get('name', 'Onbekende sectie')
        doc.add_heading(f'Sectie: {section_name}', level=1) 