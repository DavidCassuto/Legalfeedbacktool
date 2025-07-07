# test_parsing.py
import os
from src.analysis.document_parsing import extract_document_content

# Zorg dat je een test.docx bestand hebt in je ProjectFT/instance/uploads/ map
# Maak anders een dummy .docx bestand
test_docx_path = os.path.join('instance', 'uploads', 'test_document.docx') 

# Dummy test document aanmaken als het niet bestaat
if not os.path.exists(test_docx_path):
    from docx import Document
    doc = Document()
    doc.add_heading('Dit is een Kop 1', level=1)
    doc.add_paragraph('Dit is een paragraaf tekst.')
    doc.add_heading('Dit is een Kop 2', level=2)
    doc.add_paragraph('Dit is een andere paragraaf.')
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Cel A'
    table.cell(1, 1).text = 'Cel B'
    doc.save(test_docx_path)
    print(f"Dummy document gemaakt op: {test_docx_path}")

try:
    content = extract_document_content(test_docx_path)
    print("\n--- Volledige tekst ---")
    print(content['full_text'])
    print("\n--- Paragrafen ---")
    for p in content['paragraphs']:
        print(p)
    print("\n--- Headings ---")
    for h in content['headings']:
        print(h)
    print("\n--- Tabellen ---")
    for t in content['tables']:
        print(t)
except Exception as e:
    print(f"Fout tijdens het parsen: {e}")